#!/usr/bin/env python3
import os, sys, django
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'projest.settings')
django.setup()

from django.conf import settings
from docx import Document
from docx.oxml.ns import qn

TEMPLATE = os.path.join(settings.BASE_DIR, "Low Frequency_template.docx")
OUTPUT   = os.path.join(settings.BASE_DIR, "test_low_frequency_output.docx")

# --- replacer: รวมทุก run ในย่อหน้าแล้วแทนทีเดียว (กันโดน split runs) ---
def replace_in_paragraph_strict(paragraph, repl: dict):
    runs = paragraph.runs
    if not runs: return
    full = "".join(r.text for r in runs)
    new  = full
    for k, v in repl.items():
        if k in new:
            new = new.replace(k, str(v))
    if new == full: return
    for i in range(len(runs)-1, -1, -1):
        r = runs[i]
        r._element.getparent().remove(r._element)
    paragraph.add_run(new)

def replace_in_doc(doc, repl: dict):
    for p in doc.paragraphs:
        replace_in_paragraph_strict(p, repl)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph_strict(p, repl)

# --- XPath ครอบคลุม w:t ทั้งไฟล์ (ดึงข้อความที่ python-docx ไม่ expose เช่นใน text box/shape) ---
NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

def _replace_all_text_nodes(part, repl: dict):
    """
    แทนค่าทุก w:t ภายใน part (รองรับ text box/shape, header/footer)
    โดยไม่ใช้ namespaces= ซึ่ง python-docx รุ่นนี้ไม่รองรับ
    """
    if not hasattr(part, "element"):
        return
    w_t = qn('w:t')  # -> '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
    for t in part.element.iter(w_t):
        txt = t.text or ""
        new = txt
        for k, v in repl.items():
            if k in new:
                new = new.replace(k, str(v))
        if new != txt:
            t.text = new

def replace_everywhere(doc, repl: dict):
    _replace_all_text_nodes(doc, repl)                 # main document
    for sec in doc.sections:                           # headers/footers
        if getattr(sec, "header", None):
            _replace_all_text_nodes(sec.header.part, repl)
        if getattr(sec, "footer", None):
            _replace_all_text_nodes(sec.footer.part, repl)

# --- helper ระบุ placeholder ที่ยังเหลือ (debug) ---
def find_left_placeholders(doc):
    leftovers = []
    def grab(txt):
        if "{{" in txt and "}}" in txt:
            leftovers.append(txt)
    for p in doc.paragraphs: grab(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                grab(cell.text)
    return leftovers

def main():
    if not os.path.exists(TEMPLATE):
        print("❌ ไม่พบ template:", TEMPLATE); return
    doc = Document(TEMPLATE)

    # ✅ ดึงข้อมูลจริงจากฐานข้อมูล
    try:
        from calibrate.models import LowFrequencyCalibration
        from organize.models import Organize
        
        # หาข้อมูล Low Frequency calibration แรก
        calibration = LowFrequencyCalibration.objects.first()
        if not calibration:
            print("❌ ไม่พบข้อมูล Low Frequency calibration ในฐานข้อมูล")
            return
        
        print(f"✅ พบข้อมูล Low Frequency calibration: ID {calibration.pk}")
        
        # ข้อมูลเครื่องมือ
        machine = calibration.machine
        standard = calibration.std_id
        
        # ดึงข้อมูลหน่วยงาน
        try:
            organization = Organize.objects.filter(is_main_unit=True).first()
            if not organization:
                organization = Organize.objects.first()
        except:
            organization = None
        
        print(f"  - เครื่องมือ: {machine.name}")
        print(f"  - หมายเลข: {machine.serial_number}")
        print(f"  - ช่วงการวัด: {calibration.measurement_range}")
        print(f"  - วันที่สอบเทียบ: {calibration.date_calibration}")
        
        # สร้าง replacements dictionary จากข้อมูลจริง
        ctx = {
            # ข้อมูลเครื่องมือ
            "{{MODEL}}": machine.model if machine.model else '-',
            "{{MANUFACTURER}}": str(machine.manufacture) if machine.manufacture else '-',
            "{{DESCRIPTION}}": machine.name if machine.name else '-',
            "{{SERIAL_NUMBER}}": machine.serial_number if machine.serial_number else '-',
            "{{RANGE}}": machine.range if machine.range else '-',
            "{{GRADUATION}}": machine.res_uuc if machine.res_uuc else '-',
            "{{OPTION}}": machine.option if machine.option else 'N/A',
            "{{CUSTOMER_ASSET_ID}}": machine.customer_asset_id if machine.customer_asset_id else '-',
            
            # ข้อมูลการสอบเทียบ
            "{{RECEIVED_DATE}}": getattr(calibration, 'received_date', None).strftime('%d-%b-%Y') if getattr(calibration, 'received_date', None) else '-',
            "{{DATE_OF_CALIBRATION}}": calibration.date_calibration.strftime('%d-%b-%Y') if calibration.date_calibration else '-',
            "{{DUE_DATE}}": calibration.next_due.strftime('%d-%b-%Y') if calibration.next_due else '-',
            "{{ISSUE_DATE}}": getattr(calibration, 'issue_date', None).strftime('%d-%b-%Y') if getattr(calibration, 'issue_date', None) else '-',
            "{{CERTIFICATE_NUMBER}}": calibration.certificate_number if calibration.certificate_number else '-',
            "{{PROCEDURE}}": getattr(calibration, 'procedure_number', None) if getattr(calibration, 'procedure_number', None) else '-',
            
            # ข้อมูลมาตรฐาน
            "{{STANDARD_ASSET_NO}}": getattr(standard, 'asset_number', None) if standard and getattr(standard, 'asset_number', None) else '-',
            "{{STANDARD_DESCRIPTION}}": standard.name if standard else '-',
            "{{STANDARD_MAKER_MODEL}}": getattr(standard, 'description', None) if standard and getattr(standard, 'description', None) else '-',
            "{{STANDARD_SERIAL}}": standard.name if standard else '-',
            "{{STANDARD_CERTIFICATE}}": getattr(standard, 'certificate_number', None) if standard and getattr(standard, 'certificate_number', None) else '-',
            "{{STANDARD_DUE_DATE}}": getattr(standard, 'due_date', None).strftime('%d-%b-%Y') if standard and getattr(standard, 'due_date', None) else '-',
            
            # ข้อมูลผู้รับผิดชอบ
            "{{CALIBRATOR}}": str(calibration.calibrator) if calibration.calibrator else '-',
            "{{APPROVER}}": str(calibration.certificate_issuer) if calibration.certificate_issuer else '-',
            
            # ข้อมูลหน่วยงานและที่อยู่
            "{{CUSTOMER}}": organization.name if organization else 'Physical Lab, Metrology Division, DC&E (Royal Thai Air Force)',
            "{{CUSTOMER_ADDRESS}}": organization.address if organization and organization.address else '171 Building. No2025 Sanambin, Donmueang\nBangkok, 10210',
            "{{LOCATION_OF_CALIBRATION}}": organization.name if organization else 'Metrology Division, DC&E (Royal Thai Air Force)',
            "{{LOCATION_ADDRESS}}": organization.address if organization and organization.address else '171 Building. No2025 Sanambin, Donmueang\nBangkok, 10210',
            
            # ข้อมูลตารางการสอบเทียบ Low Frequency (ใช้ข้อมูลจริง)
            "{{DC_UUC_RANGE_1}}": calibration.measurement_range if calibration.measurement_range else '-',
            "{{DC_UUC_SETTING_1}}": getattr(calibration, 'dc_voltage_set', '-') if hasattr(calibration, 'dc_voltage_set') else '-',
            "{{DC_MEASURED_VALUE_1}}": getattr(calibration, 'dc_voltage_actual', '-') if hasattr(calibration, 'dc_voltage_actual') else '-',
            "{{DC_UNCERTAINTY_1}}": getattr(calibration, 'dc_voltage_error', '-') if hasattr(calibration, 'dc_voltage_error') else '-',
            "{{DC_TOLERANCE_LIMIT_1}}": getattr(calibration, 'dc_voltage_tolerance', '-') if hasattr(calibration, 'dc_voltage_tolerance') else '-',

            "{{AC_UUC_RANGE_1}}": getattr(calibration, 'ac_voltage_range', '-') if hasattr(calibration, 'ac_voltage_range') else '-',
            "{{AC_UUC_SETTING_1}}": getattr(calibration, 'ac_voltage_set', '-') if hasattr(calibration, 'ac_voltage_set') else '-',
            "{{AC_MEASURED_VALUE_1}}": getattr(calibration, 'ac_voltage_actual', '-') if hasattr(calibration, 'ac_voltage_actual') else '-',
            "{{AC_UNCERTAINTY_1}}": getattr(calibration, 'ac_voltage_error', '-') if hasattr(calibration, 'ac_voltage_error') else '-',
            "{{AC_TOLERANCE_LIMIT_1}}": getattr(calibration, 'ac_voltage_tolerance', '-') if hasattr(calibration, 'ac_voltage_tolerance') else '-',

            "{{RES_UUC_RANGE_1}}": getattr(calibration, 'resistance_range', '-') if hasattr(calibration, 'resistance_range') else '-',
            "{{RES_UUC_SETTING_1}}": getattr(calibration, 'resistance_set', '-') if hasattr(calibration, 'resistance_set') else '-',
            "{{RES_MEASURED_VALUE_1}}": getattr(calibration, 'resistance_actual', '-') if hasattr(calibration, 'resistance_actual') else '-',
            "{{RES_UNCERTAINTY_1}}": getattr(calibration, 'resistance_error', '-') if hasattr(calibration, 'resistance_error') else '-',
            "{{RES_TOLERANCE_LIMIT_1}}": getattr(calibration, 'resistance_tolerance', '-') if hasattr(calibration, 'resistance_tolerance') else '-',
        }
        
        print(f"✅ ดึงข้อมูลจริงจากฐานข้อมูลสำเร็จ")
        
    except Exception as e:
        print(f"❌ ไม่สามารถดึงข้อมูลจากฐานข้อมูลได้: {e}")
        print("🔄 ใช้ข้อมูลจำลองแทน...")
        
        # ใช้ข้อมูลจำลองเป็น fallback
        ctx = {
            # header/common
            "{{MODEL}}": "Test Model",
            "{{MANUFACTURER}}": "Test Manufacturer",
            "{{DESCRIPTION}}": "Test Description",
            "{{SERIAL_NUMBER}}": "TEST123",
            "{{RANGE}}": "0-100 Hz",
            "{{GRADUATION}}": "0.1 Hz",
            "{{CUSTOMER_ASSET_ID}}": "ASSET001",
            "{{RECEIVED_DATE}}": "01-Jan-2025",
            "{{DATE_OF_CALIBRATION}}": "02-Jan-2025",
            "{{DUE_DATE}}": "02-Jan-2026",
            "{{ISSUE_DATE}}": "03-Jan-2025",
            "{{CERTIFICATE_NUMBER}}": "CERT001",
            "{{PROCEDURE}}": "PROC001",
            "{{STANDARD_ASSET_NO}}": "STD001",
            "{{STANDARD_DESCRIPTION}}": "Standard Description",
            "{{STANDARD_MAKER_MODEL}}": "Standard Model",
            "{{STANDARD_SERIAL}}": "STD123",
            "{{STANDARD_CERTIFICATE}}": "STDCERT001",
            "{{STANDARD_DUE_DATE}}": "01-Jan-2026",
            "{{CALIBRATOR}}": "Test Calibrator",
            "{{APPROVER}}": "Test Approver",
            "{{CUSTOMER}}": "Test Customer",
            "{{CUSTOMER_ADDRESS}}": "Test Address",
            "{{LOCATION_OF_CALIBRATION}}": "Test Location",
            "{{LOCATION_ADDRESS}}": "Test Location Address",

            # DC/AC/RES row 1 (ทดสอบตัวอย่าง)
            "{{DC_UUC_RANGE_1}}": "50 mV",
            "{{DC_UUC_SETTING_1}}": "50 mV",
            "{{DC_MEASURED_VALUE_1}}": "50 mV",
            "{{DC_UNCERTAINTY_1}}": "0.0000 mV",
            "{{DC_TOLERANCE_LIMIT_1}}": "49.0035 – 50.0065 mV",

            "{{AC_UUC_RANGE_1}}": "50 mV",
            "{{AC_UUC_SETTING_1}}": "50 mV",
            "{{AC_MEASURED_VALUE_1}}": "50 mV",
            "{{AC_UNCERTAINTY_1}}": "0.0000 mV",
            "{{AC_TOLERANCE_LIMIT_1}}": "45.0080 – 50.0020 mV",

            "{{RES_UUC_RANGE_1}}": "500 Ω",
            "{{RES_UUC_SETTING_1}}": "500 Ω",
            "{{RES_MEASURED_VALUE_1}}": "500 Ω",
            "{{RES_UNCERTAINTY_1}}": "0.0000 Ω",
            "{{RES_TOLERANCE_LIMIT_1}}": "499.0077 – 50.0033 Ω",
        }

    # แทนค่าทั้งแบบ strict (paragraph/table) และแบบ XPath (เผื่อ text box/shape)
    replace_in_doc(doc, ctx)
    replace_everywhere(doc, ctx)

    # รายงาน placeholder ที่ยังเหลือ (ถ้ามี)
    leftovers = find_left_placeholders(doc)
    if leftovers:
        print("⚠️ placeholder ยังเหลืออยู่ (ตรวจสะกดชื่อในไฟล์ .docx):")
        for s in leftovers[:10]:
            print("   •", s.strip()[:120])

    doc.save(OUTPUT)
    print("✅ บันทึกไฟล์:", OUTPUT)

if __name__ == "__main__":
    main()
