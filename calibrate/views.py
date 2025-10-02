from __future__ import annotations

from typing import List

from django.contrib import messages
from django.contrib.auth.mixins import LoginRequiredMixin, PermissionRequiredMixin
from django.db import models
from django.http import HttpRequest, QueryDict
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse_lazy
from django.views.generic import ListView, CreateView, UpdateView, DeleteView

# ==== Imports: adjust to your app structure ====
from calibrate.models import (
    CalibrationPressure,
    CalibrationTorque,
    CalibrationEquipmentUsed,
    DialGaugeCalibration, 
    BalanceCalibration, 
    MicrowaveCalibration, 
    HighFrequencyCalibration, 
    LowFrequencyCalibration
)
from calibrate.forms import (
    CalibrationPressureForm,
    CalibrationTorqueForm,
    DialGaugeCalibrationForm, 
    BalanceCalibrationForm, 
    MicrowaveCalibrationForm, 
    HighFrequencyCalibrationForm, 
    LowFrequencyCalibrationForm
)
from machine.models import Machine, MachineType, CalibrationEquipment

# Additional imports for document generation
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime
import io
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse





class CalibrationPressureListView(LoginRequiredMixin, ListView):
    model = CalibrationPressure
    template_name = 'calibrate/pressure_list.html'
    context_object_name = 'calibrations'
    
    def get_queryset(self):
        # เรียงลำดับตามระดับความเร่งด่วน: ด่วนมาก -> ด่วน -> ปกติ
        return CalibrationPressure.objects.annotate(
            priority_order=models.Case(
                models.When(priority='very_urgent', then=models.Value(1)),
                models.When(priority='urgent', then=models.Value(2)),
                models.When(priority='normal', then=models.Value(3)),
                default=models.Value(4),
                output_field=models.IntegerField(),
            )
        ).order_by('priority_order', '-update')

class CalibrationPressureCreateView(LoginRequiredMixin, PermissionRequiredMixin, CreateView):
    model = CalibrationPressure
    form_class = CalibrationPressureForm
    template_name = 'calibrate/pressure_form.html'
    success_url = reverse_lazy('calibrate-dashboard')
    permission_required = 'calibrate.add_calibrationpressure'
    
    def get_form(self, form_class=None):
        form = super().get_form(form_class)
        from django.contrib.auth import get_user_model
        User = get_user_model()
        users = User.objects.filter(is_active=True).order_by('first_name', 'last_name', 'username')
        form.fields['calibrator'].queryset = users
        form.fields['certificate_issuer'].queryset = users
        return form
    
    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        if self.request.method == 'POST':
            data = kwargs.get('data')
            if data is None:
                from django.http import QueryDict
                data = QueryDict(mutable=True)
            else:
                data = data.copy()
            if 'status' not in self.request.POST:
                data['status'] = 'in_progress'
            if 'priority' not in self.request.POST:
                data['priority'] = 'normal'
            kwargs['data'] = data
        return kwargs
    
    def form_valid(self, form):
        print("=== DEBUG: ก่อนบันทึก Pressure Create ===")
        print(f"POST data: {dict(self.request.POST)}")
        print(f"Form is_valid: {form.is_valid()}")
        print(f"Form cleaned_data: {form.cleaned_data}")
        
        # ตรวจสอบ form errors
        if form.errors:
            print(f"❌ Form errors: {form.errors}")
            return self.form_invalid(form)
        
        try:
            # บันทึกข้อมูลการสอบเทียบ
            calibration = form.save()
            print(f"✅ บันทึก calibration ID: {calibration.cal_pressure_id}")
            # ลบ success message ออกตามที่ผู้ใช้ต้องการ
            return redirect(self.success_url)
        except Exception as e:
            print(f"❌ Error saving calibration: {e}")
            messages.error(self.request, f'เกิดข้อผิดพลาดในการบันทึก: {str(e)}')
            return self.form_invalid(form)

class CalibrationPressureUpdateView(LoginRequiredMixin, PermissionRequiredMixin, UpdateView):
    model = CalibrationPressure
    form_class = CalibrationPressureForm
    template_name = 'calibrate/pressure_form.html'
    success_url = reverse_lazy('calibrate-dashboard')
    permission_required = 'calibrate.change_calibrationpressure'
    
    def get_form(self, form_class=None):
        form = super().get_form(form_class)
        from django.contrib.auth import get_user_model
        User = get_user_model()
        users = User.objects.filter(is_active=True).order_by('first_name', 'last_name', 'username')
        form.fields['calibrator'].queryset = users
        form.fields['certificate_issuer'].queryset = users
        return form
    
    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        if self.request.method == 'POST':
            data = kwargs.get('data')
            if data is None:
                from django.http import QueryDict
                data = QueryDict(mutable=True)
            else:
                data = data.copy()

            # inject status ถ้าไม่ส่งมา
            if 'status' not in self.request.POST:
                data['status'] = getattr(getattr(self, 'object', None), 'status', None) or 'in_progress'

            # ✅ inject priority ถ้าไม่ส่งมา
            if 'priority' not in self.request.POST:
                data['priority'] = getattr(getattr(self, 'object', None), 'priority', None) or 'normal'

            kwargs['data'] = data
        return kwargs
    
    def form_invalid(self, form):
        print("❌ Pressure form invalid")
        print("Errors:", form.errors)
        print("Non-field:", form.non_field_errors())
        print("POST data:", dict(self.request.POST))
        return super().form_invalid(form)
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = f'แก้ไขการสอบเทียบ Pressure - {self.object.uuc_id.name if self.object.uuc_id else "เครื่องมือ"}'
        # เพิ่มข้อมูลเครื่องมือที่ใช้ในการสอบเทียบนี้
        if self.object:
            from calibrate.models import CalibrationEquipmentUsed
            used_equipment = CalibrationEquipmentUsed.objects.filter(
                calibration_type='pressure',
                calibration_id=self.object.cal_pressure_id
            ).select_related('equipment')
            context['used_equipment'] = used_equipment
        return context
    
    def form_valid(self, form):
        print("=== DEBUG: ก่อนบันทึก Pressure Update ===")
        print(f"POST data: {dict(self.request.POST)}")
        print(f"Form is_valid: {form.is_valid()}")
        print(f"Form cleaned_data: {form.cleaned_data}")
        print(f"Object ID: {self.object.cal_pressure_id if self.object else 'None'}")
        print(f"Request method: {self.request.method}")
        print(f"Request user: {self.request.user}")
        
        # ตรวจสอบ form errors
        if form.errors:
            print(f"❌ Form errors: {form.errors}")
            return self.form_invalid(form)
        
        # ตรวจสอบข้อมูลที่สำคัญก่อนบันทึก
        print("=== DEBUG: ตรวจสอบข้อมูลก่อนบันทึก ===")
        print(f"UUC Set: {form.cleaned_data.get('set', 'MISSING')}")
        print(f"Actual: {form.cleaned_data.get('actual', 'MISSING')}")
        print(f"Tolerance Start: {form.cleaned_data.get('tolerance_start', 'MISSING')}")
        print(f"Tolerance End: {form.cleaned_data.get('tolerance_end', 'MISSING')}")
        
        # ตรวจสอบข้อมูลแถวที่ 2-6
        for i in range(2, 7):
            set_field = f'set_{i}'
            actual_field = f'actual_{i}'
            tolerance_start_field = f'tolerance_start_{i}'
            tolerance_end_field = f'tolerance_end_{i}'
            
            print(f"Row {i} - Set: {form.cleaned_data.get(set_field, 'MISSING')}")
            print(f"Row {i} - Actual: {form.cleaned_data.get(actual_field, 'MISSING')}")
            print(f"Row {i} - Tolerance Start: {form.cleaned_data.get(tolerance_start_field, 'MISSING')}")
            print(f"Row {i} - Tolerance End: {form.cleaned_data.get(tolerance_end_field, 'MISSING')}")
        
        # ตรวจสอบ required fields
        required_fields = ['uuc_id', 'measurement_range', 'update', 'next_due']
        for field in required_fields:
            if field in form.cleaned_data and form.cleaned_data[field]:
                print(f"✅ {field}: {form.cleaned_data[field]}")
            else:
                print(f"⚠️ {field}: {form.cleaned_data.get(field, 'MISSING')}")
        
        # ตรวจสอบ form fields ทั้งหมด
        print("=== DEBUG: ตรวจสอบ form fields ทั้งหมด ===")
        for field_name, field_value in form.cleaned_data.items():
            if field_name.startswith(('set', 'actual', 'tolerance')):
                print(f"Form field {field_name}: {field_value} (type: {type(field_value)})")
        
        # ตรวจสอบ POST data โดยตรง
        print("=== DEBUG: ตรวจสอบ POST data โดยตรง ===")
        for key, value in self.request.POST.items():
            if key.startswith(('set', 'actual', 'tolerance')):
                print(f"POST {key}: {value}")
        
        # ดูข้อมูลเครื่องมือก่อนบันทึก
        from calibrate.models import CalibrationEquipmentUsed
        old_equipment = CalibrationEquipmentUsed.objects.filter(
            calibration_type='pressure',
            calibration_id=self.object.cal_pressure_id if self.object else None
        )
        print(f"เครื่องมือเก่าที่มีอยู่: {list(old_equipment.values('equipment__name', 'equipment__id'))}")
        
        try:
            # บันทึกข้อมูลการสอบเทียบ
            obj = form.save(commit=False)
            if obj.uuc_id_id is None:
                obj.uuc_id_id = self.get_object().uuc_id_id  # อุ้มของเดิมไว้
            obj.save()
            calibration = obj
            print(f"✅ บันทึก calibration ID: {calibration.cal_pressure_id}")
        except Exception as e:
            print(f"❌ Error saving calibration: {e}")
            return self.form_invalid(form)
        
        # รวบรวมเครื่องมือทั้งหมดที่จะบันทึก (เหมือน Low Frequency)
        equipment_ids = []
        
        # 1. จัดการข้อมูลเครื่องมือจาก selected_equipment
        selected_equipment = self.request.POST.get('selected_equipment', '')
        print(f"=== DEBUG: selected_equipment = '{selected_equipment}'")
        if selected_equipment:
            equipment_ids.extend([eid.strip() for eid in selected_equipment.split(',') if eid.strip()])
        
        # 2. จัดการข้อมูลเครื่องมือจาก form field หลัก (std_id)
        std_id_value = self.request.POST.get('std_id', '')
        if std_id_value:
            equipment_ids.append(std_id_value)
        
        # 3. จัดการข้อมูลเครื่องมือจาก form fields (std_id_*)
        for key, value in self.request.POST.items():
            if key.startswith('std_id_') and value and not key.startswith('std_id_existing_'):
                equipment_ids.append(value)
        
        # 4. จัดการข้อมูลเครื่องมือที่มีอยู่แล้ว (std_id_existing_*)
        for key, value in self.request.POST.items():
            if key.startswith('std_id_existing_') and value:
                equipment_ids.append(value)
        
        # ทำ equipment_ids ให้ unique และกรองค่าว่าง
        equipment_ids = [e for e in equipment_ids if e]
        equipment_ids = list(dict.fromkeys(equipment_ids))  # รักษา order และ unique
        
        print(f"=== DEBUG: Final equipment_ids = {equipment_ids}")
        
        # ลบเครื่องมือเก่าทั้งหมด
        from calibrate.models import CalibrationEquipmentUsed
        CalibrationEquipmentUsed.objects.filter(
            calibration_type='pressure',
            calibration_id=calibration.cal_pressure_id
        ).delete()
        
        # บันทึกเครื่องมือใหม่ทั้งหมด
        for equipment_id in equipment_ids:
            try:
                from machine.models import CalibrationEquipment
                equipment = CalibrationEquipment.objects.get(id=equipment_id)
                CalibrationEquipmentUsed.objects.create(
                    calibration_type='pressure',
                    calibration_id=calibration.cal_pressure_id,
                    equipment=equipment
                )
                print(f"=== DEBUG: Saved equipment {equipment.name} (ID: {equipment_id})")
            except CalibrationEquipment.DoesNotExist:
                print(f"=== DEBUG: Equipment ID {equipment_id} not found")
                continue
        
        # ดูข้อมูลเครื่องมือหลังบันทึก
        new_equipment = CalibrationEquipmentUsed.objects.filter(
            calibration_type='pressure',
            calibration_id=calibration.cal_pressure_id
        )
        print(f"เครื่องมือใหม่ที่บันทึกแล้ว: {list(new_equipment.values('equipment__name', 'equipment__id'))}")
        
        # ตรวจสอบข้อมูลที่บันทึกแล้ว
        print("=== DEBUG: ตรวจสอบข้อมูลหลังบันทึก ===")
        print(f"บันทึกแล้ว - UUC Set: {calibration.set}")
        print(f"บันทึกแล้ว - Actual: {calibration.actual}")
        print(f"บันทึกแล้ว - Tolerance Start: {calibration.tolerance_start}")
        print(f"บันทึกแล้ว - Tolerance End: {calibration.tolerance_end}")
        
        # ตรวจสอบข้อมูลแถวที่ 2-6 หลังบันทึก
        for i in range(2, 7):
            set_field = f'set_{i}'
            actual_field = f'actual_{i}'
            tolerance_start_field = f'tolerance_start_{i}'
            tolerance_end_field = f'tolerance_end_{i}'
            
            print(f"บันทึกแล้ว Row {i} - Set: {getattr(calibration, set_field, 'MISSING')}")
            print(f"บันทึกแล้ว Row {i} - Actual: {getattr(calibration, actual_field, 'MISSING')}")
            print(f"บันทึกแล้ว Row {i} - Tolerance Start: {getattr(calibration, tolerance_start_field, 'MISSING')}")
            print(f"บันทึกแล้ว Row {i} - Tolerance End: {getattr(calibration, tolerance_end_field, 'MISSING')}")
        
        print("=== DEBUG: หลังบันทึก Pressure ===")
        
        # ตรวจสอบสถานะอัตโนมัติ
        self.auto_check_status(calibration)
        
        # เพิ่ม success message
        from django.contrib import messages
        # ลบ success message ออกตามที่ผู้ใช้ต้องการ
        
        # Redirect ไปยัง success_url โดยตรง
        return redirect(self.success_url)
    
    def auto_check_status(self, calibration):
        """ตรวจสอบสถานะอัตโนมัติสำหรับ Pressure"""
        # ตรวจสอบว่ามีข้อมูลครบถ้วนหรือไม่
        has_data = any([
            calibration.set, calibration.m1, calibration.m2, calibration.m3, calibration.m4,
            calibration.set_2, calibration.m1_2, calibration.m2_2, calibration.m3_2, calibration.m4_2,
            calibration.set_3, calibration.m1_3, calibration.m2_3, calibration.m3_3, calibration.m4_3,
            calibration.set_4, calibration.m1_4, calibration.m2_4, calibration.m3_4, calibration.m4_4,
            calibration.set_5, calibration.m1_5, calibration.m2_5, calibration.m3_5, calibration.m4_5,
            calibration.set_6, calibration.m1_6, calibration.m2_6, calibration.m3_6, calibration.m4_6
        ])
        
        if not has_data:
            # ยังไม่มีข้อมูล = กำลังสอบเทียบ
            calibration.status = 'in_progress'
        else:
            # มีข้อมูลแล้ว = ตรวจสอบผลการสอบเทียบ
            if self.check_pressure_pass_fail(calibration):
                calibration.status = 'passed'
            else:
                calibration.status = 'failed'
        
        calibration.save()
    
    def check_pressure_pass_fail(self, calibration):
        """ตรวจสอบผลการสอบเทียบ Pressure
        
        เงื่อนไข: ค่าของช่อง UUC Set และ ค่าของช่อง Actual 
        ต้องอยู่ในระหว่าง ค่าของช่อง Tolerance Limit ทั้งคู่ ถึงจะผ่านการสอบเทียบ
        """
        # ตรวจสอบทุกแถวที่มีข้อมูล - ใช้เฉพาะ set, actual, tolerance_start, tolerance_end
        rows_to_check = [
            (calibration.set, calibration.actual, calibration.tolerance_start, calibration.tolerance_end),
            (calibration.set_2, calibration.actual_2, calibration.tolerance_start_2, calibration.tolerance_end_2),
            (calibration.set_3, calibration.actual_3, calibration.tolerance_start_3, calibration.tolerance_end_3),
            (calibration.set_4, calibration.actual_4, calibration.tolerance_start_4, calibration.tolerance_end_4),
            (calibration.set_5, calibration.actual_5, calibration.tolerance_start_5, calibration.tolerance_end_5),
            (calibration.set_6, calibration.actual_6, calibration.tolerance_start_6, calibration.tolerance_end_6)
        ]
        
        for set_val, actual, tolerance_start, tolerance_end in rows_to_check:
            if set_val and actual is not None and tolerance_start is not None and tolerance_end is not None:
                try:
                    set_float = float(set_val)
                    actual_float = float(actual)
                    tolerance_start_float = float(tolerance_start)
                    tolerance_end_float = float(tolerance_end)
                    
                    # ตรวจสอบ UUC Set อยู่ในช่วง Tolerance Limit
                    set_in_range = tolerance_start_float <= set_float <= tolerance_end_float
                    # ตรวจสอบ Actual อยู่ในช่วง Tolerance Limit
                    actual_in_range = tolerance_start_float <= actual_float <= tolerance_end_float
                    
                    # ผ่านการสอบเทียบเมื่อทั้ง UUC Set และ Actual อยู่ในช่วง Tolerance Limit
                    if not (set_in_range and actual_in_range):
                        return False  # ไม่ผ่าน
                        
                except (ValueError, TypeError):
                    return False  # ข้อมูลไม่ถูกต้อง
                    
        return True  # ผ่าน

class CalibrationPressureDeleteView(LoginRequiredMixin, PermissionRequiredMixin, DeleteView):
    model = CalibrationPressure
    template_name = 'calibrate/pressure_confirm_delete.html'
    success_url = reverse_lazy('calibrate-dashboard')
    permission_required = 'calibrate.delete_calibrationpressure'

class CalibrationTorqueListView(LoginRequiredMixin, ListView):
    model = CalibrationTorque
    template_name = 'calibrate/torque_list.html'
    context_object_name = 'calibrations'
    
    def get_queryset(self):
        # เรียงลำดับตามระดับความเร่งด่วน: ด่วนมาก -> ด่วน -> ปกติ
        return CalibrationTorque.objects.annotate(
            priority_order=models.Case(
                models.When(priority='very_urgent', then=models.Value(1)),
                models.When(priority='urgent', then=models.Value(2)),
                models.When(priority='normal', then=models.Value(3)),
                default=models.Value(4),
                output_field=models.IntegerField(),
            )
        ).order_by('priority_order', '-update')

class CalibrationTorqueCreateView(LoginRequiredMixin, PermissionRequiredMixin, CreateView):
    model = CalibrationTorque
    form_class = CalibrationTorqueForm
    template_name = 'calibrate/torque_form.html'
    success_url = reverse_lazy('calibrate-dashboard')
    permission_required = 'calibrate.add_calibrationtorque'
    
    def get_form(self, form_class=None):
        form = super().get_form(form_class)
        from django.contrib.auth import get_user_model
        User = get_user_model()
        users = User.objects.filter(is_active=True).order_by('first_name', 'last_name', 'username')
        form.fields['calibrator'].queryset = users
        form.fields['certificate_issuer'].queryset = users
        return form
    
    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        if self.request.method == 'POST':
            data = kwargs.get('data')
            if data is None:
                from django.http import QueryDict
                data = QueryDict(mutable=True)
            else:
                data = data.copy()
            if 'status' not in self.request.POST:
                data['status'] = 'in_progress'
            if 'priority' not in self.request.POST:
                data['priority'] = 'normal'
            kwargs['data'] = data
        return kwargs
    
    def form_valid(self, form):
        print("=== DEBUG: ก่อนบันทึก Torque Create ===")
        print(f"POST data: {dict(self.request.POST)}")
        print(f"Form is_valid: {form.is_valid()}")
        print(f"Form cleaned_data: {form.cleaned_data}")
        
        # ตรวจสอบ form errors
        if form.errors:
            print(f"❌ Form errors: {form.errors}")
            return self.form_invalid(form)
        
        try:
            # บันทึกข้อมูลการสอบเทียบ
            calibration = form.save()
            print(f"✅ บันทึก calibration ID: {calibration.cal_torque_id}")
            # ลบ success message ออกตามที่ผู้ใช้ต้องการ
            return redirect(self.success_url)
        except Exception as e:
            print(f"❌ Error saving calibration: {e}")
            messages.error(self.request, f'เกิดข้อผิดพลาดในการบันทึก: {str(e)}')
            return self.form_invalid(form)

class CalibrationTorqueUpdateView(LoginRequiredMixin, PermissionRequiredMixin, UpdateView):
    model = CalibrationTorque
    form_class = CalibrationTorqueForm
    template_name = 'calibrate/torque_form.html'
    success_url = reverse_lazy('calibrate-dashboard')
    permission_required = 'calibrate.change_calibrationtorque'
    
    def get_form(self, form_class=None):
        form = super().get_form(form_class)
        from django.contrib.auth import get_user_model
        User = get_user_model()
        users = User.objects.filter(is_active=True).order_by('first_name', 'last_name', 'username')
        form.fields['calibrator'].queryset = users
        form.fields['certificate_issuer'].queryset = users
        return form
    
    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        if self.request.method == 'POST':
            data = kwargs.get('data')
            if data is None:
                from django.http import QueryDict
                data = QueryDict(mutable=True)
            else:
                data = data.copy()

            # inject status ถ้าไม่ส่งมา
            if 'status' not in self.request.POST:
                data['status'] = getattr(getattr(self, 'object', None), 'status', None) or 'in_progress'

            # ✅ inject priority ถ้าไม่ส่งมา
            if 'priority' not in self.request.POST:
                data['priority'] = getattr(getattr(self, 'object', None), 'priority', None) or 'normal'

            kwargs['data'] = data
        return kwargs
    
    def form_invalid(self, form):
        print("=== DEBUG: Torque Update Form Invalid ===")
        print(f"Form errors: {form.errors}")
        print(f"Form non_field_errors: {form.non_field_errors}")
        print(f"POST data: {dict(self.request.POST)}")
        return super().form_invalid(form)
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = f'แก้ไขการสอบเทียบ Torque - {self.object.uuc_id.name if self.object.uuc_id else "เครื่องมือ"}'
        # เพิ่มข้อมูลเครื่องมือที่ใช้ในการสอบเทียบนี้
        if self.object:
            from calibrate.models import CalibrationEquipmentUsed
            used_equipment = CalibrationEquipmentUsed.objects.filter(
                calibration_type='torque',
                calibration_id=self.object.cal_torque_id
            ).select_related('equipment')
            context['used_equipment'] = used_equipment
        return context
    
    def form_valid(self, form):
        print("=== DEBUG: ก่อนบันทึก Torque Update ===")
        print(f"POST data: {dict(self.request.POST)}")
        print(f"Form is_valid: {form.is_valid()}")
        print(f"Form cleaned_data: {form.cleaned_data}")
        print(f"Object ID: {self.object.cal_torque_id if self.object else 'None'}")
        print(f"Request method: {self.request.method}")
        print(f"Request user: {self.request.user}")
        
        # ตรวจสอบ form errors
        if form.errors:
            print(f"❌ Form errors: {form.errors}")
            return self.form_invalid(form)
        
        # ตรวจสอบ required fields
        required_fields = ['uuc_id', 'measurement_range', 'update', 'next_due']
        for field in required_fields:
            if field in form.cleaned_data and form.cleaned_data[field]:
                print(f"✅ {field}: {form.cleaned_data[field]}")
            else:
                print(f"⚠️ {field}: {form.cleaned_data.get(field, 'MISSING')}")
        
        try:
            # บันทึกข้อมูลการสอบเทียบ
            calibration = form.save()
            print(f"✅ บันทึก calibration ID: {calibration.pk}")
        except Exception as e:
            print(f"❌ Error saving calibration: {e}")
            return self.form_invalid(form)
        
        # จัดการข้อมูลเครื่องมือที่ใช้สอบเทียบหลายตัว
        selected_equipment = self.request.POST.get('selected_equipment', '')
        if selected_equipment:
            print(f"Selected equipment: {selected_equipment}")
            # ลบข้อมูลเครื่องมือเก่าที่เกี่ยวข้องกับการสอบเทียบนี้
            from calibrate.models import CalibrationEquipmentUsed
            CalibrationEquipmentUsed.objects.filter(
                calibration_type='torque',
                calibration_id=calibration.cal_torque_id
            ).delete()
            
            # เพิ่มข้อมูลเครื่องมือใหม่ (ใช้ set เพื่อหลีกเลี่ยงการซ้ำ)
            equipment_ids = set()
            for eid in selected_equipment.split(','):
                eid = eid.strip()
                if eid:
                    equipment_ids.add(eid)
            
            for equipment_id in equipment_ids:
                try:
                    from machine.models import CalibrationEquipment
                    equipment = CalibrationEquipment.objects.get(id=equipment_id)
                    CalibrationEquipmentUsed.objects.get_or_create(
                        calibration_type='torque',
                        calibration_id=calibration.cal_torque_id,
                        equipment=equipment
                    )
                    print(f"Added equipment: {equipment.name}")
                except Exception as e:
                    print(f"Error adding equipment {equipment_id}: {e}")
        
        # ตรวจสอบสถานะอัตโนมัติ
        self.auto_check_status(calibration)
        
        # เพิ่ม success message
        from django.contrib import messages
        # ลบ success message ออกตามที่ผู้ใช้ต้องการ
        
        # Redirect ไปยัง success_url โดยตรง
        return redirect(self.success_url)
    
    def auto_check_status(self, calibration):
        """ตรวจสอบสถานะอัตโนมัติสำหรับ Torque"""
        # ตรวจสอบว่ามีข้อมูลครบถ้วนหรือไม่
        has_data = any([
            calibration.cwset, calibration.cw0, calibration.cw90, calibration.cw180, calibration.cw270,
            calibration.ccwset, calibration.ccw0, calibration.ccw90, calibration.ccw180, calibration.ccw270,
            calibration.cwset_2, calibration.cwset_3, calibration.ccwset_2, calibration.ccwset_3
        ])
        
        if not has_data:
            # ยังไม่มีข้อมูล = กำลังสอบเทียบ
            calibration.status = 'in_progress'
        else:
            # มีข้อมูลแล้ว = ตรวจสอบผลการสอบเทียบ
            if self.check_torque_pass_fail(calibration):
                calibration.status = 'passed'
            else:
                calibration.status = 'failed'
        
        calibration.save()
    
    def check_torque_pass_fail(self, calibration):
        """ตรวจสอบผลการสอบเทียบ Torque"""
        # ตรวจสอบ CW
        cw_rows = [
            (calibration.cwset, calibration.cw_actual, calibration.cw_tolerance_start, calibration.cw_tolerance_end),
            (calibration.cwset_2, calibration.cw_actual_2, calibration.cw_tolerance_start_2, calibration.cw_tolerance_end_2),
            (calibration.cwset_3, calibration.cw_actual_3, calibration.cw_tolerance_start_3, calibration.cw_tolerance_end_3)
        ]
        
        for set_val, actual, tolerance_start, tolerance_end in cw_rows:
            if set_val and actual is not None and tolerance_start is not None and tolerance_end is not None:
                if not (tolerance_start <= actual <= tolerance_end):
                    return False
        
        # ตรวจสอบ CCW
        ccw_rows = [
            (calibration.ccwset, calibration.ccw_actual, calibration.ccw_tolerance_start, calibration.ccw_tolerance_end),
            (calibration.ccwset_2, calibration.ccw_actual_2, calibration.ccw_tolerance_start_2, calibration.ccw_tolerance_end_2),
            (calibration.ccwset_3, calibration.ccw_actual_3, calibration.ccw_tolerance_start_3, calibration.ccw_tolerance_end_3)
        ]
        
        for set_val, actual, tolerance_start, tolerance_end in ccw_rows:
            if set_val and actual is not None and tolerance_start is not None and tolerance_end is not None:
                if not (tolerance_start <= actual <= tolerance_end):
                    return False
        
        return True

class CalibrationTorqueDeleteView(LoginRequiredMixin, PermissionRequiredMixin, DeleteView):
    model = CalibrationTorque
    template_name = 'calibrate/torque_confirm_delete.html'
    success_url = reverse_lazy('calibrate-dashboard')
    permission_required = 'calibrate.delete_calibrationtorque'

class BalanceCalibrationListView(LoginRequiredMixin, ListView):
    model = BalanceCalibration
    template_name = 'calibrate/balance_list.html'
    context_object_name = 'calibrations'
    
    def get_queryset(self):
        # เรียงลำดับตามระดับความเร่งด่วน: ด่วนมาก -> ด่วน -> ปกติ
        return BalanceCalibration.objects.annotate(
            priority_order=models.Case(
                models.When(priority='very_urgent', then=models.Value(1)),
                models.When(priority='urgent', then=models.Value(2)),
                models.When(priority='normal', then=models.Value(3)),
                default=models.Value(4),
                output_field=models.IntegerField(),
            )
        ).order_by('priority_order', '-date_calibration')

class BalanceCalibrationCreateView(LoginRequiredMixin, PermissionRequiredMixin, CreateView):
    model = BalanceCalibration
    form_class = BalanceCalibrationForm
    template_name = 'calibrate/balance_form.html'
    success_url = reverse_lazy('calibrate-dashboard')
    permission_required = 'calibrate.add_balancecalibration'
    
    def get_form(self, form_class=None):
        form = super().get_form(form_class)
        from django.contrib.auth import get_user_model
        User = get_user_model()
        users = User.objects.filter(is_active=True).order_by('first_name', 'last_name', 'username')
        form.fields['calibrator'].queryset = users
        form.fields['certificate_issuer'].queryset = users
        return form
    
    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        if self.request.method == 'POST':
            data = kwargs.get('data')
            if data is None:
                from django.http import QueryDict
                data = QueryDict(mutable=True)
            else:
                data = data.copy()
            if 'priority' not in self.request.POST:
                data['priority'] = 'normal'
            kwargs['data'] = data
        return kwargs

class BalanceCalibrationUpdateView(LoginRequiredMixin, PermissionRequiredMixin, UpdateView):
    model = BalanceCalibration
    form_class = BalanceCalibrationForm
    template_name = 'calibrate/balance_form.html'
    success_url = reverse_lazy('calibrate-dashboard')
    permission_required = 'calibrate.change_balancecalibration'
    
    def get_form(self, form_class=None):
        form = super().get_form(form_class)
        from django.contrib.auth import get_user_model
        User = get_user_model()
        users = User.objects.filter(is_active=True).order_by('first_name', 'last_name', 'username')
        form.fields['calibrator'].queryset = users
        form.fields['certificate_issuer'].queryset = users
        return form
    
    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        if self.request.method == 'POST':
            data = kwargs.get('data')
            if data is None:
                from django.http import QueryDict
                data = QueryDict(mutable=True)
            else:
                data = data.copy()
            if 'priority' not in self.request.POST:
                data['priority'] = getattr(getattr(self, 'object', None), 'priority', None) or 'normal'
            kwargs['data'] = data
        return kwargs
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # เพิ่มข้อมูลเครื่องมือที่ใช้ในการสอบเทียบนี้
        if self.object:
            from calibrate.models import CalibrationEquipmentUsed
            used_equipment = CalibrationEquipmentUsed.objects.filter(
                calibration_type='balance',
                calibration_id=self.object.pk
            ).select_related('equipment')
            context['used_equipment'] = used_equipment
        return context
    
    def form_valid(self, form):
        print("=== DEBUG: ก่อนบันทึก Balance ===")
        print(f"POST data: {dict(self.request.POST)}")
        
        # ตรวจสอบ form errors
        if form.errors:
            print(f"❌ Form errors: {form.errors}")
            return self.form_invalid(form)
        
        try:
            # บันทึกข้อมูลการสอบเทียบ
            calibration = form.save()
            print(f"✅ บันทึก calibration ID: {calibration.pk}")
        except Exception as e:
            print(f"❌ Error saving calibration: {e}")
            return self.form_invalid(form)
        
        # จัดการข้อมูลเครื่องมือที่ใช้สอบเทียบหลายตัว
        selected_equipment = self.request.POST.get('selected_equipment', '')
        if selected_equipment:
            print(f"Selected equipment: {selected_equipment}")
            # ลบข้อมูลเครื่องมือเก่าที่เกี่ยวข้องกับการสอบเทียบนี้
            from calibrate.models import CalibrationEquipmentUsed
            CalibrationEquipmentUsed.objects.filter(
                calibration_type='balance',
                calibration_id=calibration.pk
            ).delete()
            
            # เพิ่มข้อมูลเครื่องมือใหม่ (ใช้ set เพื่อหลีกเลี่ยงการซ้ำ)
            equipment_ids = set()
            for eid in selected_equipment.split(','):
                eid = eid.strip()
                if eid:
                    equipment_ids.add(eid)
            
            for equipment_id in equipment_ids:
                try:
                    from machine.models import CalibrationEquipment
                    equipment = CalibrationEquipment.objects.get(id=equipment_id)
                    CalibrationEquipmentUsed.objects.get_or_create(
                        calibration_type='balance',
                        calibration_id=calibration.pk,
                        equipment=equipment
                    )
                    print(f"Added equipment: {equipment.name}")
                except Exception as e:
                    print(f"Error adding equipment {equipment_id}: {e}")
        
        # ตรวจสอบสถานะอัตโนมัติ
        self.auto_check_status(calibration)
        
        # เพิ่ม success message
        from django.contrib import messages
        # ลบ success message ออกตามที่ผู้ใช้ต้องการ
        
        # Redirect ไปยัง success_url โดยตรง
        return redirect(self.success_url)
    
    def auto_check_status(self, calibration):
        """ตรวจสอบสถานะอัตโนมัติสำหรับ Balance"""
        # ตรวจสอบว่ามีข้อมูลครบถ้วนหรือไม่
        has_data = any([
            calibration.linear_nominal_value, calibration.linear_conventional_mass, calibration.linear_displayed_value,
            calibration.linear_nominal_value_2, calibration.linear_conventional_mass_2, calibration.linear_displayed_value_2,
            calibration.linear_nominal_value_3, calibration.linear_conventional_mass_3, calibration.linear_displayed_value_3,
            calibration.linear_nominal_value_4, calibration.linear_conventional_mass_4, calibration.linear_displayed_value_4,
            calibration.linear_nominal_value_5, calibration.linear_conventional_mass_5, calibration.linear_displayed_value_5
        ])
        
        if not has_data:
            # ยังไม่มีข้อมูล = กำลังสอบเทียบ
            calibration.status = 'in_progress'
        else:
            # มีข้อมูลแล้ว = ตรวจสอบผลการสอบเทียบ
            if self.check_balance_pass_fail(calibration):
                calibration.status = 'passed'
            else:
                calibration.status = 'failed'
        
        calibration.save()
    
    def check_balance_pass_fail(self, calibration):
        """ตรวจสอบผลการสอบเทียบ Balance
        
        เงื่อนไข:
        1. ค่าของช่อง Displayed Value ต้องไม่เกินค่าของช่อง Conventional Mass ± 0.10
        2. ค่า Error ต้องอยู่ในช่วง ± 0.10
        """
        # ตรวจสอบทุกแถวที่มีข้อมูล
        rows_to_check = [
            (calibration.linear_conventional_mass, calibration.linear_displayed_value, calibration.linear_error),
            (calibration.linear_conventional_mass_2, calibration.linear_displayed_value_2, calibration.linear_error_2),
            (calibration.linear_conventional_mass_3, calibration.linear_displayed_value_3, calibration.linear_error_3),
            (calibration.linear_conventional_mass_4, calibration.linear_displayed_value_4, calibration.linear_error_4),
            (calibration.linear_conventional_mass_5, calibration.linear_displayed_value_5, calibration.linear_error_5)
        ]
        
        tolerance = 0.10  # ค่าความคลาดเคลื่อนที่อนุญาต ± 0.10
        
        for conventional, displayed, error in rows_to_check:
            if conventional and displayed and error is not None:
                try:
                    conventional_float = float(conventional)
                    displayed_float = float(displayed)
                    error_float = float(error)
                    
                    # ตรวจสอบ Displayed Value กับ Conventional Mass ± 0.10
                    conventional_min = conventional_float - tolerance
                    conventional_max = conventional_float + tolerance
                    displayed_in_range = conventional_min <= displayed_float <= conventional_max
                    
                    # ตรวจสอบ Error อยู่ในช่วง ± 0.10
                    error_in_range = -tolerance <= error_float <= tolerance
                    
                    # ผ่านการสอบเทียบเมื่อทั้งสองเงื่อนไขเป็นจริง
                    if not (displayed_in_range and error_in_range):
                        return False  # ไม่ผ่าน
                        
                except (ValueError, TypeError):
                    return False  # ข้อมูลไม่ถูกต้อง
        
        return True  # ผ่าน

class BalanceCalibrationDeleteView(LoginRequiredMixin, PermissionRequiredMixin, DeleteView):
    model = BalanceCalibration
    template_name = 'calibrate/balance_confirm_delete.html'
    success_url = reverse_lazy('calibrate-dashboard')
    permission_required = 'calibrate.delete_balancecalibration'

@login_required
def calibration_dashboard(request):
    """หน้าหลักสำหรับเลือกประเภทการสอบเทียบ"""
    from datetime import date, timedelta
    
    # รับพารามิเตอร์การค้นหาและกรอง
    name_search = request.GET.get('name_search', '')
    serial_search = request.GET.get('serial_search', '')
    calibration_type = request.GET.get('calibration_type', '')
    status_filter = request.GET.get('status_filter', '')
    priority_filter = request.GET.get('priority_filter', '')
    
    
    pressure_calibrations = CalibrationPressure.objects.select_related('uuc_id', 'std_id', 'calibrator', 'certificate_issuer').exclude(
        status='closed'
    ).annotate(
        priority_order=models.Case(
            models.When(priority='very_urgent', then=models.Value(1)),
            models.When(priority='urgent', then=models.Value(2)),
            models.When(priority='normal', then=models.Value(3)),
            default=models.Value(4),
            output_field=models.IntegerField(),
        )
    ).order_by('priority_order', '-update')
    
    torque_calibrations = CalibrationTorque.objects.select_related('uuc_id', 'std_id', 'calibrator', 'certificate_issuer').exclude(
        status='closed'
    ).annotate(
        priority_order=models.Case(
            models.When(priority='very_urgent', then=models.Value(1)),
            models.When(priority='urgent', then=models.Value(2)),
            models.When(priority='normal', then=models.Value(3)),
            default=models.Value(4),
            output_field=models.IntegerField(),
        )
    ).order_by('priority_order', '-update')
    
    # ดึงข้อมูลการสอบเทียบ Dial Gauge
    dial_gauge_calibrations = DialGaugeCalibration.objects.select_related('machine', 'std_id', 'calibrator', 'certificate_issuer').exclude(
        status='closed'
    ).annotate(
        priority_order=models.Case(
            models.When(priority='very_urgent', then=models.Value(1)),
            models.When(priority='urgent', then=models.Value(2)),
            models.When(priority='normal', then=models.Value(3)),
            default=models.Value(4),
            output_field=models.IntegerField(),
        )
    ).order_by('priority_order', '-date_calibration')
    
    # ดึงข้อมูลการสอบเทียบ Balance
    balance_calibrations = BalanceCalibration.objects.select_related('machine', 'std_id', 'calibrator', 'certificate_issuer').exclude(
        status='closed'
    ).annotate(
        priority_order=models.Case(
            models.When(priority='very_urgent', then=models.Value(1)),
            models.When(priority='urgent', then=models.Value(2)),
            models.When(priority='normal', then=models.Value(3)),
            default=models.Value(4),
            output_field=models.IntegerField(),
        )
    ).order_by('priority_order', '-date_calibration')
    
    # ดึงข้อมูลการสอบเทียบ Microwave
    microwave_calibrations = MicrowaveCalibration.objects.select_related('machine', 'std_id', 'calibrator', 'certificate_issuer').exclude(
        status='closed'
    ).annotate(
        priority_order=models.Case(
            models.When(priority='very_urgent', then=models.Value(1)),
            models.When(priority='urgent', then=models.Value(2)),
            models.When(priority='normal', then=models.Value(3)),
            default=models.Value(4),
            output_field=models.IntegerField(),
        )
    ).order_by('priority_order', '-date_calibration')
    
    # ดึงข้อมูลการสอบเทียบ High Frequency
    high_frequency_calibrations = HighFrequencyCalibration.objects.select_related('machine', 'std_id', 'calibrator', 'certificate_issuer').exclude(
        status='closed'
    ).annotate(
        priority_order=models.Case(
            models.When(priority='very_urgent', then=models.Value(1)),
            models.When(priority='urgent', then=models.Value(2)),
            models.When(priority='normal', then=models.Value(3)),
            default=models.Value(4),
            output_field=models.IntegerField(),
        )
    ).order_by('priority_order', '-date_calibration')
    
    # ดึงข้อมูลการสอบเทียบ Low Frequency
    low_frequency_calibrations = LowFrequencyCalibration.objects.select_related('machine', 'std_id', 'calibrator', 'certificate_issuer').exclude(
        status='closed'
    ).annotate(
        priority_order=models.Case(
            models.When(priority='very_urgent', then=models.Value(1)),
            models.When(priority='urgent', then=models.Value(2)),
            models.When(priority='normal', then=models.Value(3)),
            default=models.Value(4),
            output_field=models.IntegerField(),
        )
    ).order_by('priority_order', '-date_calibration')
    
    # รวมข้อมูลการสอบเทียบทั้งหมด
    all_calibrations = []
    
    
    # เพิ่มข้อมูล Pressure calibrations
    for cal in pressure_calibrations:
        all_calibrations.append({
            'id': cal.cal_pressure_id,
            'type': 'pressure',
            'type_name': 'การสอบเทียบ Pressure',
            'machine_name': cal.uuc_id.name if cal.uuc_id else '-',
            'machine_model': cal.uuc_id.model if cal.uuc_id else '-',
            'serial_number': cal.uuc_id.serial_number if cal.uuc_id else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.update,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.update,  # วันที่สอบเทียบ
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.uuc_id.organize.name if cal.uuc_id and cal.uuc_id.organize else '-',
        })
    
    # เพิ่มข้อมูล Torque calibrations
    for cal in torque_calibrations:
        all_calibrations.append({
            'id': cal.cal_torque_id,
            'type': 'torque',
            'type_name': 'การสอบเทียบ Torque',
            'machine_name': cal.uuc_id.name if cal.uuc_id else '-',
            'machine_model': cal.uuc_id.model if cal.uuc_id else '-',
            'serial_number': cal.uuc_id.serial_number if cal.uuc_id else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.update,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.update,  # วันที่สอบเทียบ
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.uuc_id.organize.name if cal.uuc_id and cal.uuc_id.organize else '-',
        })
    
    # เพิ่มข้อมูล Dial Gauge calibrations
    for cal in dial_gauge_calibrations:
        all_calibrations.append({
            'id': cal.id,
            'type': 'dial_gauge',
            'type_name': 'การสอบเทียบ Dial Gauge',
            'machine_name': cal.machine.name if cal.machine else '-',
            'machine_model': cal.machine.model if cal.machine else '-',
            'serial_number': cal.machine.serial_number if cal.machine else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.update,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.date_calibration,  # วันที่สอบเทียบ
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.machine.organize.name if cal.machine and cal.machine.organize else '-',
        })
    
    # เพิ่มข้อมูล Balance calibrations
    for cal in balance_calibrations:
        all_calibrations.append({
            'id': cal.id,
            'type': 'balance',
            'type_name': 'การสอบเทียบ Balance',
            'machine_name': cal.machine.name if cal.machine else '-',
            'machine_model': cal.machine.model if cal.machine else '-',
            'serial_number': cal.machine.serial_number if cal.machine else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.update,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.date_calibration,  # วันที่สอบเทียบ
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.machine.organize.name if cal.machine and cal.machine.organize else '-',
        })
    
    # เพิ่มข้อมูล Microwave calibrations
    for cal in microwave_calibrations:
        all_calibrations.append({
            'id': cal.id,
            'type': 'microwave',
            'type_name': 'การสอบเทียบ Microwave',
            'machine_name': cal.machine.name if cal.machine else '-',
            'machine_model': cal.machine.model if cal.machine else '-',
            'serial_number': cal.machine.serial_number if cal.machine else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.update or cal.date_calibration,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.date_calibration,  # วันที่สอบเทียบ
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.machine.organize.name if cal.machine and cal.machine.organize else '-',
        })
    
    # เพิ่มข้อมูล High Frequency calibrations
    for cal in high_frequency_calibrations:
        all_calibrations.append({
            'id': cal.id,
            'type': 'high-frequency',
            'type_name': 'การสอบเทียบ High Frequency',
            'machine_name': cal.machine.name if cal.machine else '-',
            'machine_model': cal.machine.model if cal.machine else '-',
            'serial_number': cal.machine.serial_number if cal.machine else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.date_calibration,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.date_calibration,  # วันที่สอบเทียบ
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.machine.organize.name if cal.machine and cal.machine.organize else '-',
        })
    
    # เพิ่มข้อมูล Low Frequency calibrations
    for cal in low_frequency_calibrations:
        all_calibrations.append({
            'id': cal.id,
            'type': 'low-frequency',
            'type_name': 'การสอบเทียบ Low Frequency',
            'machine_name': cal.machine.name if cal.machine else '-',
            'machine_model': cal.machine.model if cal.machine else '-',
            'serial_number': cal.machine.serial_number if cal.machine else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.date_calibration,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.date_calibration,  # วันที่สอบเทียบ
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.machine.organize.name if cal.machine and cal.machine.organize else '-',
        })
    
    # กรองข้อมูลตามพารามิเตอร์
    filtered_calibrations = []
    for cal in all_calibrations:
        include_item = True
        
        # กรองตามชื่อเครื่องมือ
        if name_search and name_search.lower() not in cal['machine_name'].lower():
            include_item = False
        
        # กรองตาม Serial Number
        if serial_search and serial_search.lower() not in cal['serial_number'].lower():
            include_item = False
        
        # กรองตามประเภทการสอบเทียบ
        if calibration_type and cal['type'] != calibration_type:
            include_item = False
        
        # กรองตามสถานะ
        if status_filter and cal['status'] != status_filter:
            include_item = False
        
        # กรองตามความเร่งด่วน
        if priority_filter and cal['priority'] != priority_filter:
            include_item = False
        
        if include_item:
            filtered_calibrations.append(cal)
    
    # เรียงลำดับตามระดับความเร่งด่วน: ด่วนมาก -> ด่วน -> ปกติ แล้วตามด้วยวันที่สอบเทียบล่าสุด
    priority_order = {'very_urgent': 1, 'urgent': 2, 'normal': 3}
    filtered_calibrations.sort(key=lambda x: (
        priority_order.get(x['priority'], 4),  # เรียงตามระดับความเร่งด่วน
        x['calibration_date'] if x['calibration_date'] else date.min  # แล้วตามด้วยวันที่
    ))
    
    # ข้อมูลวันที่สำหรับการคำนวณสถานะ
    today = date.today()
    today_plus_30 = today + timedelta(days=30)
    
    context = {
        'pressure_calibrations_count': CalibrationPressure.objects.count(),
        'torque_calibrations_count': CalibrationTorque.objects.count(),
        'dial_gauge_calibrations_count': DialGaugeCalibration.objects.count(),
        'balance_calibrations_count': BalanceCalibration.objects.count(),
        'microwave_calibrations_count': MicrowaveCalibration.objects.count(),
        'pending_calibrations_count': (
            CalibrationPressure.objects.filter(status='pending').count() +
            CalibrationTorque.objects.filter(status='pending').count() +
            DialGaugeCalibration.objects.filter(status='pending').count() +
            BalanceCalibration.objects.filter(status='pending').count() +
            MicrowaveCalibration.objects.filter(status='pending').count()
        ),
        'all_calibrations': filtered_calibrations,
        'today': today,
        'today_plus_30': today_plus_30,
    }
    return render(request, 'calibrate/dashboard.html', context)

@login_required
def machine_calibration_list(request, machine_id):
    """แสดงรายการการสอบเทียบของเครื่องมือเฉพาะ"""
    machine = get_object_or_404(Machine, id=machine_id)
    
    # ตรวจสอบประเภทเครื่องมือและดึงข้อมูลการสอบเทียบที่เหมาะสม
    machine_type_name = machine.machine_type.name.lower()
    
    if 'pressure' in machine_type_name:
        calibrations = CalibrationPressure.objects.filter(uuc_id=machine.id)
        calibration_type = 'pressure'
    elif 'torque' in machine_type_name:
        calibrations = CalibrationTorque.objects.filter(uuc_id=machine.id)
        calibration_type = 'torque'
    else:
        calibrations = []
        calibration_type = 'unknown'
    
    context = {
        'machine': machine,
        'calibrations': calibrations,
        'calibration_type': calibration_type,
    }
    return render(request, 'calibrate/machine_calibration_list.html', context)

@login_required
def create_calibration_for_machine(request, machine_id):
    """สร้างบันทึกการสอบเทียบสำหรับเครื่องมือเฉพาะ"""
    machine = get_object_or_404(Machine, id=machine_id)
    machine_type_name = machine.machine_type.name.lower()
    
    if request.method == 'POST':
        if 'pressure' in machine_type_name:
            form = CalibrationPressureForm(request.POST)
            template = 'calibrate/pressure_form.html'
            success_url = reverse_lazy('machine-calibration-list', kwargs={'machine_id': machine_id})
        elif 'torque' in machine_type_name:
            form = CalibrationTorqueForm(request.POST)
            template = 'calibrate/torque_form.html'
            success_url = reverse_lazy('machine-calibration-list', kwargs={'machine_id': machine_id})
        elif 'balance' in machine_type_name:
            form = BalanceCalibrationForm(request.POST)
            template = 'calibrate/balance_form.html'
            success_url = reverse_lazy('machine-calibration-list', kwargs={'machine_id': machine_id})
        elif 'microwave' in machine_type_name:
            form = MicrowaveCalibrationForm(request.POST)
            template = 'calibrate/microwave_form.html'
            success_url = reverse_lazy('machine-calibration-list', kwargs={'machine_id': machine_id})
        elif 'high frequency' in machine_type_name:
            form = HighFrequencyCalibrationForm(request.POST)
            template = 'calibrate/high_frequency_form.html'
            success_url = reverse_lazy('machine-calibration-list', kwargs={'machine_id': machine_id})
        elif 'low frequency' in machine_type_name:
            form = LowFrequencyCalibrationForm(request.POST)
            template = 'calibrate/low_frequency_form.html'
            success_url = reverse_lazy('machine-calibration-list', kwargs={'machine_id': machine_id})
        elif 'dial gauge' in machine_type_name:
            form = DialGaugeCalibrationForm(request.POST)
            template = 'calibrate/dial_gauge_form.html'
            success_url = reverse_lazy('machine-calibration-list', kwargs={'machine_id': machine_id})
        else:
            messages.error(request, 'ไม่พบประเภทการสอบเทียบที่เหมาะสมสำหรับเครื่องมือนี้')
            return redirect('machine-list')
        
        if form.is_valid():
            calibration = form.save(commit=False)
            # สำหรับ Balance, Microwave, High Frequency, Low Frequency, Dial Gauge ใช้ field 'machine' แทน 'uuc_id'
            if any(x in machine_type_name for x in ['balance', 'microwave', 'high frequency', 'low frequency', 'dial gauge']):
                calibration.machine = machine
            else:
                calibration.uuc_id = machine.id
            calibration.save()
            # ลบ success message ออกตามที่ผู้ใช้ต้องการ
            return redirect(success_url)
    else:
        if 'pressure' in machine_type_name:
            form = CalibrationPressureForm(initial={'uuc_id': machine.id})
            template = 'calibrate/pressure_form.html'
        elif 'torque' in machine_type_name:
            form = CalibrationTorqueForm(initial={'uuc_id': machine.id})
            template = 'calibrate/torque_form.html'
        elif 'balance' in machine_type_name:
            form = BalanceCalibrationForm(initial={'machine': machine.id})
            template = 'calibrate/balance_form.html'
        elif 'microwave' in machine_type_name:
            form = MicrowaveCalibrationForm(initial={'machine': machine.id})
            template = 'calibrate/microwave_form.html'
        elif 'high frequency' in machine_type_name:
            form = HighFrequencyCalibrationForm(initial={'machine': machine.id})
            template = 'calibrate/high_frequency_form.html'
        elif 'low frequency' in machine_type_name:
            form = LowFrequencyCalibrationForm(initial={'machine': machine.id})
            template = 'calibrate/low_frequency_form.html'
        elif 'dial gauge' in machine_type_name:
            form = DialGaugeCalibrationForm(initial={'machine': machine.id})
            template = 'calibrate/dial_gauge_form.html'
        else:
            messages.error(request, 'ไม่พบประเภทการสอบเทียบที่เหมาะสมสำหรับเครื่องมือนี้')
            return redirect('machine-list')
    
    context = {
        'form': form,
        'machine': machine,
        'calibration_type': machine_type_name,
        'machines': Machine.objects.filter(deleted=False).order_by('name'),
    }
    return render(request, template, context)

@login_required
def calibration_by_type(request, calibration_type):
    """แสดงรายการการสอบเทียบตามประเภท"""
    if calibration_type == 'pressure':
        calibrations = CalibrationPressure.objects.all()
        template = 'calibrate/pressure_list.html'
    elif calibration_type == 'torque':
        calibrations = CalibrationTorque.objects.all()
        template = 'calibrate/torque_list.html'
    else:
        messages.error(request, 'ประเภทการสอบเทียบไม่ถูกต้อง')
        return redirect('calibrate-dashboard')
    
    context = {
        'calibrations': calibrations,
        'calibration_type': calibration_type,
    }
    return render(request, template, context)

@login_required
def select_machine_for_calibration(request):
    """หน้าดึงข้อมูลเครื่องมือเพื่อบันทึกการสอบเทียบ"""
    machines = Machine.objects.filter(deleted=False).order_by('name')
    
    # กรองตามประเภท
    machine_type_filter = request.GET.get('machine_type')
    if machine_type_filter:
        machines = machines.filter(machine_type__name__icontains=machine_type_filter)
    
    # ค้นหาตามชื่อหรือ serial number
    search_query = request.GET.get('search')
    if search_query:
        machines = machines.filter(
            models.Q(name__icontains=search_query) |
            models.Q(serial_number__icontains=search_query) |
            models.Q(model__icontains=search_query)
        )
    
    context = {
        'machines': machines,
        'machine_types': MachineType.objects.all(),
        'selected_type': machine_type_filter,
        'search_query': search_query,
    }
    return render(request, 'calibrate/select_machine.html', context)

@login_required
def create_calibration_with_machine(request, machine_id):
    """สร้างบันทึกการสอบเทียบพร้อมข้อมูลเครื่องมือ"""
    machine = get_object_or_404(Machine, id=machine_id)
    machine_type_name = machine.machine_type.name.lower()
    
    if request.method == 'POST':
        if 'pressure' in machine_type_name:
            # สำหรับ Pressure ใช้การประมวลผลข้อมูลแบบพิเศษ
            return process_pressure_calibration(request, machine)
        elif 'torque' in machine_type_name:
            # สำหรับ Torque ใช้การประมวลผลข้อมูลแบบพิเศษ
            return process_torque_calibration(request, machine)
        elif 'balance' in machine_type_name:
            form = BalanceCalibrationForm(request.POST)
            template = 'calibrate/balance_form_with_machine.html'
            success_url = reverse_lazy('calibrate-dashboard')
        elif 'microwave' in machine_type_name:
            form = MicrowaveCalibrationForm(request.POST)
            template = 'calibrate/microwave_form_with_machine.html'
            success_url = reverse_lazy('calibrate-dashboard')
        elif 'high frequency' in machine_type_name:
            form = HighFrequencyCalibrationForm(request.POST)
            template = 'calibrate/high_frequency_form_with_machine.html'
            success_url = reverse_lazy('calibrate-dashboard')
        elif 'low frequency' in machine_type_name:
            form = LowFrequencyCalibrationForm(request.POST)
            template = 'calibrate/low_frequency_form_with_machine.html'
            success_url = reverse_lazy('calibrate-dashboard')
        elif 'dial gauge' in machine_type_name:
            form = DialGaugeCalibrationForm(request.POST)
            template = 'calibrate/dial_gauge_form_with_machine.html'
            success_url = reverse_lazy('calibrate-dashboard')
        else:
            messages.error(request, 'ไม่พบประเภทการสอบเทียบที่เหมาะสมสำหรับเครื่องมือนี้')
            return redirect('select-machine-for-calibration')
        
        if form.is_valid():
            calibration = form.save(commit=False)
            # สำหรับ Balance, Microwave, High Frequency, Low Frequency, Dial Gauge ใช้ field 'machine' แทน 'uuc_id'
            if any(x in machine_type_name for x in ['balance', 'microwave', 'high frequency', 'low frequency', 'dial gauge']):
                calibration.machine = machine
            else:
                calibration.uuc_id = machine
            
            # คำนวณวันที่ครบกำหนดถัดไป (+6 เดือน)
            from datetime import datetime, timedelta
            
            if calibration.update:
                # คำนวณ 6 เดือนจากวันที่สอบเทียบ
                year = calibration.update.year
                month = calibration.update.month + 6
                if month > 12:
                    year += month // 12
                    month = month % 12
                    if month == 0:
                        month = 12
                
                # ใช้วันที่เดิม แต่เปลี่ยนปีและเดือน
                day = min(calibration.update.day, 28)  # ป้องกันปัญหาเดือนกุมภาพันธ์
                try:
                    calibration.next_due = datetime(year, month, day).date()
                except ValueError:
                    # ถ้าวันที่ไม่ถูกต้อง ให้ใช้วันที่ 28 ของเดือนนั้น
                    calibration.next_due = datetime(year, month, 28).date()
            else:
                # ถ้าไม่มีวันที่สอบเทียบ ให้ใช้วันที่ปัจจุบัน + 6 เดือน
                today = datetime.now().date()
                year = today.year
                month = today.month + 6
                if month > 12:
                    year += month // 12
                    month = month % 12
                    if month == 0:
                        month = 12
                
                day = min(today.day, 28)
                try:
                    calibration.next_due = datetime(year, month, day).date()
                except ValueError:
                    calibration.next_due = datetime(year, month, 28).date()
            
            calibration.save()
            # ลบ success message ออกตามที่ผู้ใช้ต้องการ
            return redirect(success_url)
    else:
        # เติมข้อมูลเริ่มต้นจากเครื่องมือ
        initial_data = {
            'uuc_id': machine,
        }
        
        if 'pressure' in machine_type_name:
            form = CalibrationPressureForm(initial=initial_data)
            # เพิ่มตัวเลือกผู้ใช้สำหรับฟิลด์ calibrator และ certificate_issuer
            from django.contrib.auth import get_user_model
            User = get_user_model()
            users = User.objects.filter(is_active=True).order_by('first_name', 'last_name', 'username')
            form.fields['calibrator'].queryset = users
            form.fields['certificate_issuer'].queryset = users
            template = 'calibrate/pressure_form_with_machine.html'
        elif 'torque' in machine_type_name:
            form = CalibrationTorqueForm(initial=initial_data)
            # เพิ่มตัวเลือกผู้ใช้สำหรับฟิลด์ calibrator และ certificate_issuer
            from django.contrib.auth import get_user_model
            User = get_user_model()
            users = User.objects.filter(is_active=True).order_by('first_name', 'last_name', 'username')
            form.fields['calibrator'].queryset = users
            form.fields['certificate_issuer'].queryset = users
            template = 'calibrate/torque_form_with_machine.html'
        elif 'balance' in machine_type_name:
            # สำหรับ Balance ใช้ field 'machine' แทน 'uuc_id'
            balance_initial_data = {
                'machine': machine,
            }
            form = BalanceCalibrationForm(initial=balance_initial_data)
            # เพิ่มตัวเลือกผู้ใช้สำหรับฟิลด์ calibrator และ certificate_issuer
            from django.contrib.auth import get_user_model
            User = get_user_model()
            users = User.objects.filter(is_active=True).order_by('first_name', 'last_name', 'username')
            form.fields['calibrator'].queryset = users
            form.fields['certificate_issuer'].queryset = users
            template = 'calibrate/balance_form_with_machine.html'
        elif 'microwave' in machine_type_name:
            # สำหรับ Microwave ใช้ field 'machine' แทน 'uuc_id'
            microwave_initial_data = {
                'machine': machine,
            }
            form = MicrowaveCalibrationForm(initial=microwave_initial_data)
            # เพิ่มตัวเลือกผู้ใช้สำหรับฟิลด์ calibrator และ certificate_issuer
            from django.contrib.auth import get_user_model
            User = get_user_model()
            users = User.objects.filter(is_active=True).order_by('first_name', 'last_name', 'username')
            form.fields['calibrator'].queryset = users
            form.fields['certificate_issuer'].queryset = users
            template = 'calibrate/microwave_form_with_machine.html'
        elif 'high frequency' in machine_type_name:
            # สำหรับ High Frequency ใช้ field 'machine' แทน 'uuc_id'
            high_freq_initial_data = {
                'machine': machine,
            }
            form = HighFrequencyCalibrationForm(initial=high_freq_initial_data)
            # เพิ่มตัวเลือกผู้ใช้สำหรับฟิลด์ calibrator และ certificate_issuer
            from django.contrib.auth import get_user_model
            User = get_user_model()
            users = User.objects.filter(is_active=True).order_by('first_name', 'last_name', 'username')
            form.fields['calibrator'].queryset = users
            form.fields['certificate_issuer'].queryset = users
            template = 'calibrate/high_frequency_form_with_machine.html'
        elif 'low frequency' in machine_type_name:
            # สำหรับ Low Frequency ใช้ field 'machine' แทน 'uuc_id'
            low_freq_initial_data = {
                'machine': machine,
            }
            form = LowFrequencyCalibrationForm(initial=low_freq_initial_data)
            # เพิ่มตัวเลือกผู้ใช้สำหรับฟิลด์ calibrator และ certificate_issuer
            from django.contrib.auth import get_user_model
            User = get_user_model()
            users = User.objects.filter(is_active=True).order_by('first_name', 'last_name', 'username')
            form.fields['calibrator'].queryset = users
            form.fields['certificate_issuer'].queryset = users
            template = 'calibrate/low_frequency_form_with_machine.html'
        elif 'dial gauge' in machine_type_name:
            # สำหรับ Dial Gauge ใช้ field 'machine' แทน 'uuc_id'
            dial_gauge_initial_data = {
                'machine': machine,
            }
            form = DialGaugeCalibrationForm(initial=dial_gauge_initial_data)
            # เพิ่มตัวเลือกผู้ใช้สำหรับฟิลด์ calibrator และ certificate_issuer
            from django.contrib.auth import get_user_model
            User = get_user_model()
            users = User.objects.filter(is_active=True).order_by('first_name', 'last_name', 'username')
            form.fields['calibrator'].queryset = users
            form.fields['certificate_issuer'].queryset = users
            template = 'calibrate/dial_gauge_form_with_machine.html'
        else:
            messages.error(request, 'ไม่พบประเภทการสอบเทียบที่เหมาะสมสำหรับเครื่องมือนี้')
            return redirect('select-machine-for-calibration')
    
    context = {
        'form': form,
        'machine': machine,
        'calibration_type': machine_type_name,
        'machines': Machine.objects.filter(deleted=False).order_by('name'),
    }
    return render(request, template, context)

@login_required
def process_torque_calibration(request, machine):
    """ประมวลผลข้อมูลการสอบเทียบ Torque"""
    if request.method == 'POST':
        try:
            print("=== DEBUG: เริ่มประมวลผลข้อมูล Torque ===")
            print(f"Machine ID: {machine.id}")
            print(f"Machine Name: {machine.name}")
            
            # สร้าง CalibrationTorque object
            calibration = CalibrationTorque()
            calibration.uuc_id = machine
            
            # ฟังก์ชันช่วยแปลงข้อมูล
            def safe_float(value, default=0.0):
                if value == '' or value is None:
                    return default
                try:
                    return float(value)
                except (ValueError, TypeError):
                    return default
            
            # ข้อมูล CW (Clockwise) - แถวที่ 1
            cwset_1 = safe_float(request.POST.get('cwset_1', 0))
            cw0_1 = safe_float(request.POST.get('cw0_1', 0))
            cw90_1 = safe_float(request.POST.get('cw90_1', 0))
            cw180_1 = safe_float(request.POST.get('cw180_1', 0))
            cw270_1 = safe_float(request.POST.get('cw270_1', 0))
            cw_avg_1 = safe_float(request.POST.get('cw_avg_1', 0))
            cw_error_1 = safe_float(request.POST.get('cw_error_1', 0))
            cw_reading_1 = safe_float(request.POST.get('cw_reading_1', 0))
            cw_tolerance_1 = request.POST.get('cw_tolerance_1', '')
            
            print(f"CW Set: {cwset_1}")
            print(f"CW 0: {cw0_1}")
            print(f"CW 90: {cw90_1}")
            print(f"CW 180: {cw180_1}")
            print(f"CW 270: {cw270_1}")
            print(f"CW Avg: {cw_avg_1}")
            print(f"CW Error: {cw_error_1}")
            print(f"CW Reading: {cw_reading_1}")
            print(f"CW Tolerance: {cw_tolerance_1}")
            
            calibration.cwset = cwset_1
            calibration.cw0 = cw0_1
            calibration.cw90 = cw90_1
            calibration.cw180 = cw180_1
            calibration.cw270 = cw270_1
            calibration.cw_avg = cw_avg_1
            calibration.cw_error = cw_error_1
            calibration.cw_reading = cw_reading_1
            
            # แยก tolerance string เป็น start และ end
            if ' - ' in cw_tolerance_1:
                cw_parts = cw_tolerance_1.split(' - ')
                calibration.cw_tolerance_start = safe_float(cw_parts[0]) if cw_parts[0] else None
                calibration.cw_tolerance_end = safe_float(cw_parts[1]) if cw_parts[1] else None
            
            # ข้อมูล CCW (Counter-Clockwise) - แถวที่ 1
            ccwset_1 = safe_float(request.POST.get('ccwset_1', 0))
            ccw0_1 = safe_float(request.POST.get('ccw0_1', 0))
            ccw90_1 = safe_float(request.POST.get('ccw90_1', 0))
            ccw180_1 = safe_float(request.POST.get('ccw180_1', 0))
            ccw270_1 = safe_float(request.POST.get('ccw270_1', 0))
            ccw_avg_1 = safe_float(request.POST.get('ccw_avg_1', 0))
            ccw_error_1 = safe_float(request.POST.get('ccw_error_1', 0))
            ccw_reading_1 = safe_float(request.POST.get('ccw_reading_1', 0))
            ccw_tolerance_1 = request.POST.get('ccw_tolerance_1', '')
            
            print(f"CCW Set: {ccwset_1}")
            print(f"CCW 0: {ccw0_1}")
            print(f"CCW 90: {ccw90_1}")
            print(f"CCW 180: {ccw180_1}")
            print(f"CCW 270: {ccw270_1}")
            print(f"CCW Avg: {ccw_avg_1}")
            print(f"CCW Error: {ccw_error_1}")
            print(f"CCW Reading: {ccw_reading_1}")
            print(f"CCW Tolerance: {ccw_tolerance_1}")
            
            calibration.ccwset = ccwset_1
            calibration.ccw0 = ccw0_1
            calibration.ccw90 = ccw90_1
            calibration.ccw180 = ccw180_1
            calibration.ccw270 = ccw270_1
            calibration.ccw_avg = ccw_avg_1
            calibration.ccw_error = ccw_error_1
            calibration.ccw_reading = ccw_reading_1
            
            # แยก tolerance string เป็น start และ end
            if ' - ' in ccw_tolerance_1:
                ccw_parts = ccw_tolerance_1.split(' - ')
                calibration.ccw_tolerance_start = safe_float(ccw_parts[0]) if ccw_parts[0] else None
                calibration.ccw_tolerance_end = safe_float(ccw_parts[1]) if ccw_parts[1] else None
            
            # วันที่สอบเทียบ
            update_date = request.POST.get('update')
            if update_date:
                calibration.update = update_date
                print(f"Update Date: {update_date}")
            
            # คำนวณวันที่ครบกำหนดถัดไป (+6 เดือน)
            from datetime import datetime, timedelta
            
            if calibration.update:
                # แปลง string เป็น date object
                try:
                    update_date_obj = datetime.strptime(calibration.update, '%Y-%m-%d').date()
                    # คำนวณ 6 เดือนจากวันที่สอบเทียบ
                    year = update_date_obj.year
                    month = update_date_obj.month + 6
                    if month > 12:
                        year += month // 12
                        month = month % 12
                        if month == 0:
                            month = 12
                    
                    # ใช้วันที่เดิม แต่เปลี่ยนปีและเดือน
                    day = min(update_date_obj.day, 28)  # ป้องกันปัญหาเดือนกุมภาพันธ์
                    try:
                        calibration.next_due = datetime(year, month, day).date()
                    except ValueError:
                        # ถ้าวันที่ไม่ถูกต้อง ให้ใช้วันที่ 28 ของเดือนนั้น
                        calibration.next_due = datetime(year, month, 28).date()
                    print(f"Next Due Date: {calibration.next_due}")
                except ValueError as e:
                    print(f"Error parsing date: {e}")
                    # ถ้าแปลงวันที่ไม่ได้ ให้ใช้วันที่ปัจจุบัน + 6 เดือน
                    today = datetime.now().date()
                    year = today.year
                    month = today.month + 6
                    if month > 12:
                        year += month // 12
                        month = month % 12
                        if month == 0:
                            month = 12
                    
                    day = min(today.day, 28)
                    try:
                        calibration.next_due = datetime(year, month, day).date()
                    except ValueError:
                        calibration.next_due = datetime(year, month, 28).date()
                    print(f"Next Due Date (default): {calibration.next_due}")
            else:
                # ถ้าไม่มีวันที่สอบเทียบ ให้ใช้วันที่ปัจจุบัน + 6 เดือน
                today = datetime.now().date()
                year = today.year
                month = today.month + 6
                if month > 12:
                    year += month // 12
                    month = month % 12
                    if month == 0:
                        month = 12
                
                day = min(today.day, 28)
                try:
                    calibration.next_due = datetime(year, month, day).date()
                except ValueError:
                    calibration.next_due = datetime(year, month, 28).date()
                print(f"Next Due Date (default): {calibration.next_due}")
            
            # ข้อมูลเพิ่มเติม
            status = request.POST.get('status', 'not_set')
            priority = request.POST.get('priority', 'normal')
            std_id = request.POST.get('std_id')
            calibrator_id = request.POST.get('calibrator')
            certificate_issuer_id = request.POST.get('certificate_issuer')
            
            calibration.status = status
            calibration.priority = priority
            if std_id and std_id != '':
                calibration.std_id_id = int(std_id)
            if calibrator_id and calibrator_id != '':
                calibration.calibrator_id = int(calibrator_id)
            if certificate_issuer_id and certificate_issuer_id != '':
                calibration.certificate_issuer_id = int(certificate_issuer_id)
            
            print("=== DEBUG: บันทึกข้อมูล ===")
            calibration.save()
            print(f"บันทึกสำเร็จ! ID: {calibration.cal_torque_id}")
            
            # ลบ success message ออกตามที่ผู้ใช้ต้องการ
            return redirect('calibrate-dashboard')
            
        except Exception as e:
            print(f"=== DEBUG: เกิดข้อผิดพลาด ===")
            print(f"Error: {str(e)}")
            import traceback
            traceback.print_exc()
            messages.error(request, f'เกิดข้อผิดพลาดในการบันทึกข้อมูล: {str(e)}')
            return redirect('select-machine-for-calibration')
    
    # ถ้าไม่ใช่ POST request
    messages.error(request, 'วิธีการส่งข้อมูลไม่ถูกต้อง')
    return redirect('select-machine-for-calibration')

@login_required
def process_pressure_calibration(request, machine):
    """ประมวลผลข้อมูลการสอบเทียบ Pressure"""
    if request.method == 'POST':
        try:
            print("=== DEBUG: เริ่มประมวลผลข้อมูล Pressure ===")
            print(f"Machine ID: {machine.id}")
            print(f"Machine Name: {machine.name}")
            
            # สร้าง CalibrationPressure object
            calibration = CalibrationPressure()
            calibration.uuc_id = machine
            
            # ฟังก์ชันช่วยแปลงข้อมูล
            def safe_float(value, default=0.0):
                if value == '' or value is None:
                    return default
                try:
                    return float(value)
                except (ValueError, TypeError):
                    return default
            
            # ข้อมูลจากแถวที่ 1
            set_1 = safe_float(request.POST.get('set_1', 0))
            m1_1 = safe_float(request.POST.get('m1_1', 0))
            m2_1 = safe_float(request.POST.get('m2_1', 0))
            m3_1 = safe_float(request.POST.get('m3_1', 0))
            m4_1 = safe_float(request.POST.get('m4_1', 0))
            avg_1 = safe_float(request.POST.get('avg_1', 0))
            error_1 = safe_float(request.POST.get('error_1', 0))
            tolerance_1 = request.POST.get('tolerance_1', '')
            
            print(f"Set: {set_1}")
            print(f"M1: {m1_1}")
            print(f"M2: {m2_1}")
            print(f"M3: {m3_1}")
            print(f"M4: {m4_1}")
            print(f"AVG: {avg_1}")
            print(f"Error: {error_1}")
            print(f"Tolerance: {tolerance_1}")
            
            # ใช้ข้อมูลจากแถวแรกเป็นหลัก
            calibration.set = str(set_1) if set_1 != 0 else ''
            calibration.m1 = str(m1_1) if m1_1 != 0 else ''
            calibration.m2 = str(m2_1) if m2_1 != 0 else ''
            calibration.m3 = str(m3_1) if m3_1 != 0 else ''
            calibration.m4 = str(m4_1) if m4_1 != 0 else ''
            calibration.avg = avg_1
            calibration.error = error_1
            
            # แยก tolerance string เป็น start และ end
            if ' - ' in tolerance_1:
                tolerance_parts = tolerance_1.split(' - ')
                calibration.tolerance_start = safe_float(tolerance_parts[0]) if tolerance_parts[0] else None
                calibration.tolerance_end = safe_float(tolerance_parts[1]) if tolerance_parts[1] else None
            
            # ข้อมูลทั่วไป
            update_date = request.POST.get('update')
            if update_date:
                calibration.update = update_date
                print(f"Update Date: {update_date}")
            
            # คำนวณวันที่ครบกำหนดถัดไป (+6 เดือน)
            from datetime import datetime, timedelta
            
            if calibration.update:
                # แปลง string เป็น date object
                try:
                    update_date_obj = datetime.strptime(calibration.update, '%Y-%m-%d').date()
                    # คำนวณ 6 เดือนจากวันที่สอบเทียบ
                    year = update_date_obj.year
                    month = update_date_obj.month + 6
                    if month > 12:
                        year += month // 12
                        month = month % 12
                        if month == 0:
                            month = 12
                    
                    # ใช้วันที่เดิม แต่เปลี่ยนปีและเดือน
                    day = min(update_date_obj.day, 28)  # ป้องกันปัญหาเดือนกุมภาพันธ์
                    try:
                        calibration.next_due = datetime(year, month, day).date()
                    except ValueError:
                        # ถ้าวันที่ไม่ถูกต้อง ให้ใช้วันที่ 28 ของเดือนนั้น
                        calibration.next_due = datetime(year, month, 28).date()
                    print(f"Next Due Date: {calibration.next_due}")
                except ValueError as e:
                    print(f"Error parsing date: {e}")
                    # ถ้าแปลงวันที่ไม่ได้ ให้ใช้วันที่ปัจจุบัน + 6 เดือน
                    today = datetime.now().date()
                    year = today.year
                    month = today.month + 6
                    if month > 12:
                        year += month // 12
                        month = month % 12
                        if month == 0:
                            month = 12
                    
                    day = min(today.day, 28)
                    try:
                        calibration.next_due = datetime(year, month, day).date()
                    except ValueError:
                        calibration.next_due = datetime(year, month, 28).date()
                    print(f"Next Due Date (default): {calibration.next_due}")
            else:
                # ถ้าไม่มีวันที่สอบเทียบ ให้ใช้วันที่ปัจจุบัน + 6 เดือน
                today = datetime.now().date()
                year = today.year
                month = today.month + 6
                if month > 12:
                    year += month // 12
                    month = month % 12
                    if month == 0:
                        month = 12
                
                day = min(today.day, 28)
                try:
                    calibration.next_due = datetime(year, month, day).date()
                except ValueError:
                    calibration.next_due = datetime(year, month, 28).date()
                print(f"Next Due Date (default): {calibration.next_due}")
            
            # ข้อมูลเพิ่มเติม
            status = request.POST.get('status', 'not_set')
            priority = request.POST.get('priority', 'normal')
            std_id = request.POST.get('std_id')
            calibrator_id = request.POST.get('calibrator')
            certificate_issuer_id = request.POST.get('certificate_issuer')
            
            calibration.status = status
            calibration.priority = priority
            if std_id and std_id != '':
                calibration.std_id_id = int(std_id)
            if calibrator_id and calibrator_id != '':
                calibration.calibrator_id = int(calibrator_id)
            if certificate_issuer_id and certificate_issuer_id != '':
                calibration.certificate_issuer_id = int(certificate_issuer_id)
            
            print("=== DEBUG: บันทึกข้อมูล ===")
            calibration.save()
            print(f"บันทึกสำเร็จ! ID: {calibration.cal_pressure_id}")
            
            # ลบ success message ออกตามที่ผู้ใช้ต้องการ
            return redirect('calibrate-dashboard')
            
        except Exception as e:
            print(f"=== DEBUG: เกิดข้อผิดพลาด ===")
            print(f"Error: {str(e)}")
            import traceback
            traceback.print_exc()
            messages.error(request, f'เกิดข้อผิดพลาดในการบันทึกข้อมูล: {str(e)}')
            return redirect('select-machine-for-calibration')
    
    # ถ้าไม่ใช่ POST request
    messages.error(request, 'วิธีการส่งข้อมูลไม่ถูกต้อง')
    return redirect('select-machine-for-calibration')

@login_required
def calibration_report(request):
    """หน้ารายงานสอบเทียบ"""
    from datetime import date, timedelta
    
    # ดึงข้อมูลการสอบเทียบทั้งหมด
    pressure_calibrations = CalibrationPressure.objects.select_related('uuc_id', 'std_id', 'calibrator', 'certificate_issuer').all()
    torque_calibrations = CalibrationTorque.objects.select_related('uuc_id', 'std_id', 'calibrator', 'certificate_issuer').all()
    balance_calibrations = BalanceCalibration.objects.select_related('machine', 'std_id', 'calibrator', 'certificate_issuer').all()
    high_frequency_calibrations = HighFrequencyCalibration.objects.select_related('machine', 'std_id', 'calibrator', 'certificate_issuer').all()
    low_frequency_calibrations = LowFrequencyCalibration.objects.select_related('machine', 'std_id', 'calibrator', 'certificate_issuer').all()
    microwave_calibrations = MicrowaveCalibration.objects.select_related('machine', 'std_id', 'calibrator', 'certificate_issuer').all()
    dial_gauge_calibrations = DialGaugeCalibration.objects.select_related('machine', 'std_id', 'calibrator', 'certificate_issuer').all()
    force_calibrations = []  # Empty list since CalibrationForce model no longer exists
    
    # รวมข้อมูลการสอบเทียบทั้งหมด
    all_calibrations = []
    
    
    # เพิ่มข้อมูล Pressure calibrations
    for cal in pressure_calibrations:
        all_calibrations.append({
            'id': cal.cal_pressure_id,
            'type': 'pressure',
            'type_name': 'การสอบเทียบ Pressure',
            'machine_name': cal.uuc_id.name if cal.uuc_id else '-',
            'machine_model': cal.uuc_id.model if cal.uuc_id else '-',
            'serial_number': cal.uuc_id.serial_number if cal.uuc_id else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.update,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.update,  # วันที่สอบเทียบ
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.uuc_id.organize.name if cal.uuc_id and cal.uuc_id.organize else '-',
        })
    
    # เพิ่มข้อมูล Torque calibrations
    for cal in torque_calibrations:
        all_calibrations.append({
            'id': cal.cal_torque_id,
            'type': 'torque',
            'type_name': 'การสอบเทียบ Torque',
            'machine_name': cal.uuc_id.name if cal.uuc_id else '-',
            'machine_model': cal.uuc_id.model if cal.uuc_id else '-',
            'serial_number': cal.uuc_id.serial_number if cal.uuc_id else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.update,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.update,  # วันที่สอบเทียบ
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.uuc_id.organize.name if cal.uuc_id and cal.uuc_id.organize else '-',
        })
    
    # เพิ่มข้อมูล Balance calibrations
    for cal in balance_calibrations:
        all_calibrations.append({
            'id': cal.id,
            'type': 'balance',
            'type_name': 'การสอบเทียบ Balance',
            'machine_name': cal.machine.name if cal.machine else '-',
            'machine_model': cal.machine.model if cal.machine else '-',
            'serial_number': cal.machine.serial_number if cal.machine else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.date_calibration,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.date_calibration,  # วันที่สอบเทียบ
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.machine.organize.name if cal.machine and cal.machine.organize else '-',
        })
    
    # เพิ่มข้อมูล High Frequency calibrations
    for cal in high_frequency_calibrations:
        all_calibrations.append({
            'id': cal.id,
            'type': 'high-frequency',
            'type_name': 'การสอบเทียบ High Frequency',
            'machine_name': cal.machine.name if cal.machine else '-',
            'machine_model': cal.machine.model if cal.machine else '-',
            'serial_number': cal.machine.serial_number if cal.machine else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.date_calibration,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.date_calibration,
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.machine.organize.name if cal.machine and cal.machine.organize else '-',
        })
    
    # เพิ่มข้อมูล Low Frequency calibrations
    for cal in low_frequency_calibrations:
        all_calibrations.append({
            'id': cal.id,
            'type': 'low-frequency',
            'type_name': 'การสอบเทียบ Low Frequency',
            'machine_name': cal.machine.name if cal.machine else '-',
            'machine_model': cal.machine.model if cal.machine else '-',
            'serial_number': cal.machine.serial_number if cal.machine else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.date_calibration,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.date_calibration,
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.machine.organize.name if cal.machine and cal.machine.organize else '-',
        })
    
    # เพิ่มข้อมูล Microwave calibrations
    for cal in microwave_calibrations:
        all_calibrations.append({
            'id': cal.id,
            'type': 'microwave',
            'type_name': 'การสอบเทียบ Microwave',
            'machine_name': cal.machine.name if cal.machine else '-',
            'machine_model': cal.machine.model if cal.machine else '-',
            'serial_number': cal.machine.serial_number if cal.machine else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.date_calibration,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.date_calibration,
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.machine.organize.name if cal.machine and cal.machine.organize else '-',
        })
    
    # เพิ่มข้อมูล Dial Gauge calibrations
    for cal in dial_gauge_calibrations:
        all_calibrations.append({
            'id': cal.id,
            'type': 'dial_gauge',
            'type_name': 'การสอบเทียบ Dial Gauge',
            'machine_name': cal.machine.name if cal.machine else '-',
            'machine_model': cal.machine.model if cal.machine else '-',
            'serial_number': cal.machine.serial_number if cal.machine else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.date_calibration,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.date_calibration,
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.machine.organize.name if cal.machine and cal.machine.organize else '-',
        })
    
    # เรียงลำดับตามวันที่สอบเทียบล่าสุด
    all_calibrations.sort(key=lambda x: x['calibration_date'] if x['calibration_date'] else date.min, reverse=True)
    
    # ข้อมูลวันที่สำหรับการคำนวณสถานะ
    today = date.today()
    today_plus_30 = today + timedelta(days=30)
    
    context = {
        'pressure_machines': Machine.objects.filter(machine_type__name__icontains='pressure').count(),
        'torque_machines': Machine.objects.filter(machine_type__name__icontains='torque').count(),
        'total_calibrations': (
            CalibrationPressure.objects.count() +
            CalibrationTorque.objects.count()
        ),
        'all_calibrations': all_calibrations,
        'today': today,
        'today_plus_30': today_plus_30,
    }
    return render(request, 'calibrate/calibration_report.html', context)

@login_required
def calibration_report_detail(request):
    """หน้ารายงานผลสอบเทียบแบบละเอียด"""
    from datetime import date, timedelta
    
    # รับพารามิเตอร์การกรอง
    department = request.GET.get('department', '')
    instrument_type = request.GET.get('instrument_type', '')
    start_date = request.GET.get('start_date', '')
    end_date = request.GET.get('end_date', '')
    serial_search = request.GET.get('serial_search', '')
    name_search = request.GET.get('name_search', '')
    status_filter = request.GET.get('status_filter', '')
    
    # ดึงข้อมูลหน่วยงาน (Organize) สำหรับที่อยู่
    from organize.models import Organize
    try:
        # หาหน่วยงานหลัก (main unit) หรือหน่วยงานแรก
        organization = Organize.objects.filter(is_main_unit=True).first()
        if not organization:
            organization = Organize.objects.first()
    except:
        organization = None
    
    # ดึงข้อมูลการสอบเทียบทั้งหมด
    pressure_calibrations = CalibrationPressure.objects.select_related('uuc_id', 'std_id', 'calibrator', 'certificate_issuer').all()
    torque_calibrations = CalibrationTorque.objects.select_related('uuc_id', 'std_id', 'calibrator', 'certificate_issuer').all()
    balance_calibrations = BalanceCalibration.objects.select_related('machine', 'std_id', 'calibrator', 'certificate_issuer').all()
    microwave_calibrations = MicrowaveCalibration.objects.select_related('machine', 'std_id', 'calibrator', 'certificate_issuer').all()
    high_frequency_calibrations = HighFrequencyCalibration.objects.select_related('machine', 'std_id', 'calibrator', 'certificate_issuer').all()
    low_frequency_calibrations = LowFrequencyCalibration.objects.select_related('machine', 'std_id', 'calibrator', 'certificate_issuer').all()
    dial_gauge_calibrations = DialGaugeCalibration.objects.select_related('machine', 'std_id', 'calibrator', 'certificate_issuer').all()
    force_calibrations = []  # Empty list since CalibrationForce model no longer exists
    
    # รวมข้อมูลการสอบเทียบทั้งหมด
    all_calibrations = []
    
    # เพิ่มข้อมูล Force calibrations
    for cal in force_calibrations:
        all_calibrations.append({
            'id': cal.cal_force_id,
            'type': 'force',
            'type_name': 'การสอบเทียบแรง',
            'machine_name': cal.uuc_id.name if cal.uuc_id else '-',
            'machine_model': cal.uuc_id.model if cal.uuc_id else '-',
            'serial_number': cal.uuc_id.serial_number if cal.uuc_id else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.update,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.update,  # ǹ��ͺ�º
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.uuc_id.organize.name if cal.uuc_id and cal.uuc_id.organize else '-',
        })
    
    # เพิ่มข้อมูล Pressure calibrations
    for cal in pressure_calibrations:
        all_calibrations.append({
            'id': cal.cal_pressure_id,
            'type': 'pressure',
            'type_name': 'การสอบเทียบ Pressure',
            'machine_name': cal.uuc_id.name if cal.uuc_id else '-',
            'machine_model': cal.uuc_id.model if cal.uuc_id else '-',
            'serial_number': cal.uuc_id.serial_number if cal.uuc_id else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.update,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.update,  # ǹ��ͺ�º
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.uuc_id.organize.name if cal.uuc_id and cal.uuc_id.organize else '-',
        })
    
    # เพิ่มข้อมูล Torque calibrations
    for cal in torque_calibrations:
        all_calibrations.append({
            'id': cal.cal_torque_id,
            'type': 'torque',
            'type_name': 'การสอบเทียบ Torque',
            'machine_name': cal.uuc_id.name if cal.uuc_id else '-',
            'machine_model': cal.uuc_id.model if cal.uuc_id else '-',
            'serial_number': cal.uuc_id.serial_number if cal.uuc_id else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.update,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.update,  # ǹ��ͺ�º
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.uuc_id.organize.name if cal.uuc_id and cal.uuc_id.organize else '-',
        })
    
    # เพิ่มข้อมูล Balance calibrations
    for cal in balance_calibrations:
        all_calibrations.append({
            'id': cal.id,
            'type': 'balance',
            'type_name': 'การสอบเทียบ Balance',
            'machine_name': cal.machine.name if cal.machine else '-',
            'machine_model': cal.machine.model if cal.machine else '-',
            'serial_number': cal.machine.serial_number if cal.machine else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.date_calibration,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.date_calibration,  # วันที่สอบเทียบ
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.machine.organize.name if cal.machine and cal.machine.organize else '-',
        })
    
    # เพิ่มข้อมูล Microwave calibrations
    for cal in microwave_calibrations:
        all_calibrations.append({
            'id': cal.id,
            'type': 'microwave',
            'type_name': 'การสอบเทียบ Microwave',
            'machine_name': cal.machine.name if cal.machine else '-',
            'machine_model': cal.machine.model if cal.machine else '-',
            'serial_number': cal.machine.serial_number if cal.machine else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.update or cal.date_calibration,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.date_calibration,  # วันที่สอบเทียบ
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.machine.organize.name if cal.machine and cal.machine.organize else '-',
        })
    
    # เพิ่มข้อมูล High Frequency calibrations
    for cal in high_frequency_calibrations:
        all_calibrations.append({
            'id': cal.id,
            'type': 'high-frequency',
            'type_name': 'การสอบเทียบ High Frequency',
            'machine_name': cal.machine.name if cal.machine else '-',
            'machine_model': cal.machine.model if cal.machine else '-',
            'serial_number': cal.machine.serial_number if cal.machine else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.date_calibration,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.date_calibration,  # วันที่สอบเทียบ
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.machine.organize.name if cal.machine and cal.machine.organize else '-',
        })
    
    # เพิ่มข้อมูล Low Frequency calibrations
    for cal in low_frequency_calibrations:
        all_calibrations.append({
            'id': cal.id,
            'type': 'low-frequency',
            'type_name': 'การสอบเทียบ Low Frequency',
            'machine_name': cal.machine.name if cal.machine else '-',
            'machine_model': cal.machine.model if cal.machine else '-',
            'serial_number': cal.machine.serial_number if cal.machine else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.date_calibration,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.date_calibration,  # วันที่สอบเทียบ
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.machine.organize.name if cal.machine and cal.machine.organize else '-',
        })
    
    # เพิ่มข้อมูล Dial Gauge calibrations
    for cal in dial_gauge_calibrations:
        all_calibrations.append({
            'id': cal.id,
            'type': 'dial_gauge',
            'type_name': 'การสอบเทียบ Dial Gauge',
            'machine_name': cal.machine.name if cal.machine else '-',
            'machine_model': cal.machine.model if cal.machine else '-',
            'serial_number': cal.machine.serial_number if cal.machine else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.date_calibration,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.date_calibration,  # วันที่สอบเทียบ
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.machine.organize.name if cal.machine and cal.machine.organize else '-',
        })
    
    # กรองข้อมูลตามพารามิเตอร์
    filtered_calibrations = []
    for cal in all_calibrations:
        include_item = True
        
        # กรองตาม Serial Number
        if serial_search and serial_search.lower() not in cal['serial_number'].lower():
            include_item = False
        
        # กรองตามชื่อเครื่องมือ
        if name_search and name_search.lower() not in cal['machine_name'].lower():
            include_item = False
        
        # กรองตามประเภทเครื่องมือ
        if instrument_type:
            if instrument_type == 'force' and cal['type'] != 'force':
                include_item = False
            elif instrument_type == 'pressure' and cal['type'] != 'pressure':
                include_item = False
            elif instrument_type == 'torque' and cal['type'] != 'torque':
                include_item = False
            elif instrument_type == 'balance' and cal['type'] != 'balance':
                include_item = False
        
        # กรองตามสถานะ
        if status_filter:
            if status_filter == 'cert_issued' and cal['status'] not in ['cert_issued', 'passed', 'active']:
                include_item = False
            elif status_filter == 'failed' and cal['status'] not in ['failed', 'pending', 'in_progress', 'not_set']:
                include_item = False
        
        # กรองตามวันที่
        if start_date and cal['update_date']:
            try:
                start_date_obj = datetime.strptime(start_date, '%Y-%m-%d').date()
                if cal['update_date'] < start_date_obj:
                    include_item = False
            except:
                pass
        
        if end_date and cal['update_date']:
            try:
                end_date_obj = datetime.strptime(end_date, '%Y-%m-%d').date()
                if cal['update_date'] > end_date_obj:
                    include_item = False
            except:
                pass
        
        if include_item:
            filtered_calibrations.append(cal)
    
    # เรียงลำดับตามวันที่สอบเทียบล่าสุด
    filtered_calibrations.sort(key=lambda x: x['calibration_date'] if x['calibration_date'] else date.min, reverse=True)
    
    # ข้อมูลวันที่สำหรับการคำนวณสถานะ
    today = date.today()
    today_plus_30 = today + timedelta(days=30)
    
    context = {
        'pressure_machines': Machine.objects.filter(machine_type__name__icontains='pressure').count(),
        'torque_machines': Machine.objects.filter(machine_type__name__icontains='torque').count(),
        'total_calibrations': (
            CalibrationPressure.objects.count() +
            CalibrationTorque.objects.count() +
            BalanceCalibration.objects.count() +
            MicrowaveCalibration.objects.count() +
            HighFrequencyCalibration.objects.count() +
            LowFrequencyCalibration.objects.count() +
            DialGaugeCalibration.objects.count()
        ),
        'all_calibrations': filtered_calibrations,
        'today': today,
        'today_plus_30': today_plus_30,
        'organization': organization,  # เพิ่มข้อมูลหน่วยงาน
        # พารามิเตอร์การกรอง
        'department': department,
        'instrument_type': instrument_type,
        'start_date': start_date,
        'end_date': end_date,
        'serial_search': serial_search,
        'name_search': name_search,
        'status_filter': status_filter,
    }
    return render(request, 'calibrate/calibration_report_detail.html', context)

@login_required
def export_to_word(request):
    """Export รายงานสอบเทียบเป็นไฟล์ Word"""
    from datetime import date, timedelta
    
    # รับพารามิเตอร์การกรอง (เหมือนกับใน calibration_report_detail)
    department = request.GET.get('department', '')
    instrument_type = request.GET.get('instrument_type', '')
    start_date = request.GET.get('start_date', '')
    end_date = request.GET.get('end_date', '')
    serial_search = request.GET.get('serial_search', '')
    name_search = request.GET.get('name_search', '')
    status_filter = request.GET.get('status_filter', '')
    
    # ดึงข้อมูลหน่วยงาน (Organize) สำหรับที่อยู่
    from organize.models import Organize
    try:
        # หาหน่วยงานหลัก (main unit) หรือหน่วยงานแรก
        organization = Organize.objects.filter(is_main_unit=True).first()
        if not organization:
            organization = Organize.objects.first()
    except:
        organization = None
    
    # ดึงข้อมูลการสอบเทียบทั้งหมด
    pressure_calibrations = CalibrationPressure.objects.select_related('uuc_id', 'std_id', 'calibrator', 'certificate_issuer').all()
    torque_calibrations = CalibrationTorque.objects.select_related('uuc_id', 'std_id', 'calibrator', 'certificate_issuer').all()
    
    # รวมข้อมูลการสอบเทียบทั้งหมด
    all_calibrations = []
    
    # เพิ่มข้อมูล Force calibrations
    for cal in force_calibrations:
        all_calibrations.append({
            'id': cal.cal_force_id,
            'type': 'force',
            'type_name': 'การสอบเทียบแรง',
            'machine_name': cal.uuc_id.name if cal.uuc_id else '-',
            'machine_model': cal.uuc_id.model if cal.uuc_id else '-',
            'serial_number': cal.uuc_id.serial_number if cal.uuc_id else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.update,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.update,  # ǹ��ͺ�º
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.uuc_id.organize.name if cal.uuc_id and cal.uuc_id.organize else '-',
        })
    
    # เพิ่มข้อมูล Pressure calibrations
    for cal in pressure_calibrations:
        all_calibrations.append({
            'id': cal.cal_pressure_id,
            'type': 'pressure',
            'type_name': 'การสอบเทียบ Pressure',
            'machine_name': cal.uuc_id.name if cal.uuc_id else '-',
            'machine_model': cal.uuc_id.model if cal.uuc_id else '-',
            'serial_number': cal.uuc_id.serial_number if cal.uuc_id else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.update,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.update,  # ǹ��ͺ�º
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.uuc_id.organize.name if cal.uuc_id and cal.uuc_id.organize else '-',
        })
    
    # เพิ่มข้อมูล Torque calibrations
    for cal in torque_calibrations:
        all_calibrations.append({
            'id': cal.cal_torque_id,
            'type': 'torque',
            'type_name': 'การสอบเทียบ Torque',
            'machine_name': cal.uuc_id.name if cal.uuc_id else '-',
            'machine_model': cal.uuc_id.model if cal.uuc_id else '-',
            'serial_number': cal.uuc_id.serial_number if cal.uuc_id else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.update,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.update,  # ǹ��ͺ�º
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.uuc_id.organize.name if cal.uuc_id and cal.uuc_id.organize else '-',
        })
    
    # กรองข้อมูลตามพารามิเตอร์ (เหมือนกับใน calibration_report_detail)
    filtered_calibrations = []
    for cal in all_calibrations:
        include_item = True
        
        # กรองตาม Serial Number
        if serial_search and serial_search.lower() not in cal['serial_number'].lower():
            include_item = False
        
        # กรองตามชื่อเครื่องมือ
        if name_search and name_search.lower() not in cal['machine_name'].lower():
            include_item = False
        
        # กรองตามประเภทเครื่องมือ
        if instrument_type:
            if instrument_type == 'force' and cal['type'] != 'force':
                include_item = False
            elif instrument_type == 'pressure' and cal['type'] != 'pressure':
                include_item = False
            elif instrument_type == 'torque' and cal['type'] != 'torque':
                include_item = False
            elif instrument_type == 'balance' and cal['type'] != 'balance':
                include_item = False
        
        # กรองตามสถานะ
        if status_filter:
            if status_filter == 'cert_issued' and cal['status'] not in ['cert_issued', 'passed', 'active']:
                include_item = False
            elif status_filter == 'failed' and cal['status'] not in ['failed', 'pending', 'in_progress', 'not_set']:
                include_item = False
        
        # กรองตามวันที่
        if start_date and cal['update_date']:
            try:
                start_date_obj = datetime.strptime(start_date, '%Y-%m-%d').date()
                if cal['update_date'] < start_date_obj:
                    include_item = False
            except:
                pass
        
        if end_date and cal['update_date']:
            try:
                end_date_obj = datetime.strptime(end_date, '%Y-%m-%d').date()
                if cal['update_date'] > end_date_obj:
                    include_item = False
            except:
                pass
        
        if include_item:
            filtered_calibrations.append(cal)
    
    # เรียงลำดับตามวันที่สอบเทียบล่าสุด
    filtered_calibrations.sort(key=lambda x: x['calibration_date'] if x['calibration_date'] else date.min, reverse=True)
    
    # สร้าง Word document
    doc = Document()
    
    # ตั้งค่าหน้ากระดาษ
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # หัวเรื่อง
    title = doc.add_heading('รายงานสอบเทียบ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # วันที่รายงาน
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_paragraph.add_run(f'วันที่: {datetime.now().strftime("%d/%m/%Y")}')
    date_run.font.size = Pt(12)
    
    # เพิ่มบรรทัดว่าง
    doc.add_paragraph()
    
    # สร้างตาราง
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # หัวตาราง
    header_cells = table.rows[0].cells
    headers = ['ลำดับ', 'Model', 'SERIAL No.', 'ชื่อเครื่องวัด', 'หน่วยผู้ใช้', 'จำนวน', 'หมายเหตุ']
    
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # เพิ่มข้อมูลในตาราง
    for i, cal in enumerate(filtered_calibrations, 1):
        row_cells = table.add_row().cells
        
        # ลำดับ
        row_cells[0].text = str(i)
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Model
        row_cells[1].text = cal['machine_model'] if cal['machine_model'] != '-' else '-'
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # SERIAL No.
        row_cells[2].text = cal['serial_number'] if cal['serial_number'] != '-' else '-'
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # ชื่อเครื่องวัด
        if cal['type'] == 'pressure':
            instrument_name = 'Pressure Gauge'
        elif cal['type'] == 'torque':
            instrument_name = 'Torque Wrench'
        else:
            instrument_name = cal['type_name']
        
        row_cells[3].text = instrument_name
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # หน่วยผู้ใช้
        user_unit = organization.name if organization else 'ยังไม่ได้กรอก'
        row_cells[4].text = user_unit
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # จำนวน
        row_cells[5].text = '๑'
        row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # หมายเหตุ
        if cal['status'] in ['cert_issued', 'passed', 'active']:
            remark = 'ออกใบรับรอง'
        elif cal['status'] in ['failed', 'pending', 'in_progress', 'not_set']:
            remark = 'ไม่ผ่านการสอบเทียบ'
        else:
            remark = cal['status']
        
        row_cells[6].text = remark
        row_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # เพิ่มบรรทัดว่าง
    doc.add_paragraph()
    
    # สรุปข้อมูล
    summary_paragraph = doc.add_paragraph()
    summary_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_run = summary_paragraph.add_run(f'สรุป: รายการทั้งหมด {len(filtered_calibrations)} รายการ')
    summary_run.font.size = Pt(12)
    summary_run.font.bold = True
    
    # บันทึกไฟล์
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # ส่งไฟล์กลับ
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="รายงานสอบเทียบ_{datetime.now().strftime("%Y%m%d")}.docx"'
    
    return response

@login_required
def export_to_excel(request):
    """Export รายงานสอบเทียบเป็นไฟล์ Excel"""
    from datetime import date, timedelta
    
    # รับพารามิเตอร์การกรอง (เหมือนกับใน calibration_report_detail)
    department = request.GET.get('department', '')
    instrument_type = request.GET.get('instrument_type', '')
    start_date = request.GET.get('start_date', '')
    end_date = request.GET.get('end_date', '')
    serial_search = request.GET.get('serial_search', '')
    name_search = request.GET.get('name_search', '')
    status_filter = request.GET.get('status_filter', '')
    
    # ดึงข้อมูลหน่วยงาน (Organize) สำหรับที่อยู่
    from organize.models import Organize
    try:
        # หาหน่วยงานหลัก (main unit) หรือหน่วยงานแรก
        organization = Organize.objects.filter(is_main_unit=True).first()
        if not organization:
            organization = Organize.objects.first()
    except:
        organization = None
    
    # ดึงข้อมูลการสอบเทียบทั้งหมด
    pressure_calibrations = CalibrationPressure.objects.select_related('uuc_id', 'std_id', 'calibrator', 'certificate_issuer').all()
    torque_calibrations = CalibrationTorque.objects.select_related('uuc_id', 'std_id', 'calibrator', 'certificate_issuer').all()
    
    # รวมข้อมูลการสอบเทียบทั้งหมด
    all_calibrations = []
    
    # เพิ่มข้อมูล Force calibrations
    for cal in force_calibrations:
        all_calibrations.append({
            'id': cal.cal_force_id,
            'type': 'force',
            'type_name': 'การสอบเทียบแรง',
            'machine_name': cal.uuc_id.name if cal.uuc_id else '-',
            'machine_model': cal.uuc_id.model if cal.uuc_id else '-',
            'serial_number': cal.uuc_id.serial_number if cal.uuc_id else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.update,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.update,  # ǹ��ͺ�º
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.uuc_id.organize.name if cal.uuc_id and cal.uuc_id.organize else '-',
        })
    
    # เพิ่มข้อมูล Pressure calibrations
    for cal in pressure_calibrations:
        all_calibrations.append({
            'id': cal.cal_pressure_id,
            'type': 'pressure',
            'type_name': 'การสอบเทียบ Pressure',
            'machine_name': cal.uuc_id.name if cal.uuc_id else '-',
            'machine_model': cal.uuc_id.model if cal.uuc_id else '-',
            'serial_number': cal.uuc_id.serial_number if cal.uuc_id else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.update,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.update,  # ǹ��ͺ�º
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.uuc_id.organize.name if cal.uuc_id and cal.uuc_id.organize else '-',
        })
    
    # เพิ่มข้อมูล Torque calibrations
    for cal in torque_calibrations:
        all_calibrations.append({
            'id': cal.cal_torque_id,
            'type': 'torque',
            'type_name': 'การสอบเทียบ Torque',
            'machine_name': cal.uuc_id.name if cal.uuc_id else '-',
            'machine_model': cal.uuc_id.model if cal.uuc_id else '-',
            'serial_number': cal.uuc_id.serial_number if cal.uuc_id else '-',
            'std_name': ', '.join([f"{eq.equipment.name} - {eq.equipment.model or ''} - {eq.equipment.serial_number or ''}" for eq in cal.calibration_equipment_used]) if cal.calibration_equipment_used else (cal.std_id.name if cal.std_id else '-'),
            'update_date': cal.update,
            'next_due': cal.next_due,
            'status': cal.status,
            'priority': cal.priority,
            'calibration_date': cal.update,  # ǹ��ͺ�º
            'calibrator': cal.calibrator.get_full_name() if cal.calibrator else '-',
            'certificate_issuer': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            'user_unit': cal.uuc_id.organize.name if cal.uuc_id and cal.uuc_id.organize else '-',
        })
    
    # กรองข้อมูลตามพารามิเตอร์ (เหมือนกับใน calibration_report_detail)
    filtered_calibrations = []
    for cal in all_calibrations:
        include_item = True
        
        # กรองตาม Serial Number
        if serial_search and serial_search.lower() not in cal['serial_number'].lower():
            include_item = False
        
        # กรองตามชื่อเครื่องมือ
        if name_search and name_search.lower() not in cal['machine_name'].lower():
            include_item = False
        
        # กรองตามประเภทเครื่องมือ
        if instrument_type:
            if instrument_type == 'force' and cal['type'] != 'force':
                include_item = False
            elif instrument_type == 'pressure' and cal['type'] != 'pressure':
                include_item = False
            elif instrument_type == 'torque' and cal['type'] != 'torque':
                include_item = False
            elif instrument_type == 'balance' and cal['type'] != 'balance':
                include_item = False
        
        # กรองตามสถานะ
        if status_filter:
            if status_filter == 'cert_issued' and cal['status'] not in ['cert_issued', 'passed', 'active']:
                include_item = False
            elif status_filter == 'failed' and cal['status'] not in ['failed', 'pending', 'in_progress', 'not_set']:
                include_item = False
        
        # กรองตามวันที่
        if start_date and cal['update_date']:
            try:
                start_date_obj = datetime.strptime(start_date, '%Y-%m-%d').date()
                if cal['update_date'] < start_date_obj:
                    include_item = False
            except:
                pass
        
        if end_date and cal['update_date']:
            try:
                end_date_obj = datetime.strptime(end_date, '%Y-%m-%d').date()
                if cal['update_date'] > end_date_obj:
                    include_item = False
            except:
                pass
        
        if include_item:
            filtered_calibrations.append(cal)
    
    # เรียงลำดับตามวันที่สอบเทียบล่าสุด
    filtered_calibrations.sort(key=lambda x: x['calibration_date'] if x['calibration_date'] else date.min, reverse=True)
    
    # สร้าง Excel workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "รายงานสอบเทียบ"
    
    # ตั้งค่าหัวเรื่อง
    title_font = Font(name='TH Sarabun New', size=16, bold=True)
    header_font = Font(name='TH Sarabun New', size=12, bold=True)
    cell_font = Font(name='TH Sarabun New', size=11)
    
    # สีสำหรับ header
    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    
    # หัวเรื่อง
    ws['A1'] = 'รายงานสอบเทียบ'
    ws['A1'].font = title_font
    ws.merge_cells('A1:G1')
    ws['A1'].alignment = Alignment(horizontal='center')
    
    # วันที่รายงาน
    ws['A2'] = f'วันที่: {datetime.now().strftime("%d/%m/%Y")}'
    ws['A2'].font = cell_font
    ws.merge_cells('A2:G2')
    ws['A2'].alignment = Alignment(horizontal='center')
    
    # หัวตาราง
    headers = ['ลำดับ', 'Model', 'SERIAL No.', 'ชื่อเครื่องวัด', 'หน่วยผู้ใช้', 'จำนวน', 'หมายเหตุ']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # เพิ่มข้อมูลในตาราง
    for row, cal in enumerate(filtered_calibrations, 5):
        # ลำดับ
        ws.cell(row=row, column=1, value=row-4).alignment = Alignment(horizontal='center')
        
        # Model
        model_value = cal['machine_model'] if cal['machine_model'] != '-' else '-'
        ws.cell(row=row, column=2, value=model_value).alignment = Alignment(horizontal='center')
        
        # SERIAL No.
        serial_value = cal['serial_number'] if cal['serial_number'] != '-' else '-'
        ws.cell(row=row, column=3, value=serial_value).alignment = Alignment(horizontal='center')
        
        # ชื่อเครื่องวัด
        if cal['type'] == 'pressure':
            instrument_name = 'Pressure Gauge'
        elif cal['type'] == 'torque':
            instrument_name = 'Torque Wrench'
        else:
            instrument_name = cal['type_name']
        
        ws.cell(row=row, column=4, value=instrument_name).alignment = Alignment(horizontal='center')
        
        # หน่วยผู้ใช้
        user_unit = organization.name if organization else 'ยังไม่ได้กรอก'
        ws.cell(row=row, column=5, value=user_unit).alignment = Alignment(horizontal='center')
        
        # จำนวน
        ws.cell(row=row, column=6, value='๑').alignment = Alignment(horizontal='center')
        
        # หมายเหตุ
        if cal['status'] in ['cert_issued', 'passed', 'active']:
            remark = 'ออกใบรับรอง'
        elif cal['status'] in ['failed', 'pending', 'in_progress', 'not_set']:
            remark = 'ไม่ผ่านการสอบเทียบ'
        else:
            remark = cal['status']
        
        ws.cell(row=row, column=7, value=remark).alignment = Alignment(horizontal='center')
    
    # ปรับความกว้างคอลัมน์
    column_widths = [8, 15, 15, 20, 25, 10, 20]
    for col, width in enumerate(column_widths, 1):
        ws.column_dimensions[get_column_letter(col)].width = width
    
    # เพิ่มเส้นขอบ
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    for row in range(4, len(filtered_calibrations) + 5):
        for col in range(1, 8):
            ws.cell(row=row, column=col).border = thin_border
    
    # บันทึกไฟล์
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    # ส่งไฟล์กลับ
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="รายงานสอบเทียบ_{datetime.now().strftime("%Y%m%d")}.xlsx"'
    
    return response

@login_required
def increase_priority(request, cal_type, cal_id):
    """เพิ่มระดับความเร่งด่วนของการสอบเทียบ"""
    # ตรวจสอบสิทธิ์ - เฉพาะเจ้าหน้าที่และ admin
    if not (request.user.is_staff or request.user.is_superuser):
        return JsonResponse({
            'success': False,
            'message': 'คุณไม่มีสิทธิ์ในการเพิ่มระดับความเร่งด่วน'
        })
    
    try:
        # รับข้อมูลจาก request
        import json
        data = json.loads(request.body)
        new_priority = data.get('new_priority')
        
        if not new_priority:
            return JsonResponse({
                'success': False,
                'message': 'ไม่พบข้อมูลระดับความเร่งด่วนใหม่'
            })
        
        # ตรวจสอบระดับความเร่งด่วนที่ถูกต้อง
        valid_priorities = ['normal', 'urgent', 'very_urgent']
        if new_priority not in valid_priorities:
            return JsonResponse({
                'success': False,
                'message': 'ระดับความเร่งด่วนไม่ถูกต้อง'
            })
        
        # ดึงข้อมูลการสอบเทียบตามประเภท
        if cal_type == 'pressure':
            calibration = get_object_or_404(CalibrationPressure, cal_pressure_id=cal_id)
        elif cal_type == 'torque':
            calibration = get_object_or_404(CalibrationTorque, cal_torque_id=cal_id)
        elif cal_type == 'balance':
            calibration = get_object_or_404(BalanceCalibration, pk=cal_id)
        elif cal_type == 'high-frequency':
            calibration = get_object_or_404(HighFrequencyCalibration, pk=cal_id)
        elif cal_type == 'low-frequency':
            calibration = get_object_or_404(LowFrequencyCalibration, pk=cal_id)
        elif cal_type == 'microwave':
            calibration = get_object_or_404(MicrowaveCalibration, pk=cal_id)
        elif cal_type == 'dial_gauge':
            calibration = get_object_or_404(DialGaugeCalibration, pk=cal_id)
        else:
            return JsonResponse({
                'success': False,
                'message': 'ประเภทการสอบเทียบไม่ถูกต้อง'
            })
        
        # ตรวจสอบว่าการเปลี่ยนระดับมีความหมายหรือไม่
        current_priority = calibration.priority
        if current_priority == new_priority:
            return JsonResponse({
                'success': False,
                'message': 'ระดับความเร่งด่วนนี้เป็นระดับปัจจุบันอยู่แล้ว'
            })
        
        # อัปเดตระดับความเร่งด่วน
        calibration.priority = new_priority
        calibration.save()
        
        # สร้างข้อความตอบกลับ
        priority_names = {
            'normal': 'ปกติ',
            'urgent': 'ด่วน',
            'very_urgent': 'ด่วนมาก'
        }
        
        return JsonResponse({
            'success': True,
            'message': f'เปลี่ยนระดับความเร่งด่วนเป็น "{priority_names[new_priority]}" สำเร็จ',
            'old_priority': current_priority,
            'new_priority': new_priority
        })
        
    except Exception as e:
        return JsonResponse({
            'success': False,
            'message': f'เกิดข้อผิดพลาด: {str(e)}'
        })

@login_required
def close_work(request, cal_type, cal_id):
    """ปิดงานการสอบเทียบ"""
    # ตรวจสอบสิทธิ์ - เฉพาะเจ้าหน้าที่และ admin
    if not (request.user.is_staff or request.user.is_superuser):
        return JsonResponse({
            'success': False,
            'message': 'คุณไม่มีสิทธิ์ในการปิดงาน'
        })
    
    try:
        # ดึงข้อมูลการสอบเทียบตามประเภท
        if cal_type == 'pressure':
            calibration = get_object_or_404(CalibrationPressure, cal_pressure_id=cal_id)
        elif cal_type == 'torque':
            calibration = get_object_or_404(CalibrationTorque, cal_torque_id=cal_id)
        elif cal_type == 'balance':
            calibration = get_object_or_404(BalanceCalibration, pk=cal_id)
        elif cal_type == 'high-frequency':
            calibration = get_object_or_404(HighFrequencyCalibration, pk=cal_id)
        elif cal_type == 'low-frequency':
            calibration = get_object_or_404(LowFrequencyCalibration, pk=cal_id)
        elif cal_type == 'microwave':
            calibration = get_object_or_404(MicrowaveCalibration, pk=cal_id)
        elif cal_type == 'dial_gauge':
            calibration = get_object_or_404(DialGaugeCalibration, pk=cal_id)
        else:
            return JsonResponse({
                'success': False,
                'message': 'ประเภทการสอบเทียบไม่ถูกต้อง'
            })
        
        # ตรวจสอบสถานะปัจจุบัน
        current_status = calibration.status
        if current_status == 'closed':
            return JsonResponse({
                'success': False,
                'message': 'งานนี้ถูกปิดแล้ว'
            })
        
        # อัปเดตสถานะเป็นปิดงาน
        calibration.status = 'closed'
        calibration.save()
        
        return JsonResponse({
            'success': True,
            'message': 'ปิดงานสำเร็จ',
            'old_status': current_status,
            'new_status': 'closed'
        })
        
    except Exception as e:
        return JsonResponse({
            'success': False,
            'message': f'เกิดข้อผิดพลาด: {str(e)}'
        })

@login_required
def export_certificate_excel(request, cal_id, cal_type):
    """Export ใบรับรองแบบ Excel (แบบจำลอง)"""
    try:
        # ดึงข้อมูลการสอบเทียบตามประเภท
        if cal_type == 'pressure':
            calibration = get_object_or_404(CalibrationPressure, cal_pressure_id=cal_id)
        elif cal_type == 'torque':
            calibration = get_object_or_404(CalibrationTorque, cal_torque_id=cal_id)
        elif cal_type == 'balance':
            calibration = get_object_or_404(BalanceCalibration, pk=cal_id)
        else:
            messages.error(request, 'ประเภทการสอบเทียบไม่ถูกต้อง')
            return redirect('calibrate-report-detail')
        
        # ตรวจสอบสถานะว่าผ่านการสอบเทียบหรือไม่
        if calibration.status not in ['cert_issued', 'passed', 'active']:
            messages.error(request, 'ไม่สามารถออกใบรับรองได้ เนื่องจากยังไม่ผ่านการสอบเทียบ')
            return redirect('calibrate-report-detail')
        
        # สร้างไฟล์ Excel แบบจำลอง
        wb = Workbook()
        ws = wb.active
        ws.title = "ใบรับรองการสอบเทียบ"
        
        # กำหนดขนาดคอลัมน์
        ws.column_dimensions['A'].width = 20
        ws.column_dimensions['B'].width = 30
        ws.column_dimensions['C'].width = 20
        ws.column_dimensions['D'].width = 30
        
        # หัวข้อหลัก
        ws.merge_cells('A1:D1')
        ws['A1'] = 'ใบรับรองการสอบเทียบเครื่องมือวัด'
        ws['A1'].font = Font(name='TH Sarabun New', size=18, bold=True)
        ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
        
        # ข้อมูลเครื่องมือ
        row = 3
        ws[f'A{row}'] = 'ชื่อเครื่องมือ:'
        ws[f'B{row}'] = calibration.uuc_id.name if calibration.uuc_id else '-'
        ws[f'C{row}'] = 'รุ่น:'
        ws[f'D{row}'] = calibration.uuc_id.model if calibration.uuc_id and calibration.uuc_id.model else '-'
        
        row += 1
        ws[f'A{row}'] = 'Serial Number:'
        ws[f'B{row}'] = calibration.uuc_id.serial_number if calibration.uuc_id and calibration.uuc_id.serial_number else '-'
        ws[f'C{row}'] = 'ผู้ผลิต:'
        ws[f'D{row}'] = str(calibration.uuc_id.manufacture) if calibration.uuc_id and calibration.uuc_id.manufacture else '-'
        
        row += 1
        ws[f'A{row}'] = 'วันที่สอบเทียบ:'
        ws[f'B{row}'] = calibration.update.strftime('%d/%m/%Y') if calibration.update else '-'
        ws[f'C{row}'] = 'วันที่ครบกำหนด:'
        ws[f'D{row}'] = calibration.next_due.strftime('%d/%m/%Y') if calibration.next_due else '-'
        
        row += 1
        ws[f'A{row}'] = 'ผลการสอบเทียบ:'
        ws[f'B{row}'] = 'ผ่านการสอบเทียบ'
        ws[f'B{row}'].font = Font(color='008000', bold=True)  # สีเขียว
        
        row += 2
        ws[f'A{row}'] = 'หมายเหตุ:'
        ws[f'B{row}'] = 'นี่เป็นแบบจำลองใบรับรอง - รอไฟล์ template จริง'
        ws[f'B{row}'].font = Font(italic=True, color='FF0000')  # สีแดง ตัวเอียง
        
        # กำหนดรูปแบบตัวอักษร
        for row in range(3, 8):
            for col in ['A', 'B', 'C', 'D']:
                cell = ws[f'{col}{row}']
                cell.font = Font(name='TH Sarabun New', size=12)
                cell.alignment = Alignment(vertical='center')
        
        # สร้าง response
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        
        # สร้างชื่อไฟล์ที่ปลอดภัย
        machine_name = calibration.uuc_id.name if calibration.uuc_id else "เครื่องมือ"
        safe_filename = f"ใบรับรอง_{machine_name}_{datetime.now().strftime('%Y%m%d')}.xlsx"
        
        response['Content-Disposition'] = f'attachment; filename="{safe_filename}"'
        
        wb.save(response)
        return response
        
    except Exception as e:
        messages.error(request, f'เกิดข้อผิดพลาดในการสร้างใบรับรอง: {str(e)}')
        return redirect('calibrate-report-detail')

from datetime import datetime
import os
from django.conf import settings
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from docx import Document
from docx.oxml.ns import qn

from calibrate.models import LowFrequencyCalibration, CalibrationEquipmentUsed

@login_required
def export_low_frequency_certificate_docx(request, cal_id):
    try:
        print(f"DEBUG: Attempting to export Low Frequency certificate for ID: {cal_id}")
        
        cal = get_object_or_404(
            LowFrequencyCalibration.objects.select_related('machine','std_id','calibrator','certificate_issuer'),
            pk=cal_id
        )
        print(f"DEBUG: Found calibration - Status: {cal.status}, Machine: {cal.machine}")
        
        if cal.status not in ["passed", "cert_issued"]:
            print(f"DEBUG: Status '{cal.status}' not allowed for export")
            messages.error(request, "ไม่สามารถออกใบรับรองได้ เนื่องจากยังไม่ผ่านการสอบเทียบ")
            return redirect("calibrate-report-detail")

        template_path = os.path.join(settings.BASE_DIR, 'cert_templates', 'Low Frequency_template.docx')
        doc = Document(template_path)

        m = cal.machine
        std = cal.std_id

        # องค์กร (ออปชัน)
        try:
            from organize.models import Organize
            org = Organize.objects.filter(is_main_unit=True).first() or Organize.objects.first()
        except Exception:
            org = None

        def fmt(v, dash="-"):
            return v if v not in (None, "") else dash

        def fmt_date(d):
            if not d: return "-"
            if isinstance(d, str):
                try: d = datetime.fromisoformat(d).date()
                except Exception: return d
            if hasattr(d, "date"):
                try: d = d.date()
                except Exception: pass
            try: return d.strftime("%d-%b-%Y")
            except Exception: return str(d)

        # ---------- mappings ส่วนหัว ----------
        replacements = {
            "{{MODEL}}": fmt(getattr(m, "model", None)),
            "{{MANUFACTURER}}": fmt(str(getattr(m, "manufacture", "")) if getattr(m, "manufacture", None) else None),
            "{{DESCRIPTION}}": fmt(getattr(m, "name", None)),
            "{{SERIAL_NUMBER}}": fmt(getattr(m, "serial_number", None)),
            "{{RANGE}}": fmt(getattr(m, "range", None)),
            "{{GRADUATION}}": fmt(getattr(m, "res_uuc", None)),
            "{{OPTION}}": fmt(getattr(m, "option", None), "N/A"),
            "{{CUSTOMER_ASSET_ID}}": fmt(getattr(m, "customer_asset_id", None)),

            "{{RECEIVED_DATE}}": fmt_date(getattr(cal, "received_date", None)),
            "{{DATE_OF_CALIBRATION}}": fmt_date(getattr(cal, "date_calibration", None)),
            "{{DUE_DATE}}": fmt_date(getattr(cal, "next_due", None)),
            "{{ISSUE_DATE}}": fmt_date(getattr(cal, "issue_date", None)),
            "{{CERTIFICATE_NUMBER}}": fmt(getattr(cal, "certificate_number", None)),
            "{{PROCEDURE}}": fmt(getattr(cal, "procedure_number", None)),

            "{{STANDARD_ASSET_NO}}": fmt(getattr(std, "asset_number", None)),
            "{{STANDARD_DESCRIPTION}}": fmt(getattr(std, "name", None)),
            "{{STANDARD_MAKER_MODEL}}": fmt(getattr(std, "description", None)),
            "{{STANDARD_SERIAL}}": fmt(getattr(std, "name", None)),
            "{{STANDARD_CERTIFICATE}}": fmt(getattr(std, "certificate_number", None)),
            "{{STANDARD_DUE_DATE}}": fmt_date(getattr(std, "due_date", None)),

            "{{CALIBRATOR}}": fmt(str(cal.calibrator) if cal.calibrator else None),
            "{{APPROVER}}": fmt(str(cal.certificate_issuer) if cal.certificate_issuer else None),

            "{{CUSTOMER}}": fmt(org.name if org else "Physical Lab, Metrology Division, DC&E (Royal Thai Air Force)"),
            "{{CUSTOMER_ADDRESS}}": fmt(org.address if (org and org.address) else "171 Building. No2025 Sanambin, Donmueang\nBangkok, 10210"),
            "{{LOCATION_OF_CALIBRATION}}": fmt(org.name if org else "Metrology Division, DC&E (Royal Thai Air Force)"),
            "{{LOCATION_ADDRESS}}": fmt(org.address if (org and org.address) else "171 Building. No2025 Sanambin, Donmueang\nBangkok, 10210"),
        }

        # ---------- เติม DC/AC/RES 1..5 ----------
        def add_block(prefix, model_prefix):
            def field_value(base, idx):
                suffix = "" if idx == 1 else f"_{idx}"
                field_name = f"{model_prefix}{base}{suffix}"
                value = getattr(cal, field_name, None)
                print(f"DEBUG: {field_name} = {value}")
                return fmt(value)
            for i in range(1, 6):
                replacements[f"{{{{{prefix}_UUC_RANGE_{i}}}}}"]      = field_value("uuc_range", i)
                replacements[f"{{{{{prefix}_UUC_SETTING_{i}}}}}"]    = field_value("uuc_setting", i)
                replacements[f"{{{{{prefix}_MEASURED_VALUE_{i}}}}}"] = field_value("measured_value", i)
                replacements[f"{{{{{prefix}_UNCERTAINTY_{i}}}}}"]    = field_value("uncertainty", i)
                replacements[f"{{{{{prefix}_TOLERANCE_LIMIT_{i}}}}}"]= field_value("tolerance_limit", i)

        print("DEBUG: Adding DC block")
        add_block("DC",  "dc_")
        print("DEBUG: Adding AC block")
        add_block("AC",  "ac_")
        print("DEBUG: Adding RES block")
        add_block("RES", "res_")
        
        # Debug: Show some replacements
        print(f"DEBUG: Total replacements: {len(replacements)}")
        print(f"DEBUG: DC_UUC_RANGE_2 = {replacements.get('{{DC_UUC_RANGE_2}}', 'NOT FOUND')}")
        print(f"DEBUG: DC_UUC_SETTING_2 = {replacements.get('{{DC_UUC_SETTING_2}}', 'NOT FOUND')}")
        print(f"DEBUG: AC_UUC_RANGE_2 = {replacements.get('{{AC_UUC_RANGE_2}}', 'NOT FOUND')}")
        print(f"DEBUG: RES_UUC_RANGE_2 = {replacements.get('{{RES_UUC_RANGE_2}}', 'NOT FOUND')}")

        # ---------- รายการเครื่องมือที่ใช้สอบเทียบ ----------
        equip_qs = (CalibrationEquipmentUsed.objects
                    .filter(calibration_type="low-frequency", calibration_id=cal.pk)
                    .select_related("equipment"))
        eqs = []
        if std: eqs.append(std)
        for r in equip_qs:
            if r.equipment and all(getattr(r.equipment, "id", None) != getattr(x, "id", None) for x in eqs):
                eqs.append(r.equipment)

        eq_rows = []
        for idx, e in enumerate(eqs, start=1):
            eq_rows.append({
                "no": str(idx),
                "name": fmt(getattr(e, "name", None)),
                "model": fmt(getattr(e, "description", None)),
                "asset": fmt(getattr(e, "asset_number", None)),
                "cert": fmt(getattr(e, "certificate_number", None)),
                "due": fmt_date(getattr(e, "due_date", None)),
            })

        print(f"DEBUG: Equipment rows: {eq_rows}")
        
        replacements["{{EQUIPMENT_LIST}}"] = "\n".join(
            f"{r['no']}. {r['name']} / {r['model']} / Asset:{r['asset']} / Cert:{r['cert']} / Due:{r['due']}"
            for r in eq_rows
        ) or "-"
        
        # เพิ่ม replacements สำหรับ Standard Used section
        if eq_rows:
            # สร้างรายการเครื่องมือทั้งหมด
            equipment_list = []
            for i, eq in enumerate(eq_rows):
                equipment_list.append(f"{i+1}. {eq['name']} / {eq['model']} / Asset: {eq['asset']} / Cert: {eq['cert']} / Due: {eq['due']}")
            
            # ใช้เครื่องมือตัวแรกสำหรับ placeholder หลัก
            first_eq = eq_rows[0]
            replacements["{{STANDARD_ASSET_NO}}"] = first_eq["asset"]
            replacements["{{STANDARD_DESCRIPTION}}"] = first_eq["name"]
            replacements["{{STANDARD_MAKER_MODEL}}"] = first_eq["model"]
            replacements["{{STANDARD_SERIAL}}"] = first_eq["asset"]
            replacements["{{STANDARD_CERTIFICATE}}"] = first_eq["cert"]
            replacements["{{STANDARD_DUE_DATE}}"] = first_eq["due"]
            
            # เพิ่ม placeholder สำหรับเครื่องมือตัวที่สอง (หากมี)
            if len(eq_rows) > 1:
                second_eq = eq_rows[1]
                replacements["{{STANDARD_ASSET_NO_2}}"] = second_eq["asset"]
                replacements["{{STANDARD_DESCRIPTION_2}}"] = second_eq["name"]
                replacements["{{STANDARD_MAKER_MODEL_2}}"] = second_eq["model"]
                replacements["{{STANDARD_SERIAL_2}}"] = second_eq["asset"]
                replacements["{{STANDARD_CERTIFICATE_2}}"] = second_eq["cert"]
                replacements["{{STANDARD_DUE_DATE_2}}"] = second_eq["due"]
            else:
                # หากไม่มีเครื่องมือตัวที่ 2 ให้ใช้เครื่องมือตัวแรกซ้ำ
                replacements["{{STANDARD_ASSET_NO_2}}"] = first_eq["asset"]
                replacements["{{STANDARD_DESCRIPTION_2}}"] = first_eq["name"]
                replacements["{{STANDARD_MAKER_MODEL_2}}"] = first_eq["model"]
                replacements["{{STANDARD_SERIAL_2}}"] = first_eq["asset"]
                replacements["{{STANDARD_CERTIFICATE_2}}"] = first_eq["cert"]
                replacements["{{STANDARD_DUE_DATE_2}}"] = first_eq["due"]
            
            # เพิ่ม placeholder สำหรับเครื่องมือตัวที่สาม (หากมี)
            if len(eq_rows) > 2:
                third_eq = eq_rows[2]
                replacements["{{STANDARD_ASSET_NO_3}}"] = third_eq["asset"]
                replacements["{{STANDARD_DESCRIPTION_3}}"] = third_eq["name"]
                replacements["{{STANDARD_MAKER_MODEL_3}}"] = third_eq["model"]
                replacements["{{STANDARD_SERIAL_3}}"] = third_eq["asset"]
                replacements["{{STANDARD_CERTIFICATE_3}}"] = third_eq["cert"]
                replacements["{{STANDARD_DUE_DATE_3}}"] = third_eq["due"]
            
            # เพิ่ม placeholder สำหรับเครื่องมือตัวที่สี่ (หากมี)
            if len(eq_rows) > 3:
                fourth_eq = eq_rows[3]
                replacements["{{STANDARD_ASSET_NO_4}}"] = fourth_eq["asset"]
                replacements["{{STANDARD_DESCRIPTION_4}}"] = fourth_eq["name"]
                replacements["{{STANDARD_MAKER_MODEL_4}}"] = fourth_eq["model"]
                replacements["{{STANDARD_SERIAL_4}}"] = fourth_eq["asset"]
                replacements["{{STANDARD_CERTIFICATE_4}}"] = fourth_eq["cert"]
                replacements["{{STANDARD_DUE_DATE_4}}"] = fourth_eq["due"]
            
            # เพิ่ม placeholder สำหรับเครื่องมือตัวที่ห้า (หากมี)
            if len(eq_rows) > 4:
                fifth_eq = eq_rows[4]
                replacements["{{STANDARD_ASSET_NO_5}}"] = fifth_eq["asset"]
                replacements["{{STANDARD_DESCRIPTION_5}}"] = fifth_eq["name"]
                replacements["{{STANDARD_MAKER_MODEL_5}}"] = fifth_eq["model"]
                replacements["{{STANDARD_SERIAL_5}}"] = fifth_eq["asset"]
                replacements["{{STANDARD_CERTIFICATE_5}}"] = fifth_eq["cert"]
                replacements["{{STANDARD_DUE_DATE_5}}"] = fifth_eq["due"]
            
            print(f"DEBUG: Equipment count: {len(eq_rows)}")
            print(f"DEBUG: Equipment list: {equipment_list}")
            
            # Debug placeholders
            print(f"DEBUG: STANDARD_ASSET_NO = {replacements.get('{{STANDARD_ASSET_NO}}', 'NOT FOUND')}")
            print(f"DEBUG: STANDARD_DESCRIPTION = {replacements.get('{{STANDARD_DESCRIPTION}}', 'NOT FOUND')}")
            print(f"DEBUG: STANDARD_MAKER_MODEL = {replacements.get('{{STANDARD_MAKER_MODEL}}', 'NOT FOUND')}")
            print(f"DEBUG: STANDARD_SERIAL = {replacements.get('{{STANDARD_SERIAL}}', 'NOT FOUND')}")
            
            # Debug placeholders สำหรับเครื่องมือตัวที่ 2
            print(f"DEBUG: STANDARD_ASSET_NO_2 = {replacements.get('{{STANDARD_ASSET_NO_2}}', 'NOT FOUND')}")
            print(f"DEBUG: STANDARD_DESCRIPTION_2 = {replacements.get('{{STANDARD_DESCRIPTION_2}}', 'NOT FOUND')}")
            print(f"DEBUG: STANDARD_MAKER_MODEL_2 = {replacements.get('{{STANDARD_MAKER_MODEL_2}}', 'NOT FOUND')}")
            print(f"DEBUG: STANDARD_SERIAL_2 = {replacements.get('{{STANDARD_SERIAL_2}}', 'NOT FOUND')}")
        else:
            # ใช้ข้อมูลจาก std_id
            replacements["{{STANDARD_ASSET_NO}}"] = fmt(getattr(std, "asset_number", None))
            replacements["{{STANDARD_DESCRIPTION}}"] = fmt(getattr(std, "name", None))
            replacements["{{STANDARD_MAKER_MODEL}}"] = fmt(getattr(std, "description", None))
            replacements["{{STANDARD_SERIAL}}"] = fmt(getattr(std, "asset_number", None))
            replacements["{{STANDARD_CERTIFICATE}}"] = fmt(getattr(std, "certificate_number", None))
            replacements["{{STANDARD_DUE_DATE}}"] = fmt_date(getattr(std, "due_date", None))

        # ---------- แทนค่าทั้งสองโหมด ----------
        # ฟังก์ชันสำหรับแทนค่าใน paragraph
        def replace_in_paragraph(paragraph, replacements):
            # แทนค่าใน paragraph text ก่อน
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    print(f"DEBUG: Replacing {old_text} with {new_text} in paragraph")
                    paragraph.text = paragraph.text.replace(old_text, new_text)
            
            # แทนค่าใน runs
            for run in paragraph.runs:
                for old_text, new_text in replacements.items():
                    if old_text in run.text:
                        print(f"DEBUG: Replacing {old_text} with {new_text} in run")
                        run.text = run.text.replace(old_text, new_text)
            
            # Debug: ตรวจสอบว่า placeholders ยังเหลืออยู่หรือไม่
            if "{{STANDARD" in paragraph.text:
                print(f"DEBUG: Found unreplaced STANDARD placeholders in paragraph: {paragraph.text[:100]}...")
        
        # ฟังก์ชันสำหรับแทนค่าใน shapes
        def replace_in_shapes(shapes, replacements):
            for shape in shapes:
                if hasattr(shape, 'text_frame') and shape.text_frame:
                    for paragraph in shape.text_frame.paragraphs:
                        replace_in_paragraph(paragraph, replacements)
                if hasattr(shape, 'shapes'):
                    replace_in_shapes(shape.shapes, replacements)
        
        # ฟังก์ชันสำหรับแทนค่าใน document ทั้งหมด
        def replace_in_document(doc, replacements):
            # แทนค่าใน paragraphs
            for paragraph in doc.paragraphs:
                replace_in_paragraph(paragraph, replacements)
            
            # แทนค่าใน tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_in_paragraph(paragraph, replacements)
            
            # แทนค่าใน headers และ footers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        replace_in_paragraph(paragraph, replacements)
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        replace_in_paragraph(paragraph, replacements)
                
                # Text boxes และ shapes
                if hasattr(section, 'shapes'):
                    replace_in_shapes(section.shapes, replacements)
        
        # เรียกใช้ฟังก์ชันแทนค่า
        print("DEBUG: Starting document replacement")
        print(f"DEBUG: Total replacements: {len(replacements)}")
        
        # Debug specific placeholders
        for key, value in replacements.items():
            if "STANDARD" in key:
                print(f"DEBUG: {key} = {value}")
        
        replace_in_document(doc, replacements)
        print("DEBUG: Document replacement completed")

        # ---------- เติมตารางเครื่องมือถ้าใช้ {{EQUIPMENT_TABLE}} ----------
        def find_table_cell_with_placeholder(d, placeholder="{{EQUIPMENT_TABLE}}"):
            for t in d.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if placeholder in p.text:
                                return t, p
            return None, None

        table, p = find_table_cell_with_placeholder(doc, "{{EQUIPMENT_TABLE}}")
        if table and eq_rows:
            for run in p.runs:
                if "{{EQUIPMENT_TABLE}}" in run.text:
                    run.text = run.text.replace("{{EQUIPMENT_TABLE}}", "")
            for r in eq_rows:
                cells = table.add_row().cells
                cells[0].text = r["no"]
                cells[1].text = r["name"]
                cells[2].text = r["model"]
                cells[3].text = r["asset"]
                cells[4].text = r["cert"]
                cells[5].text = r["due"]

        # ---------- ส่งไฟล์ ----------
        print(f"DEBUG: Creating response for Low Frequency certificate")
        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = f'attachment; filename="low_frequency_certificate_{cal.pk}.docx"'
        doc.save(response)
        print(f"DEBUG: Response created successfully")
        return response

    except Exception as e:
        print(f"DEBUG: Error in export_low_frequency_certificate_docx: {str(e)}")
        messages.error(request, f"เกิดข้อผิดพลาดในการสร้างใบรับรอง: {str(e)}")
        return redirect("calibrate-report-detail")

@login_required
def export_balance_certificate_docx(request, cal_id):
    """Export ใบรับรอง Balance แบบ DOCX โดยใช้ template"""
    try:
        print(f"DEBUG: Attempting to export Balance certificate for ID: {cal_id}")
        
        cal = get_object_or_404(
            BalanceCalibration.objects.select_related('machine','std_id','calibrator','certificate_issuer'),
            pk=cal_id
        )
        print(f"DEBUG: Found calibration - Status: {cal.status}, Machine: {cal.machine}")
        
        if cal.status not in ["passed", "cert_issued"]:
            print(f"DEBUG: Status '{cal.status}' not allowed for export")
            messages.error(request, "ไม่สามารถออกใบรับรองได้ เนื่องจากยังไม่ผ่านการสอบเทียบ")
            return redirect("calibrate-report-detail")

        template_path = os.path.join(settings.BASE_DIR, 'cert_templates', 'Balance_template.docx')
        doc = Document(template_path)
        
        m = cal.machine
        std = cal.std_id
        
        # องค์กร (ออปชัน)
        try:
            from organize.models import Organize
            org = Organize.objects.filter(is_main_unit=True).first() or Organize.objects.first()
        except Exception:
            org = None

        def fmt(v, dash="-"):
            return v if v not in (None, "") else dash

        def fmt_date(d):
            if not d: return "-"
            return d.strftime("%d-%b-%Y")

        # ดึงข้อมูลเครื่องมือที่ใช้ในการสอบเทียบ
        eq_rows = []
        from calibrate.models import CalibrationEquipmentUsed
        try:
            used_equipment = CalibrationEquipmentUsed.objects.filter(
                calibration_type='balance',
                calibration_id=cal.pk
            ).select_related('equipment')
            
            for i, eq in enumerate(used_equipment, 1):
                try:
                    eq_rows.append({
                        "no": str(i),
                        "name": fmt(eq.equipment.name),
                        "model": fmt(eq.equipment.model),
                        "asset": fmt(getattr(eq.equipment, 'asset_number', None) or getattr(eq.equipment, 'serial_number', '-')),
                        "cert": fmt(getattr(eq.equipment, 'certificate_number', None) or getattr(eq.equipment, 'certificate', '-')),
                        "due": fmt_date(getattr(eq.equipment, 'due_date', None)),
                    })
                except Exception as e:
                    print(f"DEBUG: Error processing equipment {i}: {e}")
                    # ถ้า error ให้ใส่ข้อมูลพื้นฐานแทน
                    eq_rows.append({
                        "no": str(i),
                        "name": fmt(getattr(eq.equipment, 'name', '-')),
                        "model": "-",
                        "asset": "-",
                        "cert": "-",
                        "due": "-",
                    })
        except Exception as e:
            print(f"DEBUG: Error getting equipment list: {e}")

        # สร้าง replacements dictionary
        replacements = {
            # ข้อมูลเครื่องมือ
            "{{MODEL}}": fmt(m.model),
            "{{MANUFACTURER}}": fmt(str(m.manufacture) if m.manufacture else None),
            "{{DESCRIPTION}}": fmt(m.name),
            "{{SERIAL_NUMBER}}": fmt(m.serial_number),
            "{{RANGE}}": fmt(m.range),
            "{{GRADUATION}}": fmt(m.res_uuc),
            "{{OPTION}}": fmt(m.option),
            "{{CUSTOMER_ASSET_ID}}": fmt(m.customer_asset_id),
            
            # ข้อมูลการสอบเทียบ
            "{{RECEIVED_DATE}}": fmt_date(getattr(cal, "received_date", None)),
            "{{DATE_OF_CALIBRATION}}": fmt_date(cal.date_calibration),
            "{{DUE_DATE}}": fmt_date(cal.next_due),
            "{{ISSUE_DATE}}": fmt_date(getattr(cal, "issue_date", None)),
            "{{CERTIFICATE_NUMBER}}": fmt(getattr(cal, "certificate_number", None)),
            "{{PROCEDURE}}": fmt(getattr(cal, "procedure_number", None)),
            
            # ข้อมูลมาตรฐาน
            "{{STANDARD_ASSET_NO}}": fmt(getattr(std, 'asset_number', None) if std else None),
            "{{STANDARD_DESCRIPTION}}": fmt(std.name if std else None),
            "{{STANDARD_MAKER_MODEL}}": fmt(getattr(std, 'description', None) if std else None),
            "{{STANDARD_SERIAL}}": fmt(std.serial_number if std else None),
            "{{STANDARD_CERTIFICATE}}": fmt(getattr(std, 'certificate_number', None) or std.certificate if std else None),
            "{{STANDARD_DUE_DATE}}": fmt_date(getattr(std, "due_date", None) if std else None),

            # ข้อมูลมาตรฐานชุดที่ 2 (ถ้ามี)
            "{{STANDARD_ASSET_NO_2}}": fmt(getattr(std, 'asset_number', None) if std else None),
            "{{STANDARD_DESCRIPTION_2}}": fmt(std.name if std else None),
            "{{STANDARD_MAKER_MODEL_2}}": fmt(getattr(std, 'description', None) if std else None),
            "{{STANDARD_SERIAL_2}}": fmt(std.serial_number if std else None),
            "{{STANDARD_CERTIFICATE_2}}": fmt(getattr(std, 'certificate_number', None) or std.certificate if std else None),
            "{{STANDARD_DUE_DATE_2}}": fmt_date(getattr(std, "due_date", None) if std else None),

            "{{CALIBRATOR}}": fmt(str(cal.calibrator) if cal.calibrator else None),
            "{{APPROVER}}": fmt(str(cal.certificate_issuer) if cal.certificate_issuer else None),

            "{{CUSTOMER}}": fmt(org.name if org else "Physical Lab, Metrology Division, DC&E (Royal Thai Air Force)"),
            "{{CUSTOMER_ADDRESS}}": fmt(org.address if (org and org.address) else "171 Building. No2025 Sanambin, Donmueang\nBangkok, 10210"),
            "{{LOCATION_OF_CALIBRATION}}": fmt(org.name if org else "Metrology Division, DC&E (Royal Thai Air Force)"),
            "{{LOCATION_ADDRESS}}": fmt(org.address if (org and org.address) else "171 Building. No2025 Sanambin, Donmueang\nBangkok, 10210"),

            # ข้อมูลตารางผลการสอบเทียบ
            "{{CONVENTIONAL_MASS}}": fmt(getattr(cal, 'linear_conventional_mass', None)),
            "{{DISPLAYED_VALUE}}": fmt(getattr(cal, 'linear_displayed_value', None)),
            "{{ERROR}}": fmt(getattr(cal, 'linear_error', None)),
            "{{UNCERTAINTY}}": fmt(getattr(cal, 'linear_uncertainty', None)),
            "{{NOMINAL_VALUE}}": fmt(getattr(cal, 'linear_nominal_value', None)),

            "{{CONVENTIONAL_MASS_2}}": fmt(getattr(cal, 'linear_conventional_mass_2', None)),
            "{{DISPLAYED_VALUE_2}}": fmt(getattr(cal, 'linear_displayed_value_2', None)),
            "{{ERROR_2}}": fmt(getattr(cal, 'linear_error_2', None)),
            "{{UNCERTAINTY_2}}": fmt(getattr(cal, 'linear_uncertainty_2', None)),
            "{{NOMINAL_VALUE_2}}": fmt(getattr(cal, 'linear_nominal_value_2', None)),

            "{{CONVENTIONAL_MASS_3}}": fmt(getattr(cal, 'linear_conventional_mass_3', None)),
            "{{DISPLAYED_VALUE_3}}": fmt(getattr(cal, 'linear_displayed_value_3', None)),
            "{{ERROR_3}}": fmt(getattr(cal, 'linear_error_3', None)),
            "{{UNCERTAINTY_3}}": fmt(getattr(cal, 'linear_uncertainty_3', None)),
            "{{NOMINAL_VALUE_3}}": fmt(getattr(cal, 'linear_nominal_value_3', None)),

            "{{CONVENTIONAL_MASS_4}}": fmt(getattr(cal, 'linear_conventional_mass_4', None)),
            "{{DISPLAYED_VALUE_4}}": fmt(getattr(cal, 'linear_displayed_value_4', None)),
            "{{ERROR_4}}": fmt(getattr(cal, 'linear_error_4', None)),
            "{{UNCERTAINTY_4}}": fmt(getattr(cal, 'linear_uncertainty_4', None)),
            "{{NOMINAL_VALUE_4}}": fmt(getattr(cal, 'linear_nominal_value_4', None)),

            "{{CONVENTIONAL_MASS_5}}": fmt(getattr(cal, 'linear_conventional_mass_5', None)),
            "{{DISPLAYED_VALUE_5}}": fmt(getattr(cal, 'linear_displayed_value_5', None)),
            "{{ERROR_5}}": fmt(getattr(cal, 'linear_error_5', None)),
            "{{UNCERTAINTY_5}}": fmt(getattr(cal, 'linear_uncertainty_5', None)),
            "{{NOMINAL_VALUE_5}}": fmt(getattr(cal, 'linear_nominal_value_5', None)),
        }
        
        # ---------- แทนค่าในเอกสาร ----------
        def replace_text_in_paragraph(paragraph, replacements):
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))
                
                # จัดการ placeholder ที่อาจจะมีปัญหา (เช่น {ERROR}} แทน {{ERROR}})
                problematic_placeholders = {
                    '{ERROR}}': '{{ERROR}}',
                    '{CONVENTIONAL_MASS}}': '{{CONVENTIONAL_MASS}}',
                    '{DISPLAYED_VALUE}}': '{{DISPLAYED_VALUE}}',
                    '{UNCERTAINTY}}': '{{UNCERTAINTY}}',
                    '{NOMINAL_VALUE}}': '{{NOMINAL_VALUE}}',
                    '{ERROR_2}}': '{{ERROR_2}}',
                    '{CONVENTIONAL_MASS_2}}': '{{CONVENTIONAL_MASS_2}}',
                    '{DISPLAYED_VALUE_2}}': '{{DISPLAYED_VALUE_2}}',
                    '{UNCERTAINTY_2}}': '{{UNCERTAINTY_2}}',
                    '{NOMINAL_VALUE_2}}': '{{NOMINAL_VALUE_2}}',
                    '{ERROR_3}}': '{{ERROR_3}}',
                    '{CONVENTIONAL_MASS_3}}': '{{CONVENTIONAL_MASS_3}}',
                    '{DISPLAYED_VALUE_3}}': '{{DISPLAYED_VALUE_3}}',
                    '{UNCERTAINTY_3}}': '{{UNCERTAINTY_3}}',
                    '{NOMINAL_VALUE_3}}': '{{NOMINAL_VALUE_3}}',
                    '{ERROR_4}}': '{{ERROR_4}}',
                    '{CONVENTIONAL_MASS_4}}': '{{CONVENTIONAL_MASS_4}}',
                    '{DISPLAYED_VALUE_4}}': '{{DISPLAYED_VALUE_4}}',
                    '{UNCERTAINTY_4}}': '{{UNCERTAINTY_4}}',
                    '{NOMINAL_VALUE_4}}': '{{NOMINAL_VALUE_4}}',
                    '{ERROR_5}}': '{{ERROR_5}}',
                    '{CONVENTIONAL_MASS_5}}': '{{CONVENTIONAL_MASS_5}}',
                    '{DISPLAYED_VALUE_5}}': '{{DISPLAYED_VALUE_5}}',
                    '{UNCERTAINTY_5}}': '{{UNCERTAINTY_5}}',
                    '{NOMINAL_VALUE_5}}': '{{NOMINAL_VALUE_5}}',
                }
                
                for wrong_placeholder, correct_placeholder in problematic_placeholders.items():
                    if wrong_placeholder in paragraph.text and correct_placeholder in replacements:
                        for run in paragraph.runs:
                            if wrong_placeholder in run.text:
                                run.text = run.text.replace(wrong_placeholder, str(replacements[correct_placeholder]))

        def replace_text_in_document(doc, replacements):
            # แทนค่าใน paragraphs
            for paragraph in doc.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)
            
            # แทนค่าใน tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, replacements)
            
            # แทนค่าใน headers และ footers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        replace_text_in_paragraph(paragraph, replacements)
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        replace_text_in_paragraph(paragraph, replacements)
            
            # จัดการ placeholder ที่มีปัญหาในตาราง
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # แทนที่ placeholder ที่มีปัญหา
                            problematic_placeholders = {
                                '{ERROR}}': '{{ERROR}}',
                                '{CONVENTIONAL_MASS}}': '{{CONVENTIONAL_MASS}}',
                                '{DISPLAYED_VALUE}}': '{{DISPLAYED_VALUE}}',
                                '{UNCERTAINTY}}': '{{UNCERTAINTY}}',
                                '{NOMINAL_VALUE}}': '{{NOMINAL_VALUE}}',
                                '{ERROR_2}}': '{{ERROR_2}}',
                                '{CONVENTIONAL_MASS_2}}': '{{CONVENTIONAL_MASS_2}}',
                                '{DISPLAYED_VALUE_2}}': '{{DISPLAYED_VALUE_2}}',
                                '{UNCERTAINTY_2}}': '{{UNCERTAINTY_2}}',
                                '{NOMINAL_VALUE_2}}': '{{NOMINAL_VALUE_2}}',
                                '{ERROR_3}}': '{{ERROR_3}}',
                                '{CONVENTIONAL_MASS_3}}': '{{CONVENTIONAL_MASS_3}}',
                                '{DISPLAYED_VALUE_3}}': '{{DISPLAYED_VALUE_3}}',
                                '{UNCERTAINTY_3}}': '{{UNCERTAINTY_3}}',
                                '{NOMINAL_VALUE_3}}': '{{NOMINAL_VALUE_3}}',
                                '{ERROR_4}}': '{{ERROR_4}}',
                                '{CONVENTIONAL_MASS_4}}': '{{CONVENTIONAL_MASS_4}}',
                                '{DISPLAYED_VALUE_4}}': '{{DISPLAYED_VALUE_4}}',
                                '{UNCERTAINTY_4}}': '{{UNCERTAINTY_4}}',
                                '{NOMINAL_VALUE_4}}': '{{NOMINAL_VALUE_4}}',
                                '{ERROR_5}}': '{{ERROR_5}}',
                                '{CONVENTIONAL_MASS_5}}': '{{CONVENTIONAL_MASS_5}}',
                                '{DISPLAYED_VALUE_5}}': '{{DISPLAYED_VALUE_5}}',
                                '{UNCERTAINTY_5}}': '{{UNCERTAINTY_5}}',
                                '{NOMINAL_VALUE_5}}': '{{NOMINAL_VALUE_5}}',
                            }
                            
                            # แทนที่ placeholder ที่ถูกต้อง
                            correct_placeholders = {
                                '{{CONVENTIONAL_MASS}}': replacements.get('{{CONVENTIONAL_MASS}}', '-'),
                                '{{DISPLAYED_VALUE}}': replacements.get('{{DISPLAYED_VALUE}}', '-'),
                                '{{ERROR}}': replacements.get('{{ERROR}}', '-'),
                                '{{UNCERTAINTY}}': replacements.get('{{UNCERTAINTY}}', '-'),
                                '{{NOMINAL_VALUE}}': replacements.get('{{NOMINAL_VALUE}}', '-'),
                                '{{CONVENTIONAL_MASS_2}}': replacements.get('{{CONVENTIONAL_MASS_2}}', '-'),
                                '{{DISPLAYED_VALUE_2}}': replacements.get('{{DISPLAYED_VALUE_2}}', '-'),
                                '{{ERROR_2}}': replacements.get('{{ERROR_2}}', '-'),
                                '{{UNCERTAINTY_2}}': replacements.get('{{UNCERTAINTY_2}}', '-'),
                                '{{NOMINAL_VALUE_2}}': replacements.get('{{NOMINAL_VALUE_2}}', '-'),
                                '{{CONVENTIONAL_MASS_3}}': replacements.get('{{CONVENTIONAL_MASS_3}}', '-'),
                                '{{DISPLAYED_VALUE_3}}': replacements.get('{{DISPLAYED_VALUE_3}}', '-'),
                                '{{ERROR_3}}': replacements.get('{{ERROR_3}}', '-'),
                                '{{UNCERTAINTY_3}}': replacements.get('{{UNCERTAINTY_3}}', '-'),
                                '{{NOMINAL_VALUE_3}}': replacements.get('{{NOMINAL_VALUE_3}}', '-'),
                                '{{CONVENTIONAL_MASS_4}}': replacements.get('{{CONVENTIONAL_MASS_4}}', '-'),
                                '{{DISPLAYED_VALUE_4}}': replacements.get('{{DISPLAYED_VALUE_4}}', '-'),
                                '{{ERROR_4}}': replacements.get('{{ERROR_4}}', '-'),
                                '{{UNCERTAINTY_4}}': replacements.get('{{UNCERTAINTY_4}}', '-'),
                                '{{NOMINAL_VALUE_4}}': replacements.get('{{NOMINAL_VALUE_4}}', '-'),
                                '{{CONVENTIONAL_MASS_5}}': replacements.get('{{CONVENTIONAL_MASS_5}}', '-'),
                                '{{DISPLAYED_VALUE_5}}': replacements.get('{{DISPLAYED_VALUE_5}}', '-'),
                                '{{ERROR_5}}': replacements.get('{{ERROR_5}}', '-'),
                                '{{UNCERTAINTY_5}}': replacements.get('{{UNCERTAINTY_5}}', '-'),
                                '{{NOMINAL_VALUE_5}}': replacements.get('{{NOMINAL_VALUE_5}}', '-'),
                            }
                            
                            # Debug: แสดงข้อมูลในตาราง
                            if paragraph.text.strip():
                                print(f"DEBUG: Table cell text: '{paragraph.text}'")
                                for placeholder in ['{{CONVENTIONAL_MASS}}', '{{DISPLAYED_VALUE}}', '{{ERROR}}']:
                                    if placeholder in paragraph.text:
                                        print(f"DEBUG: Found placeholder {placeholder} in table cell")
                            
                            # แทนที่ placeholder ที่มีปัญหา
                            for wrong_placeholder, correct_placeholder in problematic_placeholders.items():
                                if wrong_placeholder in paragraph.text and correct_placeholder in replacements:
                                    for run in paragraph.runs:
                                        if wrong_placeholder in run.text:
                                            run.text = run.text.replace(wrong_placeholder, str(replacements[correct_placeholder]))
                                            print(f"DEBUG: Replaced {wrong_placeholder} with {replacements[correct_placeholder]}")
                            
                            # แทนที่ placeholder ที่ถูกต้อง
                            for placeholder, value in correct_placeholders.items():
                                if placeholder in paragraph.text:
                                    for run in paragraph.runs:
                                        if placeholder in run.text:
                                            run.text = run.text.replace(placeholder, str(value))
                                            print(f"DEBUG: Replaced {placeholder} with {value}")

        # แทนค่าในเอกสาร
        replace_text_in_document(doc, replacements)
        print("DEBUG: Document replacement completed")
        
        # Debug: แสดงข้อมูลตารางที่ส่งไป
        print(f"DEBUG: Conventional Mass 1: {getattr(cal, 'linear_conventional_mass', None)}")
        print(f"DEBUG: Displayed Value 1: {getattr(cal, 'linear_displayed_value', None)}")
        print(f"DEBUG: Error 1: {getattr(cal, 'linear_error', None)}")
        print(f"DEBUG: Conventional Mass 2: {getattr(cal, 'linear_conventional_mass_2', None)}")
        print(f"DEBUG: Displayed Value 2: {getattr(cal, 'linear_displayed_value_2', None)}")
        print(f"DEBUG: Error 2: {getattr(cal, 'linear_error_2', None)}")
        
        # Debug: แสดงข้อมูลมาตรฐาน
        print(f"DEBUG: Standard 1 - Asset: {getattr(std, 'asset_number', None) if std else None}")
        print(f"DEBUG: Standard 1 - Name: {std.name if std else None}")
        print(f"DEBUG: Standard 1 - Serial: {std.serial_number if std else None}")
        print(f"DEBUG: Standard 2 - Asset: {getattr(std, 'asset_number', None) if std else None}")
        print(f"DEBUG: Standard 2 - Name: {std.name if std else None}")
        print(f"DEBUG: Standard 2 - Serial: {std.serial_number if std else None}")
        
        # Debug: แสดง replacements ที่ส่งไป
        print(f"DEBUG: Total replacements: {len(replacements)}")
        print(f"DEBUG: CONVENTIONAL_MASS: {replacements.get('{{CONVENTIONAL_MASS}}', 'NOT_FOUND')}")
        print(f"DEBUG: DISPLAYED_VALUE: {replacements.get('{{DISPLAYED_VALUE}}', 'NOT_FOUND')}")
        print(f"DEBUG: ERROR: {replacements.get('{{ERROR}}', 'NOT_FOUND')}")
        print(f"DEBUG: CONVENTIONAL_MASS_2: {replacements.get('{{CONVENTIONAL_MASS_2}}', 'NOT_FOUND')}")
        print(f"DEBUG: DISPLAYED_VALUE_2: {replacements.get('{{DISPLAYED_VALUE_2}}', 'NOT_FOUND')}")
        print(f"DEBUG: ERROR_2: {replacements.get('{{ERROR_2}}', 'NOT_FOUND')}")
        
        # Debug: แสดงข้อมูลดิบจากฐานข้อมูล
        print(f"DEBUG: Raw data - linear_conventional_mass: {cal.linear_conventional_mass}")
        print(f"DEBUG: Raw data - linear_displayed_value: {cal.linear_displayed_value}")
        print(f"DEBUG: Raw data - linear_error: {cal.linear_error}")
        print(f"DEBUG: Raw data - linear_conventional_mass_2: {cal.linear_conventional_mass_2}")
        print(f"DEBUG: Raw data - linear_displayed_value_2: {cal.linear_displayed_value_2}")
        print(f"DEBUG: Raw data - linear_error_2: {cal.linear_error_2}")
        
        # ---------- เติมตารางเครื่องมือถ้าใช้ {{EQUIPMENT_TABLE}} ----------
        def find_table_cell_with_placeholder(d, placeholder="{{EQUIPMENT_TABLE}}"):
            for t in d.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if placeholder in p.text:
                                return t, p
            return None, None

        table, p = find_table_cell_with_placeholder(doc, "{{EQUIPMENT_TABLE}}")
        if table and eq_rows:
            for run in p.runs:
                if "{{EQUIPMENT_TABLE}}" in run.text:
                    run.text = run.text.replace("{{EQUIPMENT_TABLE}}", "")
            for r in eq_rows:
                cells = table.add_row().cells
                cells[0].text = r["no"]
                cells[1].text = r["name"]
                cells[2].text = r["model"]
                cells[3].text = r["asset"]
                cells[4].text = r["cert"]
                cells[5].text = r["due"]

        # ---------- ส่งไฟล์ ----------
        print(f"DEBUG: Creating response for Balance certificate")
        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = f'attachment; filename="balance_certificate_{cal.pk}.docx"'
        doc.save(response)
        print(f"DEBUG: Response created successfully")
        return response
        
    except Exception as e:
        print(f"DEBUG: Error in export_balance_certificate_docx: {str(e)}")
        messages.error(request, f"เกิดข้อผิดพลาดในการสร้างใบรับรอง: {str(e)}")
        return redirect("calibrate-report-detail")

@login_required
def export_dial_gauge_certificate_docx(request, cal_id):
    """Export ใบรับรอง Dial Gauge แบบ DOCX โดยใช้ template"""
    try:
        print(f"DEBUG: Attempting to export Dial Gauge certificate for ID: {cal_id}")
        
        cal = get_object_or_404(
            DialGaugeCalibration.objects.select_related('machine','std_id','calibrator','certificate_issuer'),
            pk=cal_id
        )
        print(f"DEBUG: Found calibration - Status: {cal.status}, Machine: {cal.machine}")
        
        if cal.status not in ["passed", "cert_issued"]:
            print(f"DEBUG: Status '{cal.status}' not allowed for export")
            messages.error(request, "ไม่สามารถออกใบรับรองได้ เนื่องจากยังไม่ผ่านการสอบเทียบ")
            return redirect("calibrate-report-detail")

        template_path = os.path.join(settings.BASE_DIR, 'cert_templates', 'Dial Gauge_template.docx')
        doc = Document(template_path)
        
        m = cal.machine
        std = cal.std_id

        # องค์กร (ออปชัน)
        try:
            from organize.models import Organize
            org = Organize.objects.filter(is_main_unit=True).first() or Organize.objects.first()
        except Exception:
            org = None

        def fmt(v, dash="-"):
            return v if v not in (None, "") else dash

        def fmt_date(d):
            if not d: return "-"
            return d.strftime("%d-%b-%Y")

        # ดึงข้อมูลเครื่องมือที่ใช้ในการสอบเทียบ
        eq_rows = []
        from calibrate.models import CalibrationEquipmentUsed
        try:
            used_equipment = CalibrationEquipmentUsed.objects.filter(
                calibration_type='dial_gauge',
                calibration_id=cal.pk
            ).select_related('equipment')
            
            for i, eq in enumerate(used_equipment, 1):
                try:
                    eq_rows.append({
                        "no": str(i),
                        "name": fmt(eq.equipment.name),
                        "model": fmt(eq.equipment.model),
                        "asset": fmt(getattr(eq.equipment, 'asset_number', None) or getattr(eq.equipment, 'serial_number', '-')),
                        "cert": fmt(getattr(eq.equipment, 'certificate_number', None) or getattr(eq.equipment, 'certificate', '-')),
                        "due": fmt_date(getattr(eq.equipment, 'due_date', None)),
                    })
                except Exception as e:
                    print(f"DEBUG: Error processing equipment {i}: {e}")
                    # ถ้า error ให้ใส่ข้อมูลพื้นฐานแทน
                    eq_rows.append({
                        "no": str(i),
                        "name": fmt(getattr(eq.equipment, 'name', '-')),
                        "model": "-",
                        "asset": "-",
                        "cert": "-",
                        "due": "-",
                    })
        except Exception as e:
            print(f"DEBUG: Error getting equipment list: {e}")

        # สร้าง replacements dictionary
        replacements = {
            # ข้อมูลเครื่องมือ
            "{{MODEL}}": fmt(m.model),
            "{{MANUFACTURER}}": fmt(str(m.manufacture) if m.manufacture else None),
            "{{DESCRIPTION}}": fmt(m.name),
            "{{SERIAL_NUMBER}}": fmt(m.serial_number),
            "{{RANGE}}": fmt(m.range),
            "{{GRADUATION}}": fmt(m.res_uuc),
            "{{OPTION}}": fmt(m.option),
            "{{CUSTOMER_ASSET_ID}}": fmt(m.customer_asset_id),

            # ข้อมูลการสอบเทียบ
            "{{RECEIVED_DATE}}": fmt_date(getattr(cal, "received_date", None)),
            "{{DATE_OF_CALIBRATION}}": fmt_date(cal.date_calibration),
            "{{DUE_DATE}}": fmt_date(cal.next_due),
            "{{ISSUE_DATE}}": fmt_date(getattr(cal, "issue_date", None)),
            "{{CERTIFICATE_NUMBER}}": fmt(getattr(cal, "certificate_number", None)),
            "{{PROCEDURE}}": fmt(getattr(cal, "procedure_number", None)),

            # ข้อมูลมาตรฐาน
            "{{STANDARD_ASSET_NO}}": fmt(getattr(std, 'asset_number', None) if std else None),
            "{{STANDARD_DESCRIPTION}}": fmt(std.name if std else None),
            "{{STANDARD_MAKER_MODEL}}": fmt(getattr(std, 'description', None) if std else None),
            "{{STANDARD_SERIAL}}": fmt(std.serial_number if std else None),
            "{{STANDARD_CERTIFICATE}}": fmt(getattr(std, 'certificate_number', None) or std.certificate if std else None),
            "{{STANDARD_DUE_DATE}}": fmt_date(getattr(std, "due_date", None) if std else None),

            # ข้อมูลมาตรฐานชุดที่ 2 (ถ้ามี)
            "{{STANDARD_ASSET_NO_2}}": fmt(getattr(std, 'asset_number', None) if std else None),
            "{{STANDARD_DESCRIPTION_2}}": fmt(std.name if std else None),
            "{{STANDARD_MAKER_MODEL_2}}": fmt(getattr(std, 'description', None) if std else None),
            "{{STANDARD_SERIAL_2}}": fmt(std.serial_number if std else None),
            "{{STANDARD_CERTIFICATE_2}}": fmt(getattr(std, 'certificate_number', None) or std.certificate if std else None),
            "{{STANDARD_DUE_DATE_2}}": fmt_date(getattr(std, "due_date", None) if std else None),

            "{{CALIBRATOR}}": fmt(str(cal.calibrator) if cal.calibrator else None),
            "{{APPROVER}}": fmt(str(cal.certificate_issuer) if cal.certificate_issuer else None),

            "{{CUSTOMER}}": fmt(org.name if org else "Physical Lab, Metrology Division, DC&E (Royal Thai Air Force)"),
            "{{CUSTOMER_ADDRESS}}": fmt(org.address if (org and org.address) else "171 Building. No2025 Sanambin, Donmueang\nBangkok, 10210"),
            "{{LOCATION_OF_CALIBRATION}}": fmt(org.name if org else "Metrology Division, DC&E (Royal Thai Air Force)"),
            "{{LOCATION_ADDRESS}}": fmt(org.address if (org and org.address) else "171 Building. No2025 Sanambin, Donmueang\nBangkok, 10210"),

            # ข้อมูลตารางผลการสอบเทียบ Dial Gauge
            "{{UUC_SET}}": fmt(getattr(cal, 'uuc_set', None)),
            "{{ACTUAL}}": fmt(getattr(cal, 'actual', None)),
            "{{ERROR}}": fmt(getattr(cal, 'error', None)),
            "{{UNCERTAINTY}}": fmt(getattr(cal, 'uncertainty', None)),
            "{{TOLERANCE_LIMIT}}": fmt(getattr(cal, 'tolerance_limit', None)),

            "{{UUC_SET_2}}": fmt(getattr(cal, 'set_2', None)),
            "{{ACTUAL_2}}": fmt(getattr(cal, 'actual_2', None)),
            "{{ERROR_2}}": fmt(getattr(cal, 'error_2', None)),
            "{{UNCERTAINTY_2}}": fmt(getattr(cal, 'uncertainty', None)),
            "{{TOLERANCE_LIMIT_2}}": fmt(getattr(cal, 'tolerance_limit_2', None)),

            "{{UUC_SET_3}}": fmt(getattr(cal, 'set_3', None)),
            "{{ACTUAL_3}}": fmt(getattr(cal, 'actual_3', None)),
            "{{ERROR_3}}": fmt(getattr(cal, 'error_3', None)),
            "{{UNCERTAINTY_3}}": fmt(getattr(cal, 'uncertainty', None)),
            "{{TOLERANCE_LIMIT_3}}": fmt(getattr(cal, 'tolerance_limit_3', None)),

            "{{UUC_SET_4}}": fmt(getattr(cal, 'set_4', None)),
            "{{ACTUAL_4}}": fmt(getattr(cal, 'actual_4', None)),
            "{{ERROR_4}}": fmt(getattr(cal, 'error_4', None)),
            "{{UNCERTAINTY_4}}": fmt(getattr(cal, 'uncertainty', None)),
            "{{TOLERANCE_LIMIT_4}}": fmt(getattr(cal, 'tolerance_limit_4', None)),

            "{{UUC_SET_5}}": fmt(getattr(cal, 'set_5', None)),
            "{{ACTUAL_5}}": fmt(getattr(cal, 'actual_5', None)),
            "{{ERROR_5}}": fmt(getattr(cal, 'error_5', None)),
            "{{UNCERTAINTY_5}}": fmt(getattr(cal, 'uncertainty', None)),
            "{{TOLERANCE_LIMIT_5}}": fmt(getattr(cal, 'tolerance_limit_5', None)),
        }
        
        # ---------- แทนค่าในเอกสาร ----------
        def replace_text_in_paragraph(paragraph, replacements):
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))
                
                # จัดการ placeholder ที่อาจจะมีปัญหา (เช่น {ERROR}} แทน {{ERROR}})
                problematic_placeholders = {
                    '{ERROR}}': '{{ERROR}}',
                    '{UUC_SET}}': '{{UUC_SET}}',
                    '{ACTUAL}}': '{{ACTUAL}}',
                    '{UNCERTAINTY}}': '{{UNCERTAINTY}}',
                    '{TOLERANCE_LIMIT}}': '{{TOLERANCE_LIMIT}}',
                    '{ERROR_2}}': '{{ERROR_2}}',
                    '{UUC_SET_2}}': '{{UUC_SET_2}}',
                    '{ACTUAL_2}}': '{{ACTUAL_2}}',
                    '{UNCERTAINTY_2}}': '{{UNCERTAINTY_2}}',
                    '{TOLERANCE_LIMIT_2}}': '{{TOLERANCE_LIMIT_2}}',
                    '{ERROR_3}}': '{{ERROR_3}}',
                    '{UUC_SET_3}}': '{{UUC_SET_3}}',
                    '{ACTUAL_3}}': '{{ACTUAL_3}}',
                    '{UNCERTAINTY_3}}': '{{UNCERTAINTY_3}}',
                    '{TOLERANCE_LIMIT_3}}': '{{TOLERANCE_LIMIT_3}}',
                    '{ERROR_4}}': '{{ERROR_4}}',
                    '{UUC_SET_4}}': '{{UUC_SET_4}}',
                    '{ACTUAL_4}}': '{{ACTUAL_4}}',
                    '{UNCERTAINTY_4}}': '{{UNCERTAINTY_4}}',
                    '{TOLERANCE_LIMIT_4}}': '{{TOLERANCE_LIMIT_4}}',
                    '{ERROR_5}}': '{{ERROR_5}}',
                    '{UUC_SET_5}}': '{{UUC_SET_5}}',
                    '{ACTUAL_5}}': '{{ACTUAL_5}}',
                    '{UNCERTAINTY_5}}': '{{UNCERTAINTY_5}}',
                    '{TOLERANCE_LIMIT_5}}': '{{TOLERANCE_LIMIT_5}}',
                }
                
                # Debug: แสดงข้อมูลใน paragraph
                if paragraph.text.strip():
                    print(f"DEBUG: Paragraph text: '{paragraph.text}'")
                    for placeholder in ['{{UUC_SET}}', '{{ACTUAL}}', '{{ERROR}}']:
                        if placeholder in paragraph.text:
                            print(f"DEBUG: Found placeholder {placeholder} in paragraph")
                
                for wrong_placeholder, correct_placeholder in problematic_placeholders.items():
                    if wrong_placeholder in paragraph.text and correct_placeholder in replacements:
                        for run in paragraph.runs:
                            if wrong_placeholder in run.text:
                                run.text = run.text.replace(wrong_placeholder, str(replacements[correct_placeholder]))
                                print(f"DEBUG: Replaced {wrong_placeholder} with {replacements[correct_placeholder]}")
                
                # แทนที่ placeholder ที่ถูกต้อง
                correct_placeholders = {
                    '{{UUC_SET}}': replacements.get('{{UUC_SET}}', '-'),
                    '{{ACTUAL}}': replacements.get('{{ACTUAL}}', '-'),
                    '{{ERROR}}': replacements.get('{{ERROR}}', '-'),
                    '{{UNCERTAINTY}}': replacements.get('{{UNCERTAINTY}}', '-'),
                    '{{TOLERANCE_LIMIT}}': replacements.get('{{TOLERANCE_LIMIT}}', '-'),
                    '{{UUC_SET_2}}': replacements.get('{{UUC_SET_2}}', '-'),
                    '{{ACTUAL_2}}': replacements.get('{{ACTUAL_2}}', '-'),
                    '{{ERROR_2}}': replacements.get('{{ERROR_2}}', '-'),
                    '{{UNCERTAINTY_2}}': replacements.get('{{UNCERTAINTY_2}}', '-'),
                    '{{TOLERANCE_LIMIT_2}}': replacements.get('{{TOLERANCE_LIMIT_2}}', '-'),
                    '{{UUC_SET_3}}': replacements.get('{{UUC_SET_3}}', '-'),
                    '{{ACTUAL_3}}': replacements.get('{{ACTUAL_3}}', '-'),
                    '{{ERROR_3}}': replacements.get('{{ERROR_3}}', '-'),
                    '{{UNCERTAINTY_3}}': replacements.get('{{UNCERTAINTY_3}}', '-'),
                    '{{TOLERANCE_LIMIT_3}}': replacements.get('{{TOLERANCE_LIMIT_3}}', '-'),
                    '{{UUC_SET_4}}': replacements.get('{{UUC_SET_4}}', '-'),
                    '{{ACTUAL_4}}': replacements.get('{{ACTUAL_4}}', '-'),
                    '{{ERROR_4}}': replacements.get('{{ERROR_4}}', '-'),
                    '{{UNCERTAINTY_4}}': replacements.get('{{UNCERTAINTY_4}}', '-'),
                    '{{TOLERANCE_LIMIT_4}}': replacements.get('{{TOLERANCE_LIMIT_4}}', '-'),
                    '{{UUC_SET_5}}': replacements.get('{{UUC_SET_5}}', '-'),
                    '{{ACTUAL_5}}': replacements.get('{{ACTUAL_5}}', '-'),
                    '{{ERROR_5}}': replacements.get('{{ERROR_5}}', '-'),
                    '{{UNCERTAINTY_5}}': replacements.get('{{UNCERTAINTY_5}}', '-'),
                    '{{TOLERANCE_LIMIT_5}}': replacements.get('{{TOLERANCE_LIMIT_5}}', '-'),
                }
                
                for placeholder, value in correct_placeholders.items():
                    if placeholder in paragraph.text:
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, str(value))
                                print(f"DEBUG: Replaced {placeholder} with {value}")

        def replace_text_in_document(doc, replacements):
            # แทนค่าใน paragraphs
            for paragraph in doc.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)
            
            # แทนค่าใน tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, replacements)
            
            # แทนค่าใน headers และ footers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        replace_text_in_paragraph(paragraph, replacements)
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        replace_text_in_paragraph(paragraph, replacements)
            
            # จัดการ placeholder ที่มีปัญหาในตาราง
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # แทนที่ placeholder ที่มีปัญหา
                            problematic_placeholders = {
                                '{ERROR}}': '{{ERROR}}',
                                '{UUC_SET}}': '{{UUC_SET}}',
                                '{ACTUAL}}': '{{ACTUAL}}',
                                '{UNCERTAINTY}}': '{{UNCERTAINTY}}',
                                '{TOLERANCE_LIMIT}}': '{{TOLERANCE_LIMIT}}',
                                '{ERROR_2}}': '{{ERROR_2}}',
                                '{UUC_SET_2}}': '{{UUC_SET_2}}',
                                '{ACTUAL_2}}': '{{ACTUAL_2}}',
                                '{UNCERTAINTY_2}}': '{{UNCERTAINTY_2}}',
                                '{TOLERANCE_LIMIT_2}}': '{{TOLERANCE_LIMIT_2}}',
                                '{ERROR_3}}': '{{ERROR_3}}',
                                '{UUC_SET_3}}': '{{UUC_SET_3}}',
                                '{ACTUAL_3}}': '{{ACTUAL_3}}',
                                '{UNCERTAINTY_3}}': '{{UNCERTAINTY_3}}',
                                '{TOLERANCE_LIMIT_3}}': '{{TOLERANCE_LIMIT_3}}',
                                '{ERROR_4}}': '{{ERROR_4}}',
                                '{UUC_SET_4}}': '{{UUC_SET_4}}',
                                '{ACTUAL_4}}': '{{ACTUAL_4}}',
                                '{UNCERTAINTY_4}}': '{{UNCERTAINTY_4}}',
                                '{TOLERANCE_LIMIT_4}}': '{{TOLERANCE_LIMIT_4}}',
                                '{ERROR_5}}': '{{ERROR_5}}',
                                '{UUC_SET_5}}': '{{UUC_SET_5}}',
                                '{ACTUAL_5}}': '{{ACTUAL_5}}',
                                '{UNCERTAINTY_5}}': '{{UNCERTAINTY_5}}',
                                '{TOLERANCE_LIMIT_5}}': '{{TOLERANCE_LIMIT_5}}',
                            }
                            
                            # Debug: แสดงข้อมูลในตาราง
                            if paragraph.text.strip():
                                print(f"DEBUG: Table cell text: '{paragraph.text}'")
                                for placeholder in ['{{UUC_SET}}', '{{ACTUAL}}', '{{ERROR}}']:
                                    if placeholder in paragraph.text:
                                        print(f"DEBUG: Found placeholder {placeholder} in table cell")
                            
                            # แทนที่ placeholder ที่ถูกต้อง
                            correct_placeholders = {
                                '{{UUC_SET}}': replacements.get('{{UUC_SET}}', '-'),
                                '{{ACTUAL}}': replacements.get('{{ACTUAL}}', '-'),
                                '{{ERROR}}': replacements.get('{{ERROR}}', '-'),
                                '{{UNCERTAINTY}}': replacements.get('{{UNCERTAINTY}}', '-'),
                                '{{TOLERANCE_LIMIT}}': replacements.get('{{TOLERANCE_LIMIT}}', '-'),
                                '{{UUC_SET_2}}': replacements.get('{{UUC_SET_2}}', '-'),
                                '{{ACTUAL_2}}': replacements.get('{{ACTUAL_2}}', '-'),
                                '{{ERROR_2}}': replacements.get('{{ERROR_2}}', '-'),
                                '{{UNCERTAINTY_2}}': replacements.get('{{UNCERTAINTY_2}}', '-'),
                                '{{TOLERANCE_LIMIT_2}}': replacements.get('{{TOLERANCE_LIMIT_2}}', '-'),
                                '{{UUC_SET_3}}': replacements.get('{{UUC_SET_3}}', '-'),
                                '{{ACTUAL_3}}': replacements.get('{{ACTUAL_3}}', '-'),
                                '{{ERROR_3}}': replacements.get('{{ERROR_3}}', '-'),
                                '{{UNCERTAINTY_3}}': replacements.get('{{UNCERTAINTY_3}}', '-'),
                                '{{TOLERANCE_LIMIT_3}}': replacements.get('{{TOLERANCE_LIMIT_3}}', '-'),
                                '{{UUC_SET_4}}': replacements.get('{{UUC_SET_4}}', '-'),
                                '{{ACTUAL_4}}': replacements.get('{{ACTUAL_4}}', '-'),
                                '{{ERROR_4}}': replacements.get('{{ERROR_4}}', '-'),
                                '{{UNCERTAINTY_4}}': replacements.get('{{UNCERTAINTY_4}}', '-'),
                                '{{TOLERANCE_LIMIT_4}}': replacements.get('{{TOLERANCE_LIMIT_4}}', '-'),
                                '{{UUC_SET_5}}': replacements.get('{{UUC_SET_5}}', '-'),
                                '{{ACTUAL_5}}': replacements.get('{{ACTUAL_5}}', '-'),
                                '{{ERROR_5}}': replacements.get('{{ERROR_5}}', '-'),
                                '{{UNCERTAINTY_5}}': replacements.get('{{UNCERTAINTY_5}}', '-'),
                                '{{TOLERANCE_LIMIT_5}}': replacements.get('{{TOLERANCE_LIMIT_5}}', '-'),
                            }
                            
                            # Debug: แสดงข้อมูลในตาราง
                            if paragraph.text.strip():
                                print(f"DEBUG: Table cell text: '{paragraph.text}'")
                                for placeholder in ['{{UUC_SET}}', '{{ACTUAL}}', '{{ERROR}}']:
                                    if placeholder in paragraph.text:
                                        print(f"DEBUG: Found placeholder {placeholder} in table cell")
                            
                            # แทนที่ placeholder ที่มีปัญหา
                            for wrong_placeholder, correct_placeholder in problematic_placeholders.items():
                                if wrong_placeholder in paragraph.text and correct_placeholder in replacements:
                                    for run in paragraph.runs:
                                        if wrong_placeholder in run.text:
                                            run.text = run.text.replace(wrong_placeholder, str(replacements[correct_placeholder]))
                                            print(f"DEBUG: Replaced {wrong_placeholder} with {replacements[correct_placeholder]}")
                            
                            # แทนที่ placeholder ที่ถูกต้อง
                            for placeholder, value in correct_placeholders.items():
                                if placeholder in paragraph.text:
                                    for run in paragraph.runs:
                                        if placeholder in run.text:
                                            run.text = run.text.replace(placeholder, str(value))
                                            print(f"DEBUG: Replaced {placeholder} with {value}")

        # แทนค่าในเอกสาร
        replace_text_in_document(doc, replacements)
        print("DEBUG: Document replacement completed")
        
        # Debug: แสดงข้อมูลตารางที่ส่งไป
        print(f"DEBUG: UUC_SET: {getattr(cal, 'uuc_set', None)}")
        print(f"DEBUG: ACTUAL: {getattr(cal, 'actual', None)}")
        print(f"DEBUG: ERROR: {getattr(cal, 'error', None)}")
        print(f"DEBUG: UUC_SET_2: {getattr(cal, 'set_2', None)}")
        print(f"DEBUG: ACTUAL_2: {getattr(cal, 'actual_2', None)}")
        print(f"DEBUG: ERROR_2: {getattr(cal, 'error_2', None)}")
        
        # Debug: แสดงข้อมูลมาตรฐาน
        print(f"DEBUG: Standard 1 - Asset: {getattr(std, 'asset_number', None) if std else None}")
        print(f"DEBUG: Standard 1 - Name: {std.name if std else None}")
        print(f"DEBUG: Standard 1 - Serial: {std.serial_number if std else None}")
        print(f"DEBUG: Standard 2 - Asset: {getattr(std, 'asset_number', None) if std else None}")
        print(f"DEBUG: Standard 2 - Name: {std.name if std else None}")
        print(f"DEBUG: Standard 2 - Serial: {std.serial_number if std else None}")
        
        # Debug: แสดง replacements ที่ส่งไป
        print(f"DEBUG: Total replacements: {len(replacements)}")
        print(f"DEBUG: UUC_SET: {replacements.get('{{UUC_SET}}', 'NOT_FOUND')}")
        print(f"DEBUG: ACTUAL: {replacements.get('{{ACTUAL}}', 'NOT_FOUND')}")
        print(f"DEBUG: ERROR: {replacements.get('{{ERROR}}', 'NOT_FOUND')}")
        print(f"DEBUG: UUC_SET_2: {replacements.get('{{UUC_SET_2}}', 'NOT_FOUND')}")
        print(f"DEBUG: ACTUAL_2: {replacements.get('{{ACTUAL_2}}', 'NOT_FOUND')}")
        print(f"DEBUG: ERROR_2: {replacements.get('{{ERROR_2}}', 'NOT_FOUND')}")
        
        # Debug: แสดงข้อมูลดิบจากฐานข้อมูล
        print(f"DEBUG: Raw data - uuc_set: {cal.uuc_set}")
        print(f"DEBUG: Raw data - actual: {cal.actual}")
        print(f"DEBUG: Raw data - error: {cal.error}")
        print(f"DEBUG: Raw data - set_2: {cal.set_2}")
        print(f"DEBUG: Raw data - actual_2: {cal.actual_2}")
        print(f"DEBUG: Raw data - error_2: {cal.error_2}")
        
        # ---------- เติมตารางเครื่องมือถ้าใช้ {{EQUIPMENT_TABLE}} ----------
        def find_table_cell_with_placeholder(d, placeholder="{{EQUIPMENT_TABLE}}"):
            for t in d.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if placeholder in p.text:
                                return t, p
            return None, None

        table, p = find_table_cell_with_placeholder(doc, "{{EQUIPMENT_TABLE}}")
        if table and eq_rows:
            for run in p.runs:
                if "{{EQUIPMENT_TABLE}}" in run.text:
                    run.text = run.text.replace("{{EQUIPMENT_TABLE}}", "")
            for r in eq_rows:
                cells = table.add_row().cells
                cells[0].text = r["no"]
                cells[1].text = r["name"]
                cells[2].text = r["model"]
                cells[3].text = r["asset"]
                cells[4].text = r["cert"]
                cells[5].text = r["due"]

        # ---------- ส่งไฟล์ ----------
        print(f"DEBUG: Creating response for Dial Gauge certificate")
        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = f'attachment; filename="dial_gauge_certificate_{cal.pk}.docx"'
        doc.save(response)
        print(f"DEBUG: Response created successfully")
        return response
        
    except Exception as e:
        print(f"DEBUG: Error in export_dial_gauge_certificate_docx: {str(e)}")
        messages.error(request, f"เกิดข้อผิดพลาดในการสร้างใบรับรอง: {str(e)}")
        return redirect("calibrate-report-detail")

from docxtpl import DocxTemplate
def export_certificate(request, pk):
    # This function is no longer needed as Force Gauge is removed
    return HttpResponse("Force Gauge functionality has been removed", status=410)

    doc = DocxTemplate("Balance_template.docx")

    context = {
        "MODEL": calibration.uuc_id.name if calibration.uuc_id else "",
        "MANUFACTURER": calibration.uuc_id.manufacturer if hasattr(calibration.uuc_id, "manufacturer") else "",
        "DESCRIPTION": calibration.uuc_id.name if calibration.uuc_id else "",
        "SERIAL_NUMBER": getattr(calibration.uuc_id, "serial_number", ""),
        "RANGE": getattr(calibration.uuc_id, "range", ""),
        "GRADUATION": getattr(calibration.uuc_id, "graduation", ""),
        "OPTION": getattr(calibration.uuc_id, "option", ""),
        "CERTIFICATE_NUMBER": getattr(calibration, "certificate_number", ""),
        "CUSTOMER_ASSET_ID": getattr(calibration, "customer_asset_id", ""),
        "PROCEDURE": getattr(calibration, "procedure", ""),
        "RECEIVED_DATE": calibration.update,
        "DATE_OF_CALIBRATION": calibration.update,
        "DUE_DATE": calibration.next_due,
        "ISSUE_DATE": calibration.update,
        "CALIBRATOR": calibration.calibrator.get_full_name() if calibration.calibrator else "",
        "APPROVER": calibration.certificate_issuer.get_full_name() if calibration.certificate_issuer else "",
        "CUSTOMER_ADDRESS": getattr(calibration.uuc_id, "address", ""),
        "LOCATION_OF_CALIBRATION": "กบร.ทร.",
        "LOCATION_ADDRESS": "Somewhere Navy Base",
    }

    doc.render(context)

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="certificate_{calibration.pk}.docx"'
    doc.save(response)
    return response

class HighFrequencyCalibrationUpdateView(LoginRequiredMixin, UpdateView):
    model = HighFrequencyCalibration
    form_class = HighFrequencyCalibrationForm
    template_name = 'calibrate/high_frequency_form.html'
    success_url = reverse_lazy('calibrate-dashboard')
    
    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        if self.request.method == 'POST':
            data = kwargs.get('data')
            if data is None:
                from django.http import QueryDict
                data = QueryDict(mutable=True)
            else:
                data = data.copy()
            if 'status' not in self.request.POST:
                data['status'] = getattr(getattr(self, 'object', None), 'status', None) or 'in_progress'
            kwargs['data'] = data
        return kwargs
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = f'แก้ไขการสอบเทียบ High Frequency - {self.object.machine.name}'
        # เพิ่มข้อมูลเครื่องมือที่ใช้ในการสอบเทียบนี้
        if self.object:
            from calibrate.models import CalibrationEquipmentUsed
            used_equipment = CalibrationEquipmentUsed.objects.filter(
                calibration_type='high-frequency',
                calibration_id=self.object.pk
            ).select_related('equipment')
            context['used_equipment'] = used_equipment
        return context
    
    def form_valid(self, form):
        # บันทึกข้อมูลการสอบเทียบ
        calibration = form.save()
        
        # จัดการข้อมูลเครื่องมือที่ใช้สอบเทียบหลายตัว
        selected_equipment = self.request.POST.get('selected_equipment', '')
        if selected_equipment:
            print(f"Selected equipment: {selected_equipment}")
            # ลบข้อมูลเครื่องมือเก่าที่เกี่ยวข้องกับการสอบเทียบนี้
            from calibrate.models import CalibrationEquipmentUsed
            CalibrationEquipmentUsed.objects.filter(
                calibration_type='high-frequency',
                calibration_id=calibration.pk
            ).delete()
            
            # เพิ่มข้อมูลเครื่องมือใหม่ (ใช้ set เพื่อหลีกเลี่ยงการซ้ำ)
            equipment_ids = set()
            for eid in selected_equipment.split(','):
                eid = eid.strip()
                if eid:
                    equipment_ids.add(eid)
            
            for equipment_id in equipment_ids:
                try:
                    from machine.models import CalibrationEquipment
                    equipment = CalibrationEquipment.objects.get(id=equipment_id)
                    CalibrationEquipmentUsed.objects.get_or_create(
                        calibration_type='high-frequency',
                        calibration_id=calibration.pk,
                        equipment=equipment
                    )
                    print(f"Added equipment: {equipment.name}")
                except Exception as e:
                    print(f"Error adding equipment {equipment_id}: {e}")
        
        # ตรวจสอบสถานะอัตโนมัติ
        self.auto_check_status(calibration)
        
        # ลบ success message ออกตามที่ผู้ใช้ต้องการ
        
        # Redirect ไปยัง success_url โดยตรง
        return redirect(self.success_url)
    
    def auto_check_status(self, calibration):
        """ตรวจสอบสถานะอัตโนมัติสำหรับ High Frequency"""
        # ตรวจสอบว่ามีข้อมูลครบถ้วนหรือไม่
        has_data = any([
            calibration.freq_uuc_setting, calibration.freq_measured_value,
            calibration.freq_uuc_setting_2, calibration.freq_measured_value_2,
            calibration.freq_uuc_setting_3, calibration.freq_measured_value_3,
            calibration.freq_uuc_setting_4, calibration.freq_measured_value_4,
            calibration.freq_uuc_setting_5, calibration.freq_measured_value_5
        ])
        
        if not has_data:
            # ยังไม่มีข้อมูล = กำลังสอบเทียบ
            calibration.status = 'in_progress'
        else:
            # มีข้อมูลแล้ว = ตรวจสอบผลการสอบเทียบ
            if self.check_high_frequency_pass_fail(calibration):
                calibration.status = 'passed'
            else:
                calibration.status = 'failed'
        
        calibration.save()
    
    def check_high_frequency_pass_fail(self, calibration):
        """ตรวจสอบผลการสอบเทียบ High Frequency
        
        เงื่อนไข:
        1. UUC Setting ต้องอยู่ในช่วง Tolerance Limit
        2. Measured Value ต้องอยู่ในช่วง Tolerance Limit
        """
        # ใช้ฟังก์ชัน check_tolerance_limits จาก models.py
        results = calibration.check_tolerance_limits()
        
        if not results:
            return False  # ไม่มีข้อมูล
        
        # ตรวจสอบว่ามี error หรือไม่
        if any('error' in result for result in results):
            return False
        
        # ตรวจสอบว่าทุกแถวผ่านหรือไม่
        all_passed = all(result.get('passed', False) for result in results)
        
        return all_passed

class LowFrequencyCalibrationUpdateView(LoginRequiredMixin, UpdateView):
    model = LowFrequencyCalibration
    form_class = LowFrequencyCalibrationForm
    template_name = 'calibrate/low_frequency_form.html'
    success_url = reverse_lazy('calibrate-dashboard')
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = f'แก้ไขการสอบเทียบ Low Frequency - {self.object.machine.name}'
        return context
    
    def form_valid(self, form):
        # บันทึกข้อมูลการสอบเทียบ
        calibration = form.save()
        
        # Debug: ดูข้อมูลที่ส่งมา
        print(f"=== DEBUG: POST data for equipment (Low Frequency) ===")
        for key, value in self.request.POST.items():
            if 'std_id' in key:
                print(f"{key}: {value}")
        
        # รวบรวมเครื่องมือทั้งหมดที่จะบันทึก (ใช้ set เพื่อหลีกเลี่ยงการซ้ำ)
        equipment_ids = set()
        
        # 1. จัดการข้อมูลเครื่องมือจาก selected_equipment
        selected_equipment = self.request.POST.get('selected_equipment', '')
        print(f"=== DEBUG: selected_equipment = '{selected_equipment}'")
        if selected_equipment:
            for eid in selected_equipment.split(','):
                eid = eid.strip()
                if eid:
                    equipment_ids.add(eid)
        
        # 2. จัดการข้อมูลเครื่องมือจาก form field หลัก (std_id)
        std_id_value = self.request.POST.get('std_id', '')
        if std_id_value:
            equipment_ids.add(std_id_value)
        
        # 3. จัดการข้อมูลเครื่องมือจาก form fields (std_id_*)
        for key, value in self.request.POST.items():
            if key.startswith('std_id_') and value and not key.startswith('std_id_existing_'):
                equipment_ids.add(value)
        
        # 4. จัดการข้อมูลเครื่องมือที่มีอยู่แล้ว (std_id_existing_*)
        for key, value in self.request.POST.items():
            if key.startswith('std_id_existing_') and value:
                equipment_ids.add(value)
        
        print(f"=== DEBUG: Final equipment_ids = {list(equipment_ids)}")
        
        # ลบเครื่องมือเก่าทั้งหมด
        from calibrate.models import CalibrationEquipmentUsed
        CalibrationEquipmentUsed.objects.filter(
            calibration_type='low-frequency',
            calibration_id=calibration.pk
        ).delete()
        
        # บันทึกเครื่องมือใหม่ทั้งหมด (ใช้ get_or_create เพื่อหลีกเลี่ยง duplicate)
        for equipment_id in equipment_ids:
            try:
                from machine.models import CalibrationEquipment
                equipment = CalibrationEquipment.objects.get(id=equipment_id)
                CalibrationEquipmentUsed.objects.get_or_create(
                    calibration_type='low-frequency',
                    calibration_id=calibration.pk,
                    equipment=equipment
                )
                print(f"=== DEBUG: Saved equipment {equipment.name} (ID: {equipment_id})")
            except CalibrationEquipment.DoesNotExist:
                print(f"=== DEBUG: Equipment ID {equipment_id} not found")
                continue
        
        # ตรวจสอบสถานะอัตโนมัติ
        self.auto_check_status(calibration)
        
        # เพิ่ม success message
        from django.contrib import messages
        # ลบ success message ออกตามที่ผู้ใช้ต้องการ
        # ให้ Django จัดการ redirect ตาม success_url
        return super().form_valid(form)
    
    def auto_check_status(self, calibration):
        """ตรวจสอบสถานะอัตโนมัติสำหรับ Low Frequency"""
        # ตรวจสอบว่ามีข้อมูลครบถ้วนหรือไม่
        has_data = any([
            calibration.dc_uuc_setting, calibration.dc_measured_value,
            calibration.ac_uuc_setting, calibration.ac_measured_value,
            calibration.res_uuc_setting, calibration.res_measured_value
        ])
        
        if not has_data:
            # ยังไม่มีข้อมูล = กำลังสอบเทียบ
            calibration.status = 'in_progress'
        else:
            # มีข้อมูลแล้ว = ตรวจสอบผลการสอบเทียบ
            if self.check_low_frequency_pass_fail(calibration):
                calibration.status = 'passed'
            else:
                calibration.status = 'failed'
        
        calibration.save()
    
    def check_low_frequency_pass_fail(self, calibration):
        """ตรวจสอบผลการสอบเทียบ Low Frequency"""
        # ตรวจสอบ DC VOLTAGE
        dc_rows = [
            (calibration.dc_uuc_setting, calibration.dc_measured_value, calibration.dc_uncertainty),
            (calibration.dc_uuc_setting_2, calibration.dc_measured_value_2, calibration.dc_uncertainty_2),
            (calibration.dc_uuc_setting_3, calibration.dc_measured_value_3, calibration.dc_uncertainty_3),
            (calibration.dc_uuc_setting_4, calibration.dc_measured_value_4, calibration.dc_uncertainty_4),
            (calibration.dc_uuc_setting_5, calibration.dc_measured_value_5, calibration.dc_uncertainty_5)
        ]
        
        for setting, measured, uncertainty in dc_rows:
            if setting and measured and uncertainty:
                try:
                    setting_val = float(setting)
                    measured_val = float(measured)
                    uncertainty_val = float(uncertainty)
                    
                    if abs(measured_val - setting_val) > uncertainty_val:
                        return False
                except (ValueError, TypeError):
                    return False
        
        # ตรวจสอบ AC VOLTAGE
        ac_rows = [
            (calibration.ac_uuc_setting, calibration.ac_measured_value, calibration.ac_uncertainty),
            (calibration.ac_uuc_setting_2, calibration.ac_measured_value_2, calibration.ac_uncertainty_2),
            (calibration.ac_uuc_setting_3, calibration.ac_measured_value_3, calibration.ac_uncertainty_3),
            (calibration.ac_uuc_setting_4, calibration.ac_measured_value_4, calibration.ac_uncertainty_4),
            (calibration.ac_uuc_setting_5, calibration.ac_measured_value_5, calibration.ac_uncertainty_5)
        ]
        
        for setting, measured, uncertainty in ac_rows:
            if setting and measured and uncertainty:
                try:
                    setting_val = float(setting)
                    measured_val = float(measured)
                    uncertainty_val = float(uncertainty)
                    
                    if abs(measured_val - setting_val) > uncertainty_val:
                        return False
                except (ValueError, TypeError):
                    return False
        
        # ตรวจสอบ RESISTANCE
        res_rows = [
            (calibration.res_uuc_setting, calibration.res_measured_value, calibration.res_uncertainty),
            (calibration.res_uuc_setting_2, calibration.res_measured_value_2, calibration.res_uncertainty_2),
            (calibration.res_uuc_setting_3, calibration.res_measured_value_3, calibration.res_uncertainty_3),
            (calibration.res_uuc_setting_4, calibration.res_measured_value_4, calibration.res_uncertainty_4),
            (calibration.res_uuc_setting_5, calibration.res_measured_value_5, calibration.res_uncertainty_5)
        ]
        
        for setting, measured, uncertainty in res_rows:
            if setting and measured and uncertainty:
                try:
                    setting_val = float(setting)
                    measured_val = float(measured)
                    uncertainty_val = float(uncertainty)
                    
                    if abs(measured_val - setting_val) > uncertainty_val:
                        return False
                except (ValueError, TypeError):
                    return False
        
        return True

class MicrowaveCalibrationUpdateView(LoginRequiredMixin, UpdateView):
    model = MicrowaveCalibration
    form_class = MicrowaveCalibrationForm
    template_name = 'calibrate/microwave_form.html'
    success_url = reverse_lazy('calibrate-dashboard')
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = f'แก้ไขการสอบเทียบ Microwave - {self.object.machine.name}'
        # เพิ่มข้อมูลเครื่องมือที่ใช้ในการสอบเทียบนี้
        if self.object:
            from calibrate.models import CalibrationEquipmentUsed
            used_equipment = CalibrationEquipmentUsed.objects.filter(
                calibration_type='microwave',
                calibration_id=self.object.pk
            ).select_related('equipment')
            context['used_equipment'] = used_equipment
        return context
    
    def form_valid(self, form):
        print("=== DEBUG: ก่อนบันทึก Microwave ===")
        print(f"POST data: {dict(self.request.POST)}")
        
        # ตรวจสอบ form errors
        if form.errors:
            print(f"❌ Form errors: {form.errors}")
            return self.form_invalid(form)
        
        try:
            # บันทึกข้อมูลการสอบเทียบ
            calibration = form.save()
            print(f"✅ บันทึก calibration ID: {calibration.pk}")
        except Exception as e:
            print(f"❌ Error saving calibration: {e}")
            return self.form_invalid(form)
        
        # จัดการข้อมูลเครื่องมือที่ใช้สอบเทียบหลายตัว
        selected_equipment = self.request.POST.get('selected_equipment', '')
        if selected_equipment:
            print(f"Selected equipment: {selected_equipment}")
            # ลบข้อมูลเครื่องมือเก่าที่เกี่ยวข้องกับการสอบเทียบนี้
            from calibrate.models import CalibrationEquipmentUsed
            CalibrationEquipmentUsed.objects.filter(
                calibration_type='microwave',
                calibration_id=calibration.pk
            ).delete()
            
            # เพิ่มข้อมูลเครื่องมือใหม่ (ใช้ set เพื่อหลีกเลี่ยงการซ้ำ)
            equipment_ids = set()
            for eid in selected_equipment.split(','):
                eid = eid.strip()
                if eid:
                    equipment_ids.add(eid)
            
            for equipment_id in equipment_ids:
                try:
                    from machine.models import CalibrationEquipment
                    equipment = CalibrationEquipment.objects.get(id=equipment_id)
                    CalibrationEquipmentUsed.objects.get_or_create(
                        calibration_type='microwave',
                        calibration_id=calibration.pk,
                        equipment=equipment
                    )
                    print(f"Added equipment: {equipment.name}")
                except Exception as e:
                    print(f"Error adding equipment {equipment_id}: {e}")
        
        # ตรวจสอบสถานะอัตโนมัติ
        self.auto_check_status(calibration)
        
        # เพิ่ม success message
        from django.contrib import messages
        # ลบ success message ออกตามที่ผู้ใช้ต้องการ
        
        # Redirect ไปยัง success_url โดยตรง
        return redirect(self.success_url)
    
    def auto_check_status(self, calibration):
        """ตรวจสอบสถานะอัตโนมัติสำหรับ Microwave"""
        # ตรวจสอบว่ามีข้อมูลครบถ้วนหรือไม่
        has_data = any([
            calibration.dc_uuc_setting, calibration.dc_measured_value,
            calibration.dc_uuc_setting_2, calibration.dc_measured_value_2,
            calibration.dc_uuc_setting_3, calibration.dc_measured_value_3,
            calibration.dc_uuc_setting_4, calibration.dc_measured_value_4,
            calibration.dc_uuc_setting_5, calibration.dc_measured_value_5
        ])
        
        if not has_data:
            # ยังไม่มีข้อมูล = กำลังสอบเทียบ
            calibration.status = 'in_progress'
        else:
            # มีข้อมูลแล้ว = ตรวจสอบผลการสอบเทียบ
            if self.check_microwave_pass_fail(calibration):
                calibration.status = 'passed'
            else:
                calibration.status = 'failed'
        
        calibration.save()
    
    def check_microwave_pass_fail(self, calibration):
        """ตรวจสอบผลการสอบเทียบ Microwave"""
        # ตรวจสอบทุกแถวที่มีข้อมูล
        rows_to_check = [
            (calibration.dc_uuc_setting, calibration.dc_measured_value, calibration.dc_uncertainty),
            (calibration.dc_uuc_setting_2, calibration.dc_measured_value_2, calibration.dc_uncertainty_2),
            (calibration.dc_uuc_setting_3, calibration.dc_measured_value_3, calibration.dc_uncertainty_3),
            (calibration.dc_uuc_setting_4, calibration.dc_measured_value_4, calibration.dc_uncertainty_4),
            (calibration.dc_uuc_setting_5, calibration.dc_measured_value_5, calibration.dc_uncertainty_5)
        ]
        
        for setting, measured, uncertainty in rows_to_check:
            if setting and measured and uncertainty:
                try:
                    setting_val = float(setting)
                    measured_val = float(measured)
                    uncertainty_val = float(uncertainty)
                    
                    if abs(measured_val - setting_val) > uncertainty_val:
                        return False
                except (ValueError, TypeError):
                    return False
        
        return True

class DialGaugeCalibrationUpdateView(LoginRequiredMixin, UpdateView):
    model = DialGaugeCalibration
    form_class = DialGaugeCalibrationForm
    template_name = 'calibrate/dial_gauge_form.html'
    success_url = reverse_lazy('calibrate-dashboard')
    
    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        if self.request.method == 'POST':
            data = kwargs.get('data')
            if data is None:
                from django.http import QueryDict
                data = QueryDict(mutable=True)
            else:
                data = data.copy()
            if 'status' not in self.request.POST:
                data['status'] = getattr(getattr(self, 'object', None), 'status', None) or 'in_progress'
            kwargs['data'] = data
        return kwargs
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = f'แก้ไขการสอบเทียบ Dial Gauge - {self.object.machine.name}'
        # เพิ่มข้อมูลเครื่องมือที่ใช้ในการสอบเทียบนี้
        if self.object:
            from calibrate.models import CalibrationEquipmentUsed
            used_equipment = CalibrationEquipmentUsed.objects.filter(
                calibration_type='dial_gauge',
                calibration_id=self.object.pk
            ).select_related('equipment')
            context['used_equipment'] = used_equipment
        return context
    
    def form_valid(self, form):
        print("=== DEBUG: ก่อนบันทึก Dial Gauge ===")
        print(f"POST data: {dict(self.request.POST)}")
        
        # ตรวจสอบ form errors
        if form.errors:
            print(f"❌ Form errors: {form.errors}")
            return self.form_invalid(form)
        
        try:
            # บันทึกข้อมูลการสอบเทียบ
            calibration = form.save()
            print(f"✅ บันทึก calibration ID: {calibration.pk}")
        except Exception as e:
            print(f"❌ Error saving calibration: {e}")
            return self.form_invalid(form)
        
        # จัดการข้อมูลเครื่องมือที่ใช้สอบเทียบหลายตัว
        selected_equipment = self.request.POST.get('selected_equipment', '')
        if selected_equipment:
            print(f"Selected equipment: {selected_equipment}")
            # ลบข้อมูลเครื่องมือเก่าที่เกี่ยวข้องกับการสอบเทียบนี้
            from calibrate.models import CalibrationEquipmentUsed
            CalibrationEquipmentUsed.objects.filter(
                calibration_type='dial_gauge',
                calibration_id=calibration.pk
            ).delete()
            
            # เพิ่มข้อมูลเครื่องมือใหม่ (ใช้ set เพื่อหลีกเลี่ยงการซ้ำ)
            equipment_ids = set()
            for eid in selected_equipment.split(','):
                eid = eid.strip()
                if eid:
                    equipment_ids.add(eid)
            
            for equipment_id in equipment_ids:
                try:
                    from machine.models import CalibrationEquipment
                    equipment = CalibrationEquipment.objects.get(id=equipment_id)
                    CalibrationEquipmentUsed.objects.get_or_create(
                        calibration_type='dial_gauge',
                        calibration_id=calibration.pk,
                        equipment=equipment
                    )
                    print(f"Added equipment: {equipment.name}")
                except Exception as e:
                    print(f"Error adding equipment {equipment_id}: {e}")
        
        # ตรวจสอบสถานะอัตโนมัติ
        self.auto_check_status(calibration)
        
        # เพิ่ม success message
        from django.contrib import messages
        # ลบ success message ออกตามที่ผู้ใช้ต้องการ
        
        # Redirect ไปยัง success_url โดยตรง
        return redirect(self.success_url)
    
    def auto_check_status(self, calibration):
        """ตรวจสอบสถานะอัตโนมัติสำหรับ Dial Gauge"""
        # ตรวจสอบว่ามีข้อมูลครบถ้วนหรือไม่
        has_data = any([
            calibration.uuc_set, calibration.actual,
            calibration.set_2, calibration.actual_2,
            calibration.set_3, calibration.actual_3,
            calibration.set_4, calibration.actual_4,
            calibration.set_5, calibration.actual_5
        ])
        
        if not has_data:
            # ยังไม่มีข้อมูล = กำลังสอบเทียบ
            calibration.status = 'in_progress'
        else:
            # มีข้อมูลแล้ว = ตรวจสอบผลการสอบเทียบ
            if self.check_dial_gauge_pass_fail(calibration):
                calibration.status = 'passed'
            else:
                calibration.status = 'failed'
        
        calibration.save()
    
    def check_dial_gauge_pass_fail(self, calibration):
        """ตรวจสอบผลการสอบเทียบ Dial Gauge"""
        # ตรวจสอบทุกแถวที่มีข้อมูล
        rows_to_check = [
            (calibration.uuc_set, calibration.actual, calibration.error, calibration.tolerance_limit),
            (calibration.set_2, calibration.actual_2, calibration.error_2, calibration.tolerance_limit_2),
            (calibration.set_3, calibration.actual_3, calibration.error_3, calibration.tolerance_limit_3),
            (calibration.set_4, calibration.actual_4, calibration.error_4, calibration.tolerance_limit_4),
            (calibration.set_5, calibration.actual_5, calibration.error_5, calibration.tolerance_limit_5)
        ]
        
        for set_val, actual, error, tolerance in rows_to_check:
            if set_val and actual and error is not None and tolerance:
                try:
                    set_val_num = float(set_val)
                    actual_num = float(actual)
                    error_num = float(error)
                    
                    # ตรวจสอบว่าค่า error อยู่ในช่วงที่ยอมรับได้หรือไม่
                    if abs(error_num) > abs(set_val_num * 0.01):  # เกิน 1% = ไม่ผ่าน
                        return False
                except (ValueError, TypeError):
                    return False
        
        return True

class HighFrequencyCalibrationDeleteView(LoginRequiredMixin, DeleteView):
    model = HighFrequencyCalibration
    template_name = 'calibrate/high_frequency_confirm_delete.html'
    success_url = reverse_lazy('calibrate-dashboard')
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = f'ลบการสอบเทียบ High Frequency - {self.object.machine.name}'
        return context

class LowFrequencyCalibrationDeleteView(LoginRequiredMixin, DeleteView):
    model = LowFrequencyCalibration
    template_name = 'calibrate/low_frequency_confirm_delete.html'
    success_url = reverse_lazy('calibrate-dashboard')
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = f'ลบการสอบเทียบ Low Frequency - {self.object.machine.name}'
        return context

class MicrowaveCalibrationDeleteView(LoginRequiredMixin, DeleteView):
    model = MicrowaveCalibration
    template_name = 'calibrate/microwave_confirm_delete.html'
    success_url = reverse_lazy('calibrate-dashboard')
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = f'ลบการสอบเทียบ Microwave - {self.object.machine.name}'
        return context

class DialGaugeCalibrationDeleteView(LoginRequiredMixin, DeleteView):
    model = DialGaugeCalibration
    template_name = 'calibrate/dial_gauge_confirm_delete.html'
    success_url = reverse_lazy('calibrate-dashboard')
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = f'ลบการสอบเทียบ Dial Gauge - {self.object.machine.name}'
        return context

def replace_text_in_paragraph(paragraph, replacements):
    """แทนที่ placeholder ใน paragraph"""
    for placeholder, value in replacements.items():
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(value))
    
    # จัดการ placeholder ที่อาจจะมีปัญหา (เช่น {ERROR}} แทน {{ERROR}})
    problematic_placeholders = {
        '{ERROR}}': '{{ERROR}}',
        '{FREQ_UUC_RANGE}}': '{{FREQ_UUC_RANGE}}',
        '{FREQ_UUC_SETTING}}': '{{FREQ_UUC_SETTING}}',
        '{FREQ_MEASURED_VALUE}}': '{{FREQ_MEASURED_VALUE}}',
        '{FREQ_UNCERTAINTY}}': '{{FREQ_UNCERTAINTY}}',
        '{FREQ_TOLERANCE_LIMIT}}': '{{FREQ_TOLERANCE_LIMIT}}',
        '{VOLT_UUC_RANGE}}': '{{VOLT_UUC_RANGE}}',
        '{VOLT_UUC_SETTING}}': '{{VOLT_UUC_SETTING}}',
        '{VOLT_MEASURED_VALUE}}': '{{VOLT_MEASURED_VALUE}}',
        '{VOLT_UNCERTAINTY}}': '{{VOLT_UNCERTAINTY}}',
        '{VOLT_TOLERANCE_LIMIT}}': '{{VOLT_TOLERANCE_LIMIT}}',
        '{DATE_OF_CALIBRATION}}': '{{DATE_OF_CALIBRATION}}',
        '{CUSTOMER_ADDRESS}}': '{{CUSTOMER_ADDRESS}}',
        '{RANGE}}': '{{RANGE}}',
        '{APPROVER}}': '{{APPROVER}}',
        '{CALIBRATOR}}': '{{CALIBRATOR}}',
        '{UUC_SET}}': '{{UUC_SET}}',
        '{ACTUAL}}': '{{ACTUAL}}',
        '{ERROR}}': '{{ERROR}}',
        '{TOLERANCE_LIMIT}}': '{{TOLERANCE_LIMIT}}',
        '{UUC_SET_2}}': '{{UUC_SET_2}}',
        '{ACTUAL_2}}': '{{ACTUAL_2}}',
        '{ERROR_2}}': '{{ERROR_2}}',
        '{TOLERANCE_LIMIT_2}}': '{{TOLERANCE_LIMIT_2}}',
        '{UUC_SET_3}}': '{{UUC_SET_3}}',
        '{ACTUAL_3}}': '{{ACTUAL_3}}',
        '{ERROR_3}}': '{{ERROR_3}}',
        '{TOLERANCE_LIMIT_3}}': '{{TOLERANCE_LIMIT_3}}',
        '{UUC_SET_4}}': '{{UUC_SET_4}}',
        '{ACTUAL_4}}': '{{ACTUAL_4}}',
        '{ERROR_4}}': '{{ERROR_4}}',
        '{TOLERANCE_LIMIT_4}}': '{{TOLERANCE_LIMIT_4}}',
        '{UUC_SET_5}}': '{{UUC_SET_5}}',
        '{ACTUAL_5}}': '{{ACTUAL_5}}',
        '{ERROR_5}}': '{{ERROR_5}}',
        '{TOLERANCE_LIMIT_5}}': '{{TOLERANCE_LIMIT_5}}',
        '{UUC_SET_6}}': '{{UUC_SET_6}}',
        '{ACTUAL_6}}': '{{ACTUAL_6}}',
        '{ERROR_6}}': '{{ERROR_6}}',
        '{TOLERANCE_LIMIT_6}}': '{{TOLERANCE_LIMIT_6}}',
        '{CW_SET}}': '{{CW_SET}}',
        '{CW_0}}': '{{CW_0}}',
        '{CW_90}}': '{{CW_90}}',
        '{CW_180}}': '{{CW_180}}',
        '{CW_270}}': '{{CW_270}}',
        '{CW_AVG}}': '{{CW_AVG}}',
        '{CW_ACTUAL}}': '{{CW_ACTUAL}}',
        '{CW_ERROR}}': '{{CW_ERROR}}',
        '{CW_UNCERTAINTY}}': '{{CW_UNCERTAINTY}}',
        '{CW_TOLERANCE_LIMIT}}': '{{CW_TOLERANCE_LIMIT}}',
        '{CCW_SET}}': '{{CCW_SET}}',
        '{CCW_0}}': '{{CCW_0}}',
        '{CCW_90}}': '{{CCW_90}}',
        '{CCW_180}}': '{{CCW_180}}',
        '{CCW_270}}': '{{CCW_270}}',
        '{CCW_AVG}}': '{{CCW_AVG}}',
        '{CCW_ACTUAL}}': '{{CCW_ACTUAL}}',
        '{CCW_ERROR}}': '{{CCW_ERROR}}',
        '{CCW_UNCERTAINTY}}': '{{CCW_UNCERTAINTY}}',
        '{CCW_TOLERANCE_LIMIT}}': '{{CCW_TOLERANCE_LIMIT}}',
        '{CCW_ACTUAL_2}}': '{{CCW_ACTUAL_2}}',
        '{CCW_ERROR_2}}': '{{CCW_ERROR_2}}',
        '{CCW_TOLERANCE_LIMIT_2}}': '{{CCW_TOLERANCE_LIMIT_2}}',
        '{CCW_ACTUAL_3}}': '{{CCW_ACTUAL_3}}',
        '{CCW_ERROR_3}}': '{{CCW_ERROR_3}}',
        '{CCW_TOLERANCE_LIMIT_3}}': '{{CCW_TOLERANCE_LIMIT_3}}',
        '{ERROR_2}}': '{{ERROR_2}}',
        '{FREQ_UUC_RANGE_2}}': '{{FREQ_UUC_RANGE_2}}',
        '{FREQ_UUC_SETTING_2}}': '{{FREQ_UUC_SETTING_2}}',
        '{FREQ_MEASURED_VALUE_2}}': '{{FREQ_MEASURED_VALUE_2}}',
        '{FREQ_UNCERTAINTY_2}}': '{{FREQ_UNCERTAINTY_2}}',
        '{FREQ_TOLERANCE_LIMIT_2}}': '{{FREQ_TOLERANCE_LIMIT_2}}',
        '{VOLT_UUC_RANGE_2}}': '{{VOLT_UUC_RANGE_2}}',
        '{VOLT_UUC_SETTING_2}}': '{{VOLT_UUC_SETTING_2}}',
        '{VOLT_MEASURED_VALUE_2}}': '{{VOLT_MEASURED_VALUE_2}}',
        '{VOLT_UNCERTAINTY_2}}': '{{VOLT_UNCERTAINTY_2}}',
        '{VOLT_TOLERANCE_LIMIT_2}}': '{{VOLT_TOLERANCE_LIMIT_2}}',
        '{ERROR_3}}': '{{ERROR_3}}',
        '{FREQ_UUC_RANGE_3}}': '{{FREQ_UUC_RANGE_3}}',
        '{FREQ_UUC_SETTING_3}}': '{{FREQ_UUC_SETTING_3}}',
        '{FREQ_MEASURED_VALUE_3}}': '{{FREQ_MEASURED_VALUE_3}}',
        '{FREQ_UNCERTAINTY_3}}': '{{FREQ_UNCERTAINTY_3}}',
        '{FREQ_TOLERANCE_LIMIT_3}}': '{{FREQ_TOLERANCE_LIMIT_3}}',
        '{VOLT_UUC_RANGE_3}}': '{{VOLT_UUC_RANGE_3}}',
        '{VOLT_UUC_SETTING_3}}': '{{VOLT_UUC_SETTING_3}}',
        '{VOLT_MEASURED_VALUE_3}}': '{{VOLT_MEASURED_VALUE_3}}',
        '{VOLT_UNCERTAINTY_3}}': '{{VOLT_UNCERTAINTY_3}}',
        '{VOLT_TOLERANCE_LIMIT_3}}': '{{VOLT_TOLERANCE_LIMIT_3}}',
        '{ERROR_4}}': '{{ERROR_4}}',
        '{FREQ_UUC_RANGE_4}}': '{{FREQ_UUC_RANGE_4}}',
        '{FREQ_UUC_SETTING_4}}': '{{FREQ_UUC_SETTING_4}}',
        '{FREQ_MEASURED_VALUE_4}}': '{{FREQ_MEASURED_VALUE_4}}',
        '{FREQ_UNCERTAINTY_4}}': '{{FREQ_UNCERTAINTY_4}}',
        '{FREQ_TOLERANCE_LIMIT_4}}': '{{FREQ_TOLERANCE_LIMIT_4}}',
        '{VOLT_UUC_RANGE_4}}': '{{VOLT_UUC_RANGE_4}}',
        '{VOLT_UUC_SETTING_4}}': '{{VOLT_UUC_SETTING_4}}',
        '{VOLT_MEASURED_VALUE_4}}': '{{VOLT_MEASURED_VALUE_4}}',
        '{VOLT_UNCERTAINTY_4}}': '{{VOLT_UNCERTAINTY_4}}',
        '{VOLT_TOLERANCE_LIMIT_4}}': '{{VOLT_TOLERANCE_LIMIT_4}}',
        '{ERROR_5}}': '{{ERROR_5}}',
        '{FREQ_UUC_RANGE_5}}': '{{FREQ_UUC_RANGE_5}}',
        '{FREQ_UUC_SETTING_5}}': '{{FREQ_UUC_SETTING_5}}',
        '{FREQ_MEASURED_VALUE_5}}': '{{FREQ_MEASURED_VALUE_5}}',
        '{FREQ_UNCERTAINTY_5}}': '{{FREQ_UNCERTAINTY_5}}',
        '{FREQ_TOLERANCE_LIMIT_5}}': '{{FREQ_TOLERANCE_LIMIT_5}}',
        '{VOLT_UUC_RANGE_5}}': '{{VOLT_UUC_RANGE_5}}',
        '{VOLT_UUC_SETTING_5}}': '{{VOLT_UUC_SETTING_5}}',
        '{VOLT_MEASURED_VALUE_5}}': '{{VOLT_MEASURED_VALUE_5}}',
        '{VOLT_UNCERTAINTY_5}}': '{{VOLT_UNCERTAINTY_5}}',
        '{VOLT_TOLERANCE_LIMIT_5}}': '{{VOLT_TOLERANCE_LIMIT_5}}',
    }
    
    # Debug: แสดงข้อมูลใน paragraph
    if paragraph.text.strip():
        print(f"DEBUG: Paragraph text: '{paragraph.text}'")
        for placeholder in ['{{FREQ_UUC_RANGE}}', '{{FREQ_UUC_SETTING}}', '{{FREQ_MEASURED_VALUE}}']:
            if placeholder in paragraph.text:
                print(f"DEBUG: Found placeholder {placeholder} in paragraph")
    
    for wrong_placeholder, correct_placeholder in problematic_placeholders.items():
        if wrong_placeholder in paragraph.text and correct_placeholder in replacements:
            for run in paragraph.runs:
                if wrong_placeholder in run.text:
                    run.text = run.text.replace(wrong_placeholder, str(replacements[correct_placeholder]))
                    print(f"DEBUG: Replaced {wrong_placeholder} with {replacements[correct_placeholder]}")
    
    # แทนที่ placeholder ที่ถูกต้อง
    correct_placeholders = {
        '{{FREQ_UUC_RANGE}}': replacements.get('{{FREQ_UUC_RANGE}}', '-'),
        '{{FREQ_UUC_SETTING}}': replacements.get('{{FREQ_UUC_SETTING}}', '-'),
        '{{FREQ_MEASURED_VALUE}}': replacements.get('{{FREQ_MEASURED_VALUE}}', '-'),
        '{{FREQ_UNCERTAINTY}}': replacements.get('{{FREQ_UNCERTAINTY}}', '-'),
        '{{FREQ_TOLERANCE_LIMIT}}': replacements.get('{{FREQ_TOLERANCE_LIMIT}}', '-'),
        '{{VOLT_UUC_RANGE}}': replacements.get('{{VOLT_UUC_RANGE}}', '-'),
        '{{VOLT_UUC_SETTING}}': replacements.get('{{VOLT_UUC_SETTING}}', '-'),
        '{{VOLT_MEASURED_VALUE}}': replacements.get('{{VOLT_MEASURED_VALUE}}', '-'),
        '{{VOLT_UNCERTAINTY}}': replacements.get('{{VOLT_UNCERTAINTY}}', '-'),
        '{{VOLT_TOLERANCE_LIMIT}}': replacements.get('{{VOLT_TOLERANCE_LIMIT}}', '-'),
        '{{DATE_OF_CALIBRATION}}': replacements.get('{{DATE_OF_CALIBRATION}}', '-'),
        '{{CUSTOMER_ADDRESS}}': replacements.get('{{CUSTOMER_ADDRESS}}', '-'),
        '{{RANGE}}': replacements.get('{{RANGE}}', '-'),
        '{{APPROVER}}': replacements.get('{{APPROVER}}', '-'),
        '{{CALIBRATOR}}': replacements.get('{{CALIBRATOR}}', '-'),
        '{{UUC_SET}}': replacements.get('{{UUC_SET}}', '-'),
        '{{ACTUAL}}': replacements.get('{{ACTUAL}}', '-'),
        '{{ERROR}}': replacements.get('{{ERROR}}', '-'),
        '{{TOLERANCE_LIMIT}}': replacements.get('{{TOLERANCE_LIMIT}}', '-'),
        '{{UUC_SET_2}}': replacements.get('{{UUC_SET_2}}', '-'),
        '{{ACTUAL_2}}': replacements.get('{{ACTUAL_2}}', '-'),
        '{{ERROR_2}}': replacements.get('{{ERROR_2}}', '-'),
        '{{TOLERANCE_LIMIT_2}}': replacements.get('{{TOLERANCE_LIMIT_2}}', '-'),
        '{{UUC_SET_3}}': replacements.get('{{UUC_SET_3}}', '-'),
        '{{ACTUAL_3}}': replacements.get('{{ACTUAL_3}}', '-'),
        '{{ERROR_3}}': replacements.get('{{ERROR_3}}', '-'),
        '{{TOLERANCE_LIMIT_3}}': replacements.get('{{TOLERANCE_LIMIT_3}}', '-'),
        '{{UUC_SET_4}}': replacements.get('{{UUC_SET_4}}', '-'),
        '{{ACTUAL_4}}': replacements.get('{{ACTUAL_4}}', '-'),
        '{{ERROR_4}}': replacements.get('{{ERROR_4}}', '-'),
        '{{TOLERANCE_LIMIT_4}}': replacements.get('{{TOLERANCE_LIMIT_4}}', '-'),
        '{{UUC_SET_5}}': replacements.get('{{UUC_SET_5}}', '-'),
        '{{ACTUAL_5}}': replacements.get('{{ACTUAL_5}}', '-'),
        '{{ERROR_5}}': replacements.get('{{ERROR_5}}', '-'),
        '{{TOLERANCE_LIMIT_5}}': replacements.get('{{TOLERANCE_LIMIT_5}}', '-'),
        '{{UUC_SET_6}}': replacements.get('{{UUC_SET_6}}', '-'),
        '{{ACTUAL_6}}': replacements.get('{{ACTUAL_6}}', '-'),
        '{{ERROR_6}}': replacements.get('{{ERROR_6}}', '-'),
        '{{TOLERANCE_LIMIT_6}}': replacements.get('{{TOLERANCE_LIMIT_6}}', '-'),
        '{{CW_SET}}': replacements.get('{{CW_SET}}', '-'),
        '{{CW_0}}': replacements.get('{{CW_0}}', '-'),
        '{{CW_90}}': replacements.get('{{CW_90}}', '-'),
        '{{CW_180}}': replacements.get('{{CW_180}}', '-'),
        '{{CW_270}}': replacements.get('{{CW_270}}', '-'),
        '{{CW_AVG}}': replacements.get('{{CW_AVG}}', '-'),
        '{{CW_ACTUAL}}': replacements.get('{{CW_ACTUAL}}', '-'),
        '{{CW_ERROR}}': replacements.get('{{CW_ERROR}}', '-'),
        '{{CW_UNCERTAINTY}}': replacements.get('{{CW_UNCERTAINTY}}', '-'),
        '{{CW_TOLERANCE_LIMIT}}': replacements.get('{{CW_TOLERANCE_LIMIT}}', '-'),
        '{{CCW_SET}}': replacements.get('{{CCW_SET}}', '-'),
        '{{CCW_0}}': replacements.get('{{CCW_0}}', '-'),
        '{{CCW_90}}': replacements.get('{{CCW_90}}', '-'),
        '{{CCW_180}}': replacements.get('{{CCW_180}}', '-'),
        '{{CCW_270}}': replacements.get('{{CCW_270}}', '-'),
        '{{CCW_AVG}}': replacements.get('{{CCW_AVG}}', '-'),
        '{{CCW_ACTUAL}}': replacements.get('{{CCW_ACTUAL}}', '-'),
        '{{CCW_ERROR}}': replacements.get('{{CCW_ERROR}}', '-'),
        '{{CCW_UNCERTAINTY}}': replacements.get('{{CCW_UNCERTAINTY}}', '-'),
        '{{CCW_TOLERANCE_LIMIT}}': replacements.get('{{CCW_TOLERANCE_LIMIT}}', '-'),
        '{{CCW_ACTUAL_2}}': replacements.get('{{CCW_ACTUAL_2}}', '-'),
        '{{CCW_ERROR_2}}': replacements.get('{{CCW_ERROR_2}}', '-'),
        '{{CCW_TOLERANCE_LIMIT_2}}': replacements.get('{{CCW_TOLERANCE_LIMIT_2}}', '-'),
        '{{CCW_ACTUAL_3}}': replacements.get('{{CCW_ACTUAL_3}}', '-'),
        '{{CCW_ERROR_3}}': replacements.get('{{CCW_ERROR_3}}', '-'),
        '{{CCW_TOLERANCE_LIMIT_3}}': replacements.get('{{CCW_TOLERANCE_LIMIT_3}}', '-'),
        '{{FREQ_UUC_RANGE_2}}': replacements.get('{{FREQ_UUC_RANGE_2}}', '-'),
        '{{FREQ_UUC_SETTING_2}}': replacements.get('{{FREQ_UUC_SETTING_2}}', '-'),
        '{{FREQ_MEASURED_VALUE_2}}': replacements.get('{{FREQ_MEASURED_VALUE_2}}', '-'),
        '{{FREQ_UNCERTAINTY_2}}': replacements.get('{{FREQ_UNCERTAINTY_2}}', '-'),
        '{{FREQ_TOLERANCE_LIMIT_2}}': replacements.get('{{FREQ_TOLERANCE_LIMIT_2}}', '-'),
        '{{VOLT_UUC_RANGE_2}}': replacements.get('{{VOLT_UUC_RANGE_2}}', '-'),
        '{{VOLT_UUC_SETTING_2}}': replacements.get('{{VOLT_UUC_SETTING_2}}', '-'),
        '{{VOLT_MEASURED_VALUE_2}}': replacements.get('{{VOLT_MEASURED_VALUE_2}}', '-'),
        '{{VOLT_UNCERTAINTY_2}}': replacements.get('{{VOLT_UNCERTAINTY_2}}', '-'),
        '{{VOLT_TOLERANCE_LIMIT_2}}': replacements.get('{{VOLT_TOLERANCE_LIMIT_2}}', '-'),
        '{{FREQ_UUC_RANGE_3}}': replacements.get('{{FREQ_UUC_RANGE_3}}', '-'),
        '{{FREQ_UUC_SETTING_3}}': replacements.get('{{FREQ_UUC_SETTING_3}}', '-'),
        '{{FREQ_MEASURED_VALUE_3}}': replacements.get('{{FREQ_MEASURED_VALUE_3}}', '-'),
        '{{FREQ_UNCERTAINTY_3}}': replacements.get('{{FREQ_UNCERTAINTY_3}}', '-'),
        '{{FREQ_TOLERANCE_LIMIT_3}}': replacements.get('{{FREQ_TOLERANCE_LIMIT_3}}', '-'),
        '{{VOLT_UUC_RANGE_3}}': replacements.get('{{VOLT_UUC_RANGE_3}}', '-'),
        '{{VOLT_UUC_SETTING_3}}': replacements.get('{{VOLT_UUC_SETTING_3}}', '-'),
        '{{VOLT_MEASURED_VALUE_3}}': replacements.get('{{VOLT_MEASURED_VALUE_3}}', '-'),
        '{{VOLT_UNCERTAINTY_3}}': replacements.get('{{VOLT_UNCERTAINTY_3}}', '-'),
        '{{VOLT_TOLERANCE_LIMIT_3}}': replacements.get('{{VOLT_TOLERANCE_LIMIT_3}}', '-'),
        '{{FREQ_UUC_RANGE_4}}': replacements.get('{{FREQ_UUC_RANGE_4}}', '-'),
        '{{FREQ_UUC_SETTING_4}}': replacements.get('{{FREQ_UUC_SETTING_4}}', '-'),
        '{{FREQ_MEASURED_VALUE_4}}': replacements.get('{{FREQ_MEASURED_VALUE_4}}', '-'),
        '{{FREQ_UNCERTAINTY_4}}': replacements.get('{{FREQ_UNCERTAINTY_4}}', '-'),
        '{{FREQ_TOLERANCE_LIMIT_4}}': replacements.get('{{FREQ_TOLERANCE_LIMIT_4}}', '-'),
        '{{VOLT_UUC_RANGE_4}}': replacements.get('{{VOLT_UUC_RANGE_4}}', '-'),
        '{{VOLT_UUC_SETTING_4}}': replacements.get('{{VOLT_UUC_SETTING_4}}', '-'),
        '{{VOLT_MEASURED_VALUE_4}}': replacements.get('{{VOLT_MEASURED_VALUE_4}}', '-'),
        '{{VOLT_UNCERTAINTY_4}}': replacements.get('{{VOLT_UNCERTAINTY_4}}', '-'),
        '{{VOLT_TOLERANCE_LIMIT_4}}': replacements.get('{{VOLT_TOLERANCE_LIMIT_4}}', '-'),
        '{{FREQ_UUC_RANGE_5}}': replacements.get('{{FREQ_UUC_RANGE_5}}', '-'),
        '{{FREQ_UUC_SETTING_5}}': replacements.get('{{FREQ_UUC_SETTING_5}}', '-'),
        '{{FREQ_MEASURED_VALUE_5}}': replacements.get('{{FREQ_MEASURED_VALUE_5}}', '-'),
        '{{FREQ_UNCERTAINTY_5}}': replacements.get('{{FREQ_UNCERTAINTY_5}}', '-'),
        '{{FREQ_TOLERANCE_LIMIT_5}}': replacements.get('{{FREQ_TOLERANCE_LIMIT_5}}', '-'),
        '{{VOLT_UUC_RANGE_5}}': replacements.get('{{VOLT_UUC_RANGE_5}}', '-'),
        '{{VOLT_UUC_SETTING_5}}': replacements.get('{{VOLT_UUC_SETTING_5}}', '-'),
        '{{VOLT_MEASURED_VALUE_5}}': replacements.get('{{VOLT_MEASURED_VALUE_5}}', '-'),
        '{{VOLT_UNCERTAINTY_5}}': replacements.get('{{VOLT_UNCERTAINTY_5}}', '-'),
        '{{VOLT_TOLERANCE_LIMIT_5}}': replacements.get('{{VOLT_TOLERANCE_LIMIT_5}}', '-'),
    }
    
    for placeholder, value in correct_placeholders.items():
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(value))
                    print(f"DEBUG: Replaced {placeholder} with {value}")

def replace_text_in_document(doc, replacements):
    """แทนค่าในเอกสาร"""
    # แทนค่าใน paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
    
    # แทนค่าใน tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)
    
    # แทนค่าใน headers และ footers
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)
        if section.footer:
            for paragraph in section.footer.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)
    
    # จัดการ placeholder ที่มีปัญหาในตาราง
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # แทนที่ placeholder ที่มีปัญหา
                    problematic_placeholders = {
                        '{ERROR}}': '{{ERROR}}',
                        '{FREQ_UUC_RANGE}}': '{{FREQ_UUC_RANGE}}',
                        '{FREQ_UUC_SETTING}}': '{{FREQ_UUC_SETTING}}',
                        '{FREQ_MEASURED_VALUE}}': '{{FREQ_MEASURED_VALUE}}',
                        '{FREQ_UNCERTAINTY}}': '{{FREQ_UNCERTAINTY}}',
                        '{FREQ_TOLERANCE_LIMIT}}': '{{FREQ_TOLERANCE_LIMIT}}',
                        '{VOLT_UUC_RANGE}}': '{{VOLT_UUC_RANGE}}',
                        '{VOLT_UUC_SETTING}}': '{{VOLT_UUC_SETTING}}',
                        '{VOLT_MEASURED_VALUE}}': '{{VOLT_MEASURED_VALUE}}',
                        '{VOLT_UNCERTAINTY}}': '{{VOLT_UNCERTAINTY}}',
                        '{VOLT_TOLERANCE_LIMIT}}': '{{VOLT_TOLERANCE_LIMIT}}',
                        '{DATE_OF_CALIBRATION}}': '{{DATE_OF_CALIBRATION}}',
                        '{CUSTOMER_ADDRESS}}': '{{CUSTOMER_ADDRESS}}',
                        '{RANGE}}': '{{RANGE}}',
                        '{APPROVER}}': '{{APPROVER}}',
                        '{CALIBRATOR}}': '{{CALIBRATOR}}',
                        '{UUC_SET}}': '{{UUC_SET}}',
                        '{ACTUAL}}': '{{ACTUAL}}',
                        '{ERROR}}': '{{ERROR}}',
                        '{TOLERANCE_LIMIT}}': '{{TOLERANCE_LIMIT}}',
                        '{UUC_SET_2}}': '{{UUC_SET_2}}',
                        '{ACTUAL_2}}': '{{ACTUAL_2}}',
                        '{ERROR_2}}': '{{ERROR_2}}',
                        '{TOLERANCE_LIMIT_2}}': '{{TOLERANCE_LIMIT_2}}',
                        '{UUC_SET_3}}': '{{UUC_SET_3}}',
                        '{ACTUAL_3}}': '{{ACTUAL_3}}',
                        '{ERROR_3}}': '{{ERROR_3}}',
                        '{TOLERANCE_LIMIT_3}}': '{{TOLERANCE_LIMIT_3}}',
                        '{UUC_SET_4}}': '{{UUC_SET_4}}',
                        '{ACTUAL_4}}': '{{ACTUAL_4}}',
                        '{ERROR_4}}': '{{ERROR_4}}',
                        '{TOLERANCE_LIMIT_4}}': '{{TOLERANCE_LIMIT_4}}',
                        '{UUC_SET_5}}': '{{UUC_SET_5}}',
                        '{ACTUAL_5}}': '{{ACTUAL_5}}',
                        '{ERROR_5}}': '{{ERROR_5}}',
                        '{TOLERANCE_LIMIT_5}}': '{{TOLERANCE_LIMIT_5}}',
                        '{UUC_SET_6}}': '{{UUC_SET_6}}',
                        '{ACTUAL_6}}': '{{ACTUAL_6}}',
                        '{ERROR_6}}': '{{ERROR_6}}',
                        '{TOLERANCE_LIMIT_6}}': '{{TOLERANCE_LIMIT_6}}',
                        '{CW_SET}}': '{{CW_SET}}',
                        '{CW_0}}': '{{CW_0}}',
                        '{CW_90}}': '{{CW_90}}',
                        '{CW_180}}': '{{CW_180}}',
                        '{CW_270}}': '{{CW_270}}',
                        '{CW_AVG}}': '{{CW_AVG}}',
                        '{CW_ACTUAL}}': '{{CW_ACTUAL}}',
                        '{CW_ERROR}}': '{{CW_ERROR}}',
                        '{CW_UNCERTAINTY}}': '{{CW_UNCERTAINTY}}',
                        '{CW_TOLERANCE_LIMIT}}': '{{CW_TOLERANCE_LIMIT}}',
                        '{CCW_SET}}': '{{CCW_SET}}',
                        '{CCW_0}}': '{{CCW_0}}',
                        '{CCW_90}}': '{{CCW_90}}',
                        '{CCW_180}}': '{{CCW_180}}',
                        '{CCW_270}}': '{{CCW_270}}',
                        '{CCW_AVG}}': '{{CCW_AVG}}',
                        '{CCW_ACTUAL}}': '{{CCW_ACTUAL}}',
                        '{CCW_ERROR}}': '{{CCW_ERROR}}',
                        '{CCW_UNCERTAINTY}}': '{{CCW_UNCERTAINTY}}',
                        '{CCW_TOLERANCE_LIMIT}}': '{{CCW_TOLERANCE_LIMIT}}',
                        '{CCW_ACTUAL_2}}': '{{CCW_ACTUAL_2}}',
                        '{CCW_ERROR_2}}': '{{CCW_ERROR_2}}',
                        '{CCW_TOLERANCE_LIMIT_2}}': '{{CCW_TOLERANCE_LIMIT_2}}',
                        '{CCW_ACTUAL_3}}': '{{CCW_ACTUAL_3}}',
                        '{CCW_ERROR_3}}': '{{CCW_ERROR_3}}',
                        '{CCW_TOLERANCE_LIMIT_3}}': '{{CCW_TOLERANCE_LIMIT_3}}',
                        '{ERROR_2}}': '{{ERROR_2}}',
                        '{FREQ_UUC_RANGE_2}}': '{{FREQ_UUC_RANGE_2}}',
                        '{FREQ_UUC_SETTING_2}}': '{{FREQ_UUC_SETTING_2}}',
                        '{FREQ_MEASURED_VALUE_2}}': '{{FREQ_MEASURED_VALUE_2}}',
                        '{FREQ_UNCERTAINTY_2}}': '{{FREQ_UNCERTAINTY_2}}',
                        '{FREQ_TOLERANCE_LIMIT_2}}': '{{FREQ_TOLERANCE_LIMIT_2}}',
                        '{VOLT_UUC_RANGE_2}}': '{{VOLT_UUC_RANGE_2}}',
                        '{VOLT_UUC_SETTING_2}}': '{{VOLT_UUC_SETTING_2}}',
                        '{VOLT_MEASURED_VALUE_2}}': '{{VOLT_MEASURED_VALUE_2}}',
                        '{VOLT_UNCERTAINTY_2}}': '{{VOLT_UNCERTAINTY_2}}',
                        '{VOLT_TOLERANCE_LIMIT_2}}': '{{VOLT_TOLERANCE_LIMIT_2}}',
                        '{ERROR_3}}': '{{ERROR_3}}',
                        '{FREQ_UUC_RANGE_3}}': '{{FREQ_UUC_RANGE_3}}',
                        '{FREQ_UUC_SETTING_3}}': '{{FREQ_UUC_SETTING_3}}',
                        '{FREQ_MEASURED_VALUE_3}}': '{{FREQ_MEASURED_VALUE_3}}',
                        '{FREQ_UNCERTAINTY_3}}': '{{FREQ_UNCERTAINTY_3}}',
                        '{FREQ_TOLERANCE_LIMIT_3}}': '{{FREQ_TOLERANCE_LIMIT_3}}',
                        '{VOLT_UUC_RANGE_3}}': '{{VOLT_UUC_RANGE_3}}',
                        '{VOLT_UUC_SETTING_3}}': '{{VOLT_UUC_SETTING_3}}',
                        '{VOLT_MEASURED_VALUE_3}}': '{{VOLT_MEASURED_VALUE_3}}',
                        '{VOLT_UNCERTAINTY_3}}': '{{VOLT_UNCERTAINTY_3}}',
                        '{VOLT_TOLERANCE_LIMIT_3}}': '{{VOLT_TOLERANCE_LIMIT_3}}',
                        '{ERROR_4}}': '{{ERROR_4}}',
                        '{FREQ_UUC_RANGE_4}}': '{{FREQ_UUC_RANGE_4}}',
                        '{FREQ_UUC_SETTING_4}}': '{{FREQ_UUC_SETTING_4}}',
                        '{FREQ_MEASURED_VALUE_4}}': '{{FREQ_MEASURED_VALUE_4}}',
                        '{FREQ_UNCERTAINTY_4}}': '{{FREQ_UNCERTAINTY_4}}',
                        '{FREQ_TOLERANCE_LIMIT_4}}': '{{FREQ_TOLERANCE_LIMIT_4}}',
                        '{VOLT_UUC_RANGE_4}}': '{{VOLT_UUC_RANGE_4}}',
                        '{VOLT_UUC_SETTING_4}}': '{{VOLT_UUC_SETTING_4}}',
                        '{VOLT_MEASURED_VALUE_4}}': '{{VOLT_MEASURED_VALUE_4}}',
                        '{VOLT_UNCERTAINTY_4}}': '{{VOLT_UNCERTAINTY_4}}',
                        '{VOLT_TOLERANCE_LIMIT_4}}': '{{VOLT_TOLERANCE_LIMIT_4}}',
                        '{ERROR_5}}': '{{ERROR_5}}',
                        '{FREQ_UUC_RANGE_5}}': '{{FREQ_UUC_RANGE_5}}',
                        '{FREQ_UUC_SETTING_5}}': '{{FREQ_UUC_SETTING_5}}',
                        '{FREQ_MEASURED_VALUE_5}}': '{{FREQ_MEASURED_VALUE_5}}',
                        '{FREQ_UNCERTAINTY_5}}': '{{FREQ_UNCERTAINTY_5}}',
                        '{FREQ_TOLERANCE_LIMIT_5}}': '{{FREQ_TOLERANCE_LIMIT_5}}',
                        '{VOLT_UUC_RANGE_5}}': '{{VOLT_UUC_RANGE_5}}',
                        '{VOLT_UUC_SETTING_5}}': '{{VOLT_UUC_SETTING_5}}',
                        '{VOLT_MEASURED_VALUE_5}}': '{{VOLT_MEASURED_VALUE_5}}',
                        '{VOLT_UNCERTAINTY_5}}': '{{VOLT_UNCERTAINTY_5}}',
                        '{VOLT_TOLERANCE_LIMIT_5}}': '{{VOLT_TOLERANCE_LIMIT_5}}',
                    }
                    
                    # Debug: แสดงข้อมูลในตาราง
                    if paragraph.text.strip():
                        print(f"DEBUG: Table cell text: '{paragraph.text}'")
                        for placeholder in ['{{FREQ_UUC_RANGE}}', '{{FREQ_UUC_SETTING}}', '{{FREQ_MEASURED_VALUE}}']:
                            if placeholder in paragraph.text:
                                print(f"DEBUG: Found placeholder {placeholder} in table cell")
                    
                    # แทนที่ placeholder ที่ถูกต้อง
                    correct_placeholders = {
                        '{{FREQ_UUC_RANGE}}': replacements.get('{{FREQ_UUC_RANGE}}', '-'),
                        '{{FREQ_UUC_SETTING}}': replacements.get('{{FREQ_UUC_SETTING}}', '-'),
                        '{{FREQ_MEASURED_VALUE}}': replacements.get('{{FREQ_MEASURED_VALUE}}', '-'),
                        '{{FREQ_UNCERTAINTY}}': replacements.get('{{FREQ_UNCERTAINTY}}', '-'),
                        '{{FREQ_TOLERANCE_LIMIT}}': replacements.get('{{FREQ_TOLERANCE_LIMIT}}', '-'),
                        '{{VOLT_UUC_RANGE}}': replacements.get('{{VOLT_UUC_RANGE}}', '-'),
                        '{{VOLT_UUC_SETTING}}': replacements.get('{{VOLT_UUC_SETTING}}', '-'),
                        '{{VOLT_MEASURED_VALUE}}': replacements.get('{{VOLT_MEASURED_VALUE}}', '-'),
                        '{{VOLT_UNCERTAINTY}}': replacements.get('{{VOLT_UNCERTAINTY}}', '-'),
                        '{{VOLT_TOLERANCE_LIMIT}}': replacements.get('{{VOLT_TOLERANCE_LIMIT}}', '-'),
                        '{{DATE_OF_CALIBRATION}}': replacements.get('{{DATE_OF_CALIBRATION}}', '-'),
                        '{{CUSTOMER_ADDRESS}}': replacements.get('{{CUSTOMER_ADDRESS}}', '-'),
                        '{{RANGE}}': replacements.get('{{RANGE}}', '-'),
                        '{{APPROVER}}': replacements.get('{{APPROVER}}', '-'),
                        '{{CALIBRATOR}}': replacements.get('{{CALIBRATOR}}', '-'),
                        '{{UUC_SET}}': replacements.get('{{UUC_SET}}', '-'),
                        '{{ACTUAL}}': replacements.get('{{ACTUAL}}', '-'),
                        '{{ERROR}}': replacements.get('{{ERROR}}', '-'),
                        '{{TOLERANCE_LIMIT}}': replacements.get('{{TOLERANCE_LIMIT}}', '-'),
                        '{{UUC_SET_2}}': replacements.get('{{UUC_SET_2}}', '-'),
                        '{{ACTUAL_2}}': replacements.get('{{ACTUAL_2}}', '-'),
                        '{{ERROR_2}}': replacements.get('{{ERROR_2}}', '-'),
                        '{{TOLERANCE_LIMIT_2}}': replacements.get('{{TOLERANCE_LIMIT_2}}', '-'),
                        '{{UUC_SET_3}}': replacements.get('{{UUC_SET_3}}', '-'),
                        '{{ACTUAL_3}}': replacements.get('{{ACTUAL_3}}', '-'),
                        '{{ERROR_3}}': replacements.get('{{ERROR_3}}', '-'),
                        '{{TOLERANCE_LIMIT_3}}': replacements.get('{{TOLERANCE_LIMIT_3}}', '-'),
                        '{{UUC_SET_4}}': replacements.get('{{UUC_SET_4}}', '-'),
                        '{{ACTUAL_4}}': replacements.get('{{ACTUAL_4}}', '-'),
                        '{{ERROR_4}}': replacements.get('{{ERROR_4}}', '-'),
                        '{{TOLERANCE_LIMIT_4}}': replacements.get('{{TOLERANCE_LIMIT_4}}', '-'),
                        '{{UUC_SET_5}}': replacements.get('{{UUC_SET_5}}', '-'),
                        '{{ACTUAL_5}}': replacements.get('{{ACTUAL_5}}', '-'),
                        '{{ERROR_5}}': replacements.get('{{ERROR_5}}', '-'),
                        '{{TOLERANCE_LIMIT_5}}': replacements.get('{{TOLERANCE_LIMIT_5}}', '-'),
                        '{{UUC_SET_6}}': replacements.get('{{UUC_SET_6}}', '-'),
                        '{{ACTUAL_6}}': replacements.get('{{ACTUAL_6}}', '-'),
                        '{{ERROR_6}}': replacements.get('{{ERROR_6}}', '-'),
                        '{{TOLERANCE_LIMIT_6}}': replacements.get('{{TOLERANCE_LIMIT_6}}', '-'),
                        '{{CW_SET}}': replacements.get('{{CW_SET}}', '-'),
                        '{{CW_0}}': replacements.get('{{CW_0}}', '-'),
                        '{{CW_90}}': replacements.get('{{CW_90}}', '-'),
                        '{{CW_180}}': replacements.get('{{CW_180}}', '-'),
                        '{{CW_270}}': replacements.get('{{CW_270}}', '-'),
                        '{{CW_AVG}}': replacements.get('{{CW_AVG}}', '-'),
                        '{{CW_ACTUAL}}': replacements.get('{{CW_ACTUAL}}', '-'),
                        '{{CW_ERROR}}': replacements.get('{{CW_ERROR}}', '-'),
                        '{{CW_UNCERTAINTY}}': replacements.get('{{CW_UNCERTAINTY}}', '-'),
                        '{{CW_TOLERANCE_LIMIT}}': replacements.get('{{CW_TOLERANCE_LIMIT}}', '-'),
                        '{{CCW_SET}}': replacements.get('{{CCW_SET}}', '-'),
                        '{{CCW_0}}': replacements.get('{{CCW_0}}', '-'),
                        '{{CCW_90}}': replacements.get('{{CCW_90}}', '-'),
                        '{{CCW_180}}': replacements.get('{{CCW_180}}', '-'),
                        '{{CCW_270}}': replacements.get('{{CCW_270}}', '-'),
                        '{{CCW_AVG}}': replacements.get('{{CCW_AVG}}', '-'),
                        '{{CCW_ACTUAL}}': replacements.get('{{CCW_ACTUAL}}', '-'),
                        '{{CCW_ERROR}}': replacements.get('{{CCW_ERROR}}', '-'),
                        '{{CCW_UNCERTAINTY}}': replacements.get('{{CCW_UNCERTAINTY}}', '-'),
                        '{{CCW_TOLERANCE_LIMIT}}': replacements.get('{{CCW_TOLERANCE_LIMIT}}', '-'),
                        '{{CCW_ACTUAL_2}}': replacements.get('{{CCW_ACTUAL_2}}', '-'),
                        '{{CCW_ERROR_2}}': replacements.get('{{CCW_ERROR_2}}', '-'),
                        '{{CCW_TOLERANCE_LIMIT_2}}': replacements.get('{{CCW_TOLERANCE_LIMIT_2}}', '-'),
                        '{{CCW_ACTUAL_3}}': replacements.get('{{CCW_ACTUAL_3}}', '-'),
                        '{{CCW_ERROR_3}}': replacements.get('{{CCW_ERROR_3}}', '-'),
                        '{{CCW_TOLERANCE_LIMIT_3}}': replacements.get('{{CCW_TOLERANCE_LIMIT_3}}', '-'),
                        '{{FREQ_UUC_RANGE_2}}': replacements.get('{{FREQ_UUC_RANGE_2}}', '-'),
                        '{{FREQ_UUC_SETTING_2}}': replacements.get('{{FREQ_UUC_SETTING_2}}', '-'),
                        '{{FREQ_MEASURED_VALUE_2}}': replacements.get('{{FREQ_MEASURED_VALUE_2}}', '-'),
                        '{{FREQ_UNCERTAINTY_2}}': replacements.get('{{FREQ_UNCERTAINTY_2}}', '-'),
                        '{{FREQ_TOLERANCE_LIMIT_2}}': replacements.get('{{FREQ_TOLERANCE_LIMIT_2}}', '-'),
                        '{{VOLT_UUC_RANGE_2}}': replacements.get('{{VOLT_UUC_RANGE_2}}', '-'),
                        '{{VOLT_UUC_SETTING_2}}': replacements.get('{{VOLT_UUC_SETTING_2}}', '-'),
                        '{{VOLT_MEASURED_VALUE_2}}': replacements.get('{{VOLT_MEASURED_VALUE_2}}', '-'),
                        '{{VOLT_UNCERTAINTY_2}}': replacements.get('{{VOLT_UNCERTAINTY_2}}', '-'),
                        '{{VOLT_TOLERANCE_LIMIT_2}}': replacements.get('{{VOLT_TOLERANCE_LIMIT_2}}', '-'),
                        '{{FREQ_UUC_RANGE_3}}': replacements.get('{{FREQ_UUC_RANGE_3}}', '-'),
                        '{{FREQ_UUC_SETTING_3}}': replacements.get('{{FREQ_UUC_SETTING_3}}', '-'),
                        '{{FREQ_MEASURED_VALUE_3}}': replacements.get('{{FREQ_MEASURED_VALUE_3}}', '-'),
                        '{{FREQ_UNCERTAINTY_3}}': replacements.get('{{FREQ_UNCERTAINTY_3}}', '-'),
                        '{{FREQ_TOLERANCE_LIMIT_3}}': replacements.get('{{FREQ_TOLERANCE_LIMIT_3}}', '-'),
                        '{{VOLT_UUC_RANGE_3}}': replacements.get('{{VOLT_UUC_RANGE_3}}', '-'),
                        '{{VOLT_UUC_SETTING_3}}': replacements.get('{{VOLT_UUC_SETTING_3}}', '-'),
                        '{{VOLT_MEASURED_VALUE_3}}': replacements.get('{{VOLT_MEASURED_VALUE_3}}', '-'),
                        '{{VOLT_UNCERTAINTY_3}}': replacements.get('{{VOLT_UNCERTAINTY_3}}', '-'),
                        '{{VOLT_TOLERANCE_LIMIT_3}}': replacements.get('{{VOLT_TOLERANCE_LIMIT_3}}', '-'),
                        '{{FREQ_UUC_RANGE_4}}': replacements.get('{{FREQ_UUC_RANGE_4}}', '-'),
                        '{{FREQ_UUC_SETTING_4}}': replacements.get('{{FREQ_UUC_SETTING_4}}', '-'),
                        '{{FREQ_MEASURED_VALUE_4}}': replacements.get('{{FREQ_MEASURED_VALUE_4}}', '-'),
                        '{{FREQ_UNCERTAINTY_4}}': replacements.get('{{FREQ_UNCERTAINTY_4}}', '-'),
                        '{{FREQ_TOLERANCE_LIMIT_4}}': replacements.get('{{FREQ_TOLERANCE_LIMIT_4}}', '-'),
                        '{{VOLT_UUC_RANGE_4}}': replacements.get('{{VOLT_UUC_RANGE_4}}', '-'),
                        '{{VOLT_UUC_SETTING_4}}': replacements.get('{{VOLT_UUC_SETTING_4}}', '-'),
                        '{{VOLT_MEASURED_VALUE_4}}': replacements.get('{{VOLT_MEASURED_VALUE_4}}', '-'),
                        '{{VOLT_UNCERTAINTY_4}}': replacements.get('{{VOLT_UNCERTAINTY_4}}', '-'),
                        '{{VOLT_TOLERANCE_LIMIT_4}}': replacements.get('{{VOLT_TOLERANCE_LIMIT_4}}', '-'),
                        '{{FREQ_UUC_RANGE_5}}': replacements.get('{{FREQ_UUC_RANGE_5}}', '-'),
                        '{{FREQ_UUC_SETTING_5}}': replacements.get('{{FREQ_UUC_SETTING_5}}', '-'),
                        '{{FREQ_MEASURED_VALUE_5}}': replacements.get('{{FREQ_MEASURED_VALUE_5}}', '-'),
                        '{{FREQ_UNCERTAINTY_5}}': replacements.get('{{FREQ_UNCERTAINTY_5}}', '-'),
                        '{{FREQ_TOLERANCE_LIMIT_5}}': replacements.get('{{FREQ_TOLERANCE_LIMIT_5}}', '-'),
                        '{{VOLT_UUC_RANGE_5}}': replacements.get('{{VOLT_UUC_RANGE_5}}', '-'),
                        '{{VOLT_UUC_SETTING_5}}': replacements.get('{{VOLT_UUC_SETTING_5}}', '-'),
                        '{{VOLT_MEASURED_VALUE_5}}': replacements.get('{{VOLT_MEASURED_VALUE_5}}', '-'),
                        '{{VOLT_UNCERTAINTY_5}}': replacements.get('{{VOLT_UNCERTAINTY_5}}', '-'),
                        '{{VOLT_TOLERANCE_LIMIT_5}}': replacements.get('{{VOLT_TOLERANCE_LIMIT_5}}', '-'),
                    }
                    
                    # แทนที่ placeholder ที่มีปัญหา
                    for wrong_placeholder, correct_placeholder in problematic_placeholders.items():
                        if wrong_placeholder in paragraph.text and correct_placeholder in replacements:
                            for run in paragraph.runs:
                                if wrong_placeholder in run.text:
                                    run.text = run.text.replace(wrong_placeholder, str(replacements[correct_placeholder]))
                                    print(f"DEBUG: Replaced {wrong_placeholder} with {replacements[correct_placeholder]}")
                    
                    # แทนที่ placeholder ที่ถูกต้อง
                    for placeholder, value in correct_placeholders.items():
                        if placeholder in paragraph.text:
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(value))
                                    print(f"DEBUG: Replaced {placeholder} with {value}")

def export_high_frequency_certificate_docx(request, cal_id):
    """Export High Frequency certificate as DOCX"""
    try:
        print(f"DEBUG: Attempting to export High Frequency certificate for ID: {cal_id}")
        
        # ดึงข้อมูลการสอบเทียบ
        cal = get_object_or_404(HighFrequencyCalibration, id=cal_id)
        print(f"DEBUG: Found calibration - Status: {cal.status}, Machine: {cal.machine.name}")
        
        # ดึงข้อมูลเครื่องมือมาตรฐาน
        std = cal.std_id
        if std:
            print(f"DEBUG: Standard equipment found: {std.name}")
        else:
            print("DEBUG: No standard equipment found")
        
        # สร้าง path ไปยัง template
        template_path = os.path.join(settings.BASE_DIR, 'cert_templates', 'High Frequency_template.docx')
        print(f"DEBUG: Template path: {template_path}")
        
        # เปิด template
        doc = Document(template_path)
        
        # สร้าง replacements dictionary
        replacements = {
            # ข้อมูลเครื่องมือ
            '{{MODEL}}': getattr(cal.machine, 'model', '-'),
            '{{MANUFACTURER}}': getattr(cal.machine, 'manufacturer', '-'),
            '{{DESCRIPTION}}': getattr(cal.machine, 'description', '-'),
            '{{SERIAL_NUMBER}}': getattr(cal.machine, 'serial_number', '-'),
            '{{ASSET_NUMBER}}': getattr(cal.machine, 'asset_number', '-'),
            '{{RANGE}}': getattr(cal, 'measurement_range', '-'),
            
            # ข้อมูลการสอบเทียบ
            '{{CALIBRATION_DATE}}': cal.date_calibration.strftime('%d/%m/%Y') if cal.date_calibration else '-',
            '{{DATE_OF_CALIBRATION}}': cal.date_calibration.strftime('%d/%m/%Y') if cal.date_calibration else '-',
            '{{CERTIFICATE_NUMBER}}': cal.certificate_number or '-',
            '{{DUE_DATE}}': cal.next_due.strftime('%d/%m/%Y') if cal.next_due else '-',
            '{{STATUS}}': cal.get_status_display(),
            
            # ข้อมูลผู้รับผิดชอบ
            '{{CALIBRATOR}}': cal.calibrator.username if cal.calibrator else '-',
            '{{APPROVER}}': cal.certificate_issuer.username if cal.certificate_issuer else '-',
            
            # ข้อมูลหน่วยงาน
            '{{CUSTOMER}}': getattr(cal.machine.organize, 'name', '-') if cal.machine.organize else '-',
            '{{CUSTOMER_ADDRESS}}': getattr(cal.machine.organize, 'address', '-') if cal.machine.organize else '-',
            '{{LOCATION_NAME}}': getattr(cal.machine.organize, 'name', '-') if cal.machine.organize else '-',
            '{{LOCATION_ADDRESS}}': getattr(cal.machine.organize, 'address', '-') if cal.machine.organize else '-',
            
            # ข้อมูลตารางผลการสอบเทียบ Frequency (5 แถว)
            '{{FREQ_UUC_RANGE}}': cal.freq_uuc_range or '-',
            '{{FREQ_UUC_SETTING}}': cal.freq_uuc_setting or '-',
            '{{FREQ_MEASURED_VALUE}}': cal.freq_measured_value or '-',
            '{{FREQ_UNCERTAINTY}}': cal.freq_uncertainty or '-',
            '{{FREQ_TOLERANCE_LIMIT}}': cal.freq_tolerance_limit or '-',
            
                    '{{FREQ_UUC_RANGE_2}}': cal.freq_uuc_range_2 or '-',
                    '{{FREQ_UUC_SETTING_2}}': cal.freq_uuc_setting_2 or '-',
                    '{{VOLT_UUC_RANGE_2}}': cal.freq_uuc_setting_2 or '-',  # แก้ไข placeholder ที่ผิดใน template
                    '{{FREQ_MEASURED_VALUE_2}}': cal.freq_measured_value_2 or '-',
                    '{{FREQ_UNCERTAINTY_2}}': cal.freq_uncertainty_2 or '-',
                    '{{FREQ_TOLERANCE_LIMIT_2}}': cal.freq_tolerance_limit_2 or '-',
            
            '{{FREQ_UUC_RANGE_3}}': cal.freq_uuc_range_3 or '-',
            '{{FREQ_UUC_SETTING_3}}': cal.freq_uuc_setting_3 or '-',
            '{{FREQ_MEASURED_VALUE_3}}': cal.freq_measured_value_3 or '-',
            '{{FREQ_UNCERTAINTY_3}}': cal.freq_uncertainty_3 or '-',
            '{{FREQ_TOLERANCE_LIMIT_3}}': cal.freq_tolerance_limit_3 or '-',
            
            '{{FREQ_UUC_RANGE_4}}': cal.freq_uuc_range_4 or '-',
            '{{FREQ_UUC_SETTING_4}}': cal.freq_uuc_setting_4 or '-',
            '{{FREQ_MEASURED_VALUE_4}}': cal.freq_measured_value_4 or '-',
            '{{FREQ_UNCERTAINTY_4}}': cal.freq_uncertainty_4 or '-',
            '{{FREQ_TOLERANCE_LIMIT_4}}': cal.freq_tolerance_limit_4 or '-',
            
            '{{FREQ_UUC_RANGE_5}}': cal.freq_uuc_range_5 or '-',
            '{{FREQ_UUC_SETTING_5}}': cal.freq_uuc_setting_5 or '-',
            '{{FREQ_MEASURED_VALUE_5}}': cal.freq_measured_value_5 or '-',
            '{{FREQ_UNCERTAINTY_5}}': cal.freq_uncertainty_5 or '-',
            '{{FREQ_TOLERANCE_LIMIT_5}}': cal.freq_tolerance_limit_5 or '-',
            
            # ข้อมูลตารางผลการสอบเทียบ Voltage (5 แถว)
            '{{VOLT_UUC_RANGE}}': cal.volt_uuc_range or '-',
            '{{VOLT_UUC_SETTING}}': cal.volt_uuc_setting or '-',
            '{{VOLT_MEASURED_VALUE}}': cal.volt_measured_value or '-',
            '{{VOLT_UNCERTAINTY}}': cal.volt_uncertainty or '-',
            '{{VOLT_TOLERANCE_LIMIT}}': cal.volt_tolerance_limit or '-',
            
            '{{VOLT_UUC_RANGE_2}}': cal.volt_uuc_range_2 or '-',
            '{{VOLT_UUC_SETTING_2}}': cal.volt_uuc_setting_2 or '-',
            '{{VOLT_MEASURED_VALUE_2}}': cal.volt_measured_value_2 or '-',
            '{{VOLT_UNCERTAINTY_2}}': cal.volt_uncertainty_2 or '-',
            '{{VOLT_TOLERANCE_LIMIT_2}}': cal.volt_tolerance_limit_2 or '-',
            
            '{{VOLT_UUC_RANGE_3}}': cal.volt_uuc_range_3 or '-',
            '{{VOLT_UUC_SETTING_3}}': cal.volt_uuc_setting_3 or '-',
            '{{VOLT_MEASURED_VALUE_3}}': cal.volt_measured_value_3 or '-',
            '{{VOLT_UNCERTAINTY_3}}': cal.volt_uncertainty_3 or '-',
            '{{VOLT_TOLERANCE_LIMIT_3}}': cal.volt_tolerance_limit_3 or '-',
            
            '{{VOLT_UUC_RANGE_4}}': cal.volt_uuc_range_4 or '-',
            '{{VOLT_UUC_SETTING_4}}': cal.volt_uuc_setting_4 or '-',
            '{{VOLT_MEASURED_VALUE_4}}': cal.volt_measured_value_4 or '-',
            '{{VOLT_UNCERTAINTY_4}}': cal.volt_uncertainty_4 or '-',
            '{{VOLT_TOLERANCE_LIMIT_4}}': cal.volt_tolerance_limit_4 or '-',
            
            '{{VOLT_UUC_RANGE_5}}': cal.volt_uuc_range_5 or '-',
            '{{VOLT_UUC_SETTING_5}}': cal.volt_uuc_setting_5 or '-',
            '{{VOLT_MEASURED_VALUE_5}}': cal.volt_measured_value_5 or '-',
            '{{VOLT_UNCERTAINTY_5}}': cal.volt_uncertainty_5 or '-',
            '{{VOLT_TOLERANCE_LIMIT_5}}': cal.volt_tolerance_limit_5 or '-',
        }
        
        # ข้อมูลมาตรฐาน
        if std:
            replacements.update({
                '{{STANDARD_ASSET_NO}}': getattr(std, 'asset_number', None) or getattr(std, 'serial_number', '-'),
                '{{STANDARD_DESCRIPTION}}': std.name or '-',
                '{{STANDARD_MAKER_MODEL}}': getattr(std, 'maker_model', '-'),
                '{{STANDARD_SERIAL}}': std.serial_number or '-',
                '{{STANDARD_CERTIFICATE}}': getattr(std, 'certificate_number', '-'),
                '{{STANDARD_DUE_DATE}}': getattr(std, 'due_date', '-'),
            })
        else:
            replacements.update({
                '{{STANDARD_ASSET_NO}}': '-',
                '{{STANDARD_DESCRIPTION}}': '-',
                '{{STANDARD_MAKER_MODEL}}': '-',
                '{{STANDARD_SERIAL}}': '-',
                '{{STANDARD_CERTIFICATE}}': '-',
                '{{STANDARD_DUE_DATE}}': '-',
            })
        
        # ข้อมูลมาตรฐานชุดที่ 2 (ถ้ามี)
        try:
            std2 = cal.calibration_equipment_used.filter(equipment_type='standard').exclude(equipment=std).first()
            if std2:
                eq2 = std2.equipment
                replacements.update({
                    '{{STANDARD_ASSET_NO_2}}': getattr(eq2, 'asset_number', None) or getattr(eq2, 'serial_number', '-'),
                    '{{STANDARD_DESCRIPTION_2}}': eq2.name or '-',
                    '{{STANDARD_MAKER_MODEL_2}}': getattr(eq2, 'maker_model', '-'),
                    '{{STANDARD_SERIAL_2}}': eq2.serial_number or '-',
                    '{{STANDARD_CERTIFICATE_2}}': getattr(eq2, 'certificate_number', '-'),
                    '{{STANDARD_DUE_DATE_2}}': getattr(eq2, 'due_date', '-'),
                })
            else:
                replacements.update({
                    '{{STANDARD_ASSET_NO_2}}': '-',
                    '{{STANDARD_DESCRIPTION_2}}': '-',
                    '{{STANDARD_MAKER_MODEL_2}}': '-',
                    '{{STANDARD_SERIAL_2}}': '-',
                    '{{STANDARD_CERTIFICATE_2}}': '-',
                    '{{STANDARD_DUE_DATE_2}}': '-',
                })
        except Exception as e:
            print(f"DEBUG: Error getting second standard: {e}")
            replacements.update({
                '{{STANDARD_ASSET_NO_2}}': '-',
                '{{STANDARD_DESCRIPTION_2}}': '-',
                '{{STANDARD_MAKER_MODEL_2}}': '-',
                '{{STANDARD_SERIAL_2}}': '-',
                '{{STANDARD_CERTIFICATE_2}}': '-',
                '{{STANDARD_DUE_DATE_2}}': '-',
            })
        
        # สร้างตารางเครื่องมือที่ใช้
        equipment_table = ""
        try:
            equipment_used = cal.calibration_equipment_used.all()
            for i, eq_used in enumerate(equipment_used, 1):
                try:
                    eq = eq_used.equipment
                    asset_no = getattr(eq, 'asset_number', None) or getattr(eq, 'serial_number', '-')
                    name = eq.name or '-'
                    maker_model = getattr(eq, 'maker_model', '-')
                    serial = eq.serial_number or '-'
                    certificate = getattr(eq, 'certificate_number', '-')
                    due_date = getattr(eq, 'due_date', '-')
                    
                    equipment_table += f"{asset_no}\t{name}\t{maker_model}\t{serial}\t{certificate}\t{due_date}\n"
                except Exception as e:
                    print(f"DEBUG: Error getting equipment {i}: {e}")
                    equipment_table += "-\t-\t-\t-\t-\t-\n"
        except Exception as e:
            print(f"DEBUG: Error getting equipment list: {e}")
            equipment_table = "-\t-\t-\t-\t-\t-\n"
        
        replacements['{{EQUIPMENT_TABLE}}'] = equipment_table
        
        print(f"DEBUG: Total replacements: {len(replacements)}")
        
        # Debug: แสดงข้อมูลตารางที่ส่งไป
        print(f"DEBUG: FREQ_UUC_RANGE: {getattr(cal, 'freq_uuc_range', None)}")
        print(f"DEBUG: FREQ_UUC_SETTING: {getattr(cal, 'freq_uuc_setting', None)}")
        print(f"DEBUG: FREQ_MEASURED_VALUE: {getattr(cal, 'freq_measured_value', None)}")
        print(f"DEBUG: FREQ_UUC_RANGE_2: {getattr(cal, 'freq_uuc_range_2', None)}")
        print(f"DEBUG: FREQ_UUC_SETTING_2: {getattr(cal, 'freq_uuc_setting_2', None)}")
        print(f"DEBUG: FREQ_MEASURED_VALUE_2: {getattr(cal, 'freq_measured_value_2', None)}")
        
        # Debug: แสดงข้อมูล APPROVER และ CALIBRATOR
        print(f"DEBUG: APPROVER: {replacements.get('{{APPROVER}}', 'NOT_FOUND')}")
        print(f"DEBUG: CALIBRATOR: {replacements.get('{{CALIBRATOR}}', 'NOT_FOUND')}")
        print(f"DEBUG: Certificate issuer: {cal.certificate_issuer}")
        print(f"DEBUG: Calibrator: {cal.calibrator}")
        
        # Debug: แสดงข้อมูลมาตรฐาน
        print(f"DEBUG: Standard 1 - Asset: {getattr(std, 'asset_number', None) if std else None}")
        print(f"DEBUG: Standard 1 - Name: {std.name if std else None}")
        print(f"DEBUG: Standard 1 - Serial: {std.serial_number if std else None}")
        print(f"DEBUG: Standard 2 - Asset: {getattr(std, 'asset_number', None) if std else None}")
        print(f"DEBUG: Standard 2 - Name: {std.name if std else None}")
        print(f"DEBUG: Standard 2 - Serial: {std.serial_number if std else None}")
        
        # Debug: แสดงข้อมูลดิบ
        print(f"DEBUG: Raw data - freq_uuc_range: {cal.freq_uuc_range}")
        print(f"DEBUG: Raw data - freq_uuc_setting: {cal.freq_uuc_setting}")
        print(f"DEBUG: Raw data - freq_measured_value: {cal.freq_measured_value}")
        print(f"DEBUG: Raw data - freq_uuc_range_2: {cal.freq_uuc_range_2}")
        print(f"DEBUG: Raw data - freq_uuc_setting_2: {cal.freq_uuc_setting_2}")
        print(f"DEBUG: Raw data - freq_measured_value_2: {cal.freq_measured_value_2}")
        
        # แทนค่าในเอกสาร
        replace_text_in_document(doc, replacements)
        print("DEBUG: Document replacement completed")
        
        # สร้าง response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="High_Frequency_Certificate_{cal_id}.docx"'
        
        # บันทึกเอกสาร
        doc.save(response)
        print("DEBUG: Creating response for High Frequency certificate")
        print("DEBUG: Response created successfully")
        
        return response
        
    except Exception as e:
        print(f"DEBUG: Error in export_high_frequency_certificate_docx: {e}")
        return HttpResponse(f"Error: {e}", status=500)

def export_pressure_certificate_docx(request, cal_id):
    """Export Pressure certificate as DOCX"""
    try:
        print(f"DEBUG: Attempting to export Pressure certificate for ID: {cal_id}")
        
        # ดึงข้อมูลการสอบเทียบ
        cal = get_object_or_404(CalibrationPressure, cal_pressure_id=cal_id)
        print(f"DEBUG: Found calibration - Status: {cal.status}, Machine: {cal.uuc_id.name}")
        
        # ดึงข้อมูลเครื่องมือมาตรฐาน
        std = cal.std_id
        if std:
            print(f"DEBUG: Standard equipment found: {std.name}")
        else:
            print("DEBUG: No standard equipment found")
        
        # สร้าง path ไปยัง template
        template_path = os.path.join(settings.BASE_DIR, 'cert_templates', 'pressure_template.docx')
        print(f"DEBUG: Template path: {template_path}")
        
        # เปิด template
        doc = Document(template_path)
        
        # สร้าง replacements dictionary
        replacements = {
            # ข้อมูลเครื่องมือ
            '{{MODEL}}': getattr(cal.uuc_id, 'model', '-'),
            '{{MANUFACTURER}}': getattr(cal.uuc_id, 'manufacturer', '-'),
            '{{DESCRIPTION}}': getattr(cal.uuc_id, 'description', '-'),
            '{{SERIAL_NUMBER}}': getattr(cal.uuc_id, 'serial_number', '-'),
            '{{ASSET_NUMBER}}': getattr(cal.uuc_id, 'asset_number', '-'),
            '{{RANGE}}': getattr(cal, 'measurement_range', '-'),
            
            # ข้อมูลการสอบเทียบ
            '{{CALIBRATION_DATE}}': cal.update.strftime('%d/%m/%Y') if cal.update else '-',
            '{{DATE_OF_CALIBRATION}}': cal.update.strftime('%d/%m/%Y') if cal.update else '-',
            '{{CERTIFICATE_NUMBER}}': cal.certificate_number or '-',
            '{{DUE_DATE}}': cal.next_due.strftime('%d/%m/%Y') if cal.next_due else '-',
            '{{STATUS}}': cal.get_status_display(),
            
            # ข้อมูลผู้รับผิดชอบ
            '{{CALIBRATOR}}': cal.calibrator.username if cal.calibrator else '-',
            '{{APPROVER}}': cal.certificate_issuer.username if cal.certificate_issuer else '-',
            
            # ข้อมูลหน่วยงาน
            '{{CUSTOMER}}': getattr(cal.uuc_id.organize, 'name', '-') if cal.uuc_id.organize else '-',
            '{{CUSTOMER_ADDRESS}}': getattr(cal.uuc_id.organize, 'address', '-') if cal.uuc_id.organize else '-',
            '{{LOCATION_NAME}}': getattr(cal.uuc_id.organize, 'name', '-') if cal.uuc_id.organize else '-',
            '{{LOCATION_ADDRESS}}': getattr(cal.uuc_id.organize, 'address', '-') if cal.uuc_id.organize else '-',
            
            # ข้อมูลตารางผลการสอบเทียบ Pressure (6 แถว)
            # แถวที่ 1
            '{{UUC_SET}}': cal.set or '-',
            '{{ACTUAL}}': cal.actual or '-',
            '{{ERROR}}': cal.error or '-',
            '{{TOLERANCE_LIMIT}}': f"{cal.tolerance_start or '-'} - {cal.tolerance_end or '-'}" if cal.tolerance_start and cal.tolerance_end else '-',
            
            # แถวที่ 2
            '{{UUC_SET_2}}': cal.set_2 or '-',
            '{{ACTUAL_2}}': cal.actual_2 or '-',
            '{{ERROR_2}}': cal.error_2 or '-',
            '{{TOLERANCE_LIMIT_2}}': f"{cal.tolerance_start_2 or '-'} - {cal.tolerance_end_2 or '-'}" if cal.tolerance_start_2 and cal.tolerance_end_2 else '-',
            
            # แถวที่ 3
            '{{UUC_SET_3}}': cal.set_3 or '-',
            '{{ACTUAL_3}}': cal.actual_3 or '-',
            '{{ERROR_3}}': cal.error_3 or '-',
            '{{TOLERANCE_LIMIT_3}}': f"{cal.tolerance_start_3 or '-'} - {cal.tolerance_end_3 or '-'}" if cal.tolerance_start_3 and cal.tolerance_end_3 else '-',
            
            # แถวที่ 4
            '{{UUC_SET_4}}': cal.set_4 or '-',
            '{{ACTUAL_4}}': cal.actual_4 or '-',
            '{{ERROR_4}}': cal.error_4 or '-',
            '{{TOLERANCE_LIMIT_4}}': f"{cal.tolerance_start_4 or '-'} - {cal.tolerance_end_4 or '-'}" if cal.tolerance_start_4 and cal.tolerance_end_4 else '-',
            
            # แถวที่ 5
            '{{UUC_SET_5}}': cal.set_5 or '-',
            '{{ACTUAL_5}}': cal.actual_5 or '-',
            '{{ERROR_5}}': cal.error_5 or '-',
            '{{TOLERANCE_LIMIT_5}}': f"{cal.tolerance_start_5 or '-'} - {cal.tolerance_end_5 or '-'}" if cal.tolerance_start_5 and cal.tolerance_end_5 else '-',
            
            # แถวที่ 6
            '{{UUC_SET_6}}': cal.set_6 or '-',
            '{{ACTUAL_6}}': cal.actual_6 or '-',
            '{{ERROR_6}}': cal.error_6 or '-',
            '{{TOLERANCE_LIMIT_6}}': f"{cal.tolerance_start_6 or '-'} - {cal.tolerance_end_6 or '-'}" if cal.tolerance_start_6 and cal.tolerance_end_6 else '-',
        }
        
        # ข้อมูลมาตรฐาน
        if std:
            replacements.update({
                '{{STANDARD_ASSET_NO}}': getattr(std, 'asset_number', None) or getattr(std, 'serial_number', '-'),
                '{{STANDARD_DESCRIPTION}}': std.name or '-',
                '{{STANDARD_MAKER_MODEL}}': getattr(std, 'maker_model', '-'),
                '{{STANDARD_SERIAL}}': std.serial_number or '-',
                '{{STANDARD_CERTIFICATE}}': getattr(std, 'certificate_number', '-'),
                '{{STANDARD_DUE_DATE}}': getattr(std, 'due_date', '-'),
            })
        else:
            replacements.update({
                '{{STANDARD_ASSET_NO}}': '-',
                '{{STANDARD_DESCRIPTION}}': '-',
                '{{STANDARD_MAKER_MODEL}}': '-',
                '{{STANDARD_SERIAL}}': '-',
                '{{STANDARD_CERTIFICATE}}': '-',
                '{{STANDARD_DUE_DATE}}': '-',
            })
        
        # ข้อมูลมาตรฐานชุดที่ 2 (ถ้ามี)
        try:
            std2 = cal.calibration_equipment_used.filter(equipment_type='standard').exclude(equipment=std).first()
            if std2:
                eq2 = std2.equipment
                replacements.update({
                    '{{STANDARD_ASSET_NO_2}}': getattr(eq2, 'asset_number', None) or getattr(eq2, 'serial_number', '-'),
                    '{{STANDARD_DESCRIPTION_2}}': eq2.name or '-',
                    '{{STANDARD_MAKER_MODEL_2}}': getattr(eq2, 'maker_model', '-'),
                    '{{STANDARD_SERIAL_2}}': eq2.serial_number or '-',
                    '{{STANDARD_CERTIFICATE_2}}': getattr(eq2, 'certificate_number', '-'),
                    '{{STANDARD_DUE_DATE_2}}': getattr(eq2, 'due_date', '-'),
                })
            else:
                replacements.update({
                    '{{STANDARD_ASSET_NO_2}}': '-',
                    '{{STANDARD_DESCRIPTION_2}}': '-',
                    '{{STANDARD_MAKER_MODEL_2}}': '-',
                    '{{STANDARD_SERIAL_2}}': '-',
                    '{{STANDARD_CERTIFICATE_2}}': '-',
                    '{{STANDARD_DUE_DATE_2}}': '-',
                })
        except Exception as e:
            print(f"DEBUG: Error getting second standard: {e}")
            replacements.update({
                '{{STANDARD_ASSET_NO_2}}': '-',
                '{{STANDARD_DESCRIPTION_2}}': '-',
                '{{STANDARD_MAKER_MODEL_2}}': '-',
                '{{STANDARD_SERIAL_2}}': '-',
                '{{STANDARD_CERTIFICATE_2}}': '-',
                '{{STANDARD_DUE_DATE_2}}': '-',
            })
        
        # สร้างตารางเครื่องมือที่ใช้
        equipment_table = ""
        try:
            equipment_used = cal.calibration_equipment_used.all()
            for i, eq_used in enumerate(equipment_used, 1):
                try:
                    eq = eq_used.equipment
                    asset_no = getattr(eq, 'asset_number', None) or getattr(eq, 'serial_number', '-')
                    name = eq.name or '-'
                    maker_model = getattr(eq, 'maker_model', '-')
                    serial = eq.serial_number or '-'
                    certificate = getattr(eq, 'certificate_number', '-')
                    due_date = getattr(eq, 'due_date', '-')
                    
                    equipment_table += f"{asset_no}\t{name}\t{maker_model}\t{serial}\t{certificate}\t{due_date}\n"
                except Exception as e:
                    print(f"DEBUG: Error getting equipment {i}: {e}")
                    equipment_table += "-\t-\t-\t-\t-\t-\n"
        except Exception as e:
            print(f"DEBUG: Error getting equipment list: {e}")
            equipment_table = "-\t-\t-\t-\t-\t-\n"
        
        replacements['{{EQUIPMENT_TABLE}}'] = equipment_table
        
        print(f"DEBUG: Total replacements: {len(replacements)}")
        
        # Debug: แสดงข้อมูลตารางที่ส่งไป
        print(f"DEBUG: SET: {getattr(cal, 'set', None)}")
        print(f"DEBUG: M1: {getattr(cal, 'm1', None)}")
        print(f"DEBUG: M2: {getattr(cal, 'm2', None)}")
        print(f"DEBUG: M3: {getattr(cal, 'm3', None)}")
        print(f"DEBUG: M4: {getattr(cal, 'm4', None)}")
        print(f"DEBUG: AVG: {getattr(cal, 'avg', None)}")
        print(f"DEBUG: ACTUAL: {getattr(cal, 'actual', None)}")
        print(f"DEBUG: ERROR: {getattr(cal, 'error', None)}")
        
        # Debug: แสดงข้อมูล APPROVER และ CALIBRATOR
        print(f"DEBUG: APPROVER: {replacements.get('{{APPROVER}}', 'NOT_FOUND')}")
        print(f"DEBUG: CALIBRATOR: {replacements.get('{{CALIBRATOR}}', 'NOT_FOUND')}")
        print(f"DEBUG: Certificate issuer: {cal.certificate_issuer}")
        print(f"DEBUG: Calibrator: {cal.calibrator}")
        
        # แทนค่าในเอกสาร
        replace_text_in_document(doc, replacements)
        print("DEBUG: Document replacement completed")
        
        # สร้าง response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="Pressure_Certificate_{cal_id}.docx"'
        
        # บันทึกเอกสาร
        doc.save(response)
        print("DEBUG: Creating response for Pressure certificate")
        print("DEBUG: Response created successfully")
        
        return response
        
    except Exception as e:
        print(f"DEBUG: Error in export_pressure_certificate_docx: {e}")
        return HttpResponse(f"Error: {e}", status=500)

def export_torque_certificate_docx(request, cal_id):
    """Export Torque certificate as DOCX"""
    try:
        print(f"DEBUG: Attempting to export Torque certificate for ID: {cal_id}")
        
        # ดึงข้อมูลการสอบเทียบ
        cal = get_object_or_404(CalibrationTorque, cal_torque_id=cal_id)
        print(f"DEBUG: Found calibration - Status: {cal.status}, Machine: {cal.uuc_id.name}")
        
        # ดึงข้อมูลเครื่องมือมาตรฐาน
        std = cal.std_id
        if std:
            print(f"DEBUG: Standard equipment found: {std.name}")
        else:
            print("DEBUG: No standard equipment found")
        
        # สร้าง path ไปยัง template
        template_path = os.path.join(settings.BASE_DIR, 'cert_templates', 'Torque_template.docx')
        print(f"DEBUG: Template path: {template_path}")
        
        # เปิด template
        doc = Document(template_path)
        
        # สร้าง replacements dictionary
        replacements = {
            # ข้อมูลเครื่องมือ
            '{{MODEL}}': getattr(cal.uuc_id, 'model', '-'),
            '{{MANUFACTURER}}': getattr(cal.uuc_id, 'manufacturer', '-'),
            '{{DESCRIPTION}}': getattr(cal.uuc_id, 'description', '-'),
            '{{SERIAL_NUMBER}}': getattr(cal.uuc_id, 'serial_number', '-'),
            '{{ASSET_NUMBER}}': getattr(cal.uuc_id, 'asset_number', '-'),
            '{{RANGE}}': getattr(cal, 'measurement_range', '-'),
            
            # ข้อมูลการสอบเทียบ
            '{{CALIBRATION_DATE}}': cal.update.strftime('%d/%m/%Y') if cal.update else '-',
            '{{DATE_OF_CALIBRATION}}': cal.update.strftime('%d/%m/%Y') if cal.update else '-',
            '{{CERTIFICATE_NUMBER}}': cal.certificate_number or '-',
            '{{DUE_DATE}}': cal.next_due.strftime('%d/%m/%Y') if cal.next_due else '-',
            '{{STATUS}}': cal.get_status_display(),
            
            # ข้อมูลผู้รับผิดชอบ
            '{{CALIBRATOR}}': cal.calibrator.username if cal.calibrator else '-',
            '{{APPROVER}}': cal.certificate_issuer.username if cal.certificate_issuer else '-',
            
            # ข้อมูลหน่วยงาน
            '{{CUSTOMER}}': getattr(cal.uuc_id.organize, 'name', '-') if cal.uuc_id.organize else '-',
            '{{CUSTOMER_ADDRESS}}': getattr(cal.uuc_id.organize, 'address', '-') if cal.uuc_id.organize else '-',
            '{{LOCATION_NAME}}': getattr(cal.uuc_id.organize, 'name', '-') if cal.uuc_id.organize else '-',
            '{{LOCATION_ADDRESS}}': getattr(cal.uuc_id.organize, 'address', '-') if cal.uuc_id.organize else '-',
            
            # ข้อมูลตารางผลการสอบเทียบ Torque CW (3 แถว)
            # แถวที่ 1
            '{{CW_SET}}': str(cal.cwset) if cal.cwset is not None else '-',
            '{{CW_0}}': str(cal.cw0) if cal.cw0 is not None else '-',
            '{{CW_90}}': str(cal.cw90) if cal.cw90 is not None else '-',
            '{{CW_180}}': str(cal.cw180) if cal.cw180 is not None else '-',
            '{{CW_270}}': str(cal.cw270) if cal.cw270 is not None else '-',
            '{{CW_AVG}}': str(cal.cw_avg) if cal.cw_avg is not None else '-',
            '{{CW_ACTUAL}}': str(cal.cw_actual) if cal.cw_actual is not None else '-',
            '{{CW_ERROR}}': str(cal.cw_error) if cal.cw_error is not None else '0.0',
            '{{CW_UNCERTAINTY}}': str(cal.cw_uncen) if cal.cw_uncen is not None else '-',
            '{{CW_TOLERANCE_LIMIT}}': f"{cal.cw_tolerance_start} - {cal.cw_tolerance_end}" if cal.cw_tolerance_start is not None and cal.cw_tolerance_end is not None else '-',
            
            # แถวที่ 2
            '{{CW_SET_2}}': str(cal.cwset_2) if cal.cwset_2 is not None else '-',
            '{{CW_ACTUAL_2}}': str(cal.cw_actual_2) if cal.cw_actual_2 is not None else '-',
            '{{CW_ERROR_2}}': str(cal.cw_error_2) if cal.cw_error_2 is not None else '0.0',
            '{{CW_TOLERANCE_LIMIT_2}}': f"{cal.cw_tolerance_start_2} - {cal.cw_tolerance_end_2}" if cal.cw_tolerance_start_2 is not None and cal.cw_tolerance_end_2 is not None else '-',
            
            # แถวที่ 3
            '{{CW_SET_3}}': str(cal.cwset_3) if cal.cwset_3 is not None else '-',
            '{{CW_ACTUAL_3}}': str(cal.cw_actual_3) if cal.cw_actual_3 is not None else '-',
            '{{CW_ERROR_3}}': str(cal.cw_error_3) if cal.cw_error_3 is not None else '0.0',
            '{{CW_TOLERANCE_LIMIT_3}}': f"{cal.cw_tolerance_start_3} - {cal.cw_tolerance_end_3}" if cal.cw_tolerance_start_3 is not None and cal.cw_tolerance_end_3 is not None else '-',
            
            # ข้อมูลตารางผลการสอบเทียบ Torque CCW (3 แถว)
            # แถวที่ 1
            '{{CCW_SET}}': str(cal.ccwset) if cal.ccwset is not None else '-',
            '{{CCW_0}}': str(cal.ccw0) if cal.ccw0 is not None else '-',
            '{{CCW_90}}': str(cal.ccw90) if cal.ccw90 is not None else '-',
            '{{CCW_180}}': str(cal.ccw180) if cal.ccw180 is not None else '-',
            '{{CCW_270}}': str(cal.ccw270) if cal.ccw270 is not None else '-',
            '{{CCW_AVG}}': str(cal.ccw_avg) if cal.ccw_avg is not None else '-',
            '{{CCW_ACTUAL}}': str(cal.ccw_actual) if cal.ccw_actual is not None else '-',
            '{{CCW_ERROR}}': str(cal.ccw_error) if cal.ccw_error is not None else '0.0',
            '{{CCW_UNCERTAINTY}}': str(cal.ccw_uncen) if cal.ccw_uncen is not None else '-',
            '{{CCW_TOLERANCE_LIMIT}}': f"{cal.ccw_tolerance_start} - {cal.ccw_tolerance_end}" if cal.ccw_tolerance_start is not None and cal.ccw_tolerance_end is not None else '-',
            
            # แถวที่ 2
            '{{CCW_SET_2}}': str(cal.ccwset_2) if cal.ccwset_2 is not None else '-',
            '{{CCW_ACTUAL_2}}': str(cal.ccw_actual_2) if cal.ccw_actual_2 is not None else '-',
            '{{CCW_ERROR_2}}': str(cal.ccw_error_2) if cal.ccw_error_2 is not None else '0.0',
            '{{CCW_TOLERANCE_LIMIT_2}}': f"{cal.ccw_tolerance_start_2} - {cal.ccw_tolerance_end_2}" if cal.ccw_tolerance_start_2 is not None and cal.ccw_tolerance_end_2 is not None else '-',
            
            # แถวที่ 3
            '{{CCW_SET_3}}': str(cal.ccwset_3) if cal.ccwset_3 is not None else '-',
            '{{CCW_ACTUAL_3}}': str(cal.ccw_actual_3) if cal.ccw_actual_3 is not None else '-',
            '{{CCW_ERROR_3}}': str(cal.ccw_error_3) if cal.ccw_error_3 is not None else '0.0',
            '{{CCW_TOLERANCE_LIMIT_3}}': f"{cal.ccw_tolerance_start_3} - {cal.ccw_tolerance_end_3}" if cal.ccw_tolerance_start_3 is not None and cal.ccw_tolerance_end_3 is not None else '-',
        }
        
        # ข้อมูลมาตรฐาน
        if std:
            replacements.update({
                '{{STANDARD_ASSET_NO}}': getattr(std, 'asset_number', None) or getattr(std, 'serial_number', '-'),
                '{{STANDARD_DESCRIPTION}}': std.name or '-',
                '{{STANDARD_MAKER_MODEL}}': getattr(std, 'maker_model', '-'),
                '{{STANDARD_SERIAL}}': std.serial_number or '-',
                '{{STANDARD_CERTIFICATE}}': getattr(std, 'certificate_number', '-'),
                '{{STANDARD_DUE_DATE}}': getattr(std, 'due_date', '-'),
            })
        else:
            replacements.update({
                '{{STANDARD_ASSET_NO}}': '-',
                '{{STANDARD_DESCRIPTION}}': '-',
                '{{STANDARD_MAKER_MODEL}}': '-',
                '{{STANDARD_SERIAL}}': '-',
                '{{STANDARD_CERTIFICATE}}': '-',
                '{{STANDARD_DUE_DATE}}': '-',
            })
        
        # ข้อมูลมาตรฐานชุดที่ 2 (ถ้ามี)
        try:
            std2 = cal.calibration_equipment_used.filter(equipment_type='standard').exclude(equipment=std).first()
            if std2:
                eq2 = std2.equipment
                replacements.update({
                    '{{STANDARD_ASSET_NO_2}}': getattr(eq2, 'asset_number', None) or getattr(eq2, 'serial_number', '-'),
                    '{{STANDARD_DESCRIPTION_2}}': eq2.name or '-',
                    '{{STANDARD_MAKER_MODEL_2}}': getattr(eq2, 'maker_model', '-'),
                    '{{STANDARD_SERIAL_2}}': eq2.serial_number or '-',
                    '{{STANDARD_CERTIFICATE_2}}': getattr(eq2, 'certificate_number', '-'),
                    '{{STANDARD_DUE_DATE_2}}': getattr(eq2, 'due_date', '-'),
                })
            else:
                replacements.update({
                    '{{STANDARD_ASSET_NO_2}}': '-',
                    '{{STANDARD_DESCRIPTION_2}}': '-',
                    '{{STANDARD_MAKER_MODEL_2}}': '-',
                    '{{STANDARD_SERIAL_2}}': '-',
                    '{{STANDARD_CERTIFICATE_2}}': '-',
                    '{{STANDARD_DUE_DATE_2}}': '-',
                })
        except Exception as e:
            print(f"DEBUG: Error getting second standard: {e}")
            replacements.update({
                '{{STANDARD_ASSET_NO_2}}': '-',
                '{{STANDARD_DESCRIPTION_2}}': '-',
                '{{STANDARD_MAKER_MODEL_2}}': '-',
                '{{STANDARD_SERIAL_2}}': '-',
                '{{STANDARD_CERTIFICATE_2}}': '-',
                '{{STANDARD_DUE_DATE_2}}': '-',
            })
        
        # สร้างตารางเครื่องมือที่ใช้
        equipment_table = ""
        try:
            equipment_used = cal.calibration_equipment_used.all()
            for i, eq_used in enumerate(equipment_used, 1):
                try:
                    eq = eq_used.equipment
                    asset_no = getattr(eq, 'asset_number', None) or getattr(eq, 'serial_number', '-')
                    name = eq.name or '-'
                    maker_model = getattr(eq, 'maker_model', '-')
                    serial = eq.serial_number or '-'
                    certificate = getattr(eq, 'certificate_number', '-')
                    due_date = getattr(eq, 'due_date', '-')
                    
                    equipment_table += f"{asset_no}\t{name}\t{maker_model}\t{serial}\t{certificate}\t{due_date}\n"
                except Exception as e:
                    print(f"DEBUG: Error getting equipment {i}: {e}")
                    equipment_table += "-\t-\t-\t-\t-\t-\n"
        except Exception as e:
            print(f"DEBUG: Error getting equipment list: {e}")
            equipment_table = "-\t-\t-\t-\t-\t-\n"
        
        replacements['{{EQUIPMENT_TABLE}}'] = equipment_table
        
        print(f"DEBUG: Total replacements: {len(replacements)}")
        
        # Debug: แสดงข้อมูลตารางที่ส่งไป
        print(f"DEBUG: CW_SET: {getattr(cal, 'cwset', None)}")
        print(f"DEBUG: CW_0: {getattr(cal, 'cw0', None)}")
        print(f"DEBUG: CW_90: {getattr(cal, 'cw90', None)}")
        print(f"DEBUG: CW_180: {getattr(cal, 'cw180', None)}")
        print(f"DEBUG: CW_270: {getattr(cal, 'cw270', None)}")
        print(f"DEBUG: CW_AVG: {getattr(cal, 'cw_avg', None)}")
        print(f"DEBUG: CW_ACTUAL: {getattr(cal, 'cw_actual', None)}")
        print(f"DEBUG: CW_ERROR: {getattr(cal, 'cw_error', None)}")
        print(f"DEBUG: CW_TOLERANCE_START: {getattr(cal, 'cw_tolerance_start', None)}")
        print(f"DEBUG: CW_TOLERANCE_END: {getattr(cal, 'cw_tolerance_end', None)}")
        
        print(f"DEBUG: CCW_SET: {getattr(cal, 'ccwset', None)}")
        print(f"DEBUG: CCW_0: {getattr(cal, 'ccw0', None)}")
        print(f"DEBUG: CCW_90: {getattr(cal, 'ccw90', None)}")
        print(f"DEBUG: CCW_180: {getattr(cal, 'ccw180', None)}")
        print(f"DEBUG: CCW_270: {getattr(cal, 'ccw270', None)}")
        print(f"DEBUG: CCW_AVG: {getattr(cal, 'ccw_avg', None)}")
        print(f"DEBUG: CCW_ACTUAL: {getattr(cal, 'ccw_actual', None)}")
        print(f"DEBUG: CCW_ERROR: {getattr(cal, 'ccw_error', None)}")
        print(f"DEBUG: CCW_TOLERANCE_START: {getattr(cal, 'ccw_tolerance_start', None)}")
        print(f"DEBUG: CCW_TOLERANCE_END: {getattr(cal, 'ccw_tolerance_end', None)}")
        
        # Debug: แสดงค่าที่ส่งไปใน replacements
        print(f"DEBUG: CW_SET replacement: {replacements.get('{{CW_SET}}', 'NOT_FOUND')}")
        print(f"DEBUG: CW_ACTUAL replacement: {replacements.get('{{CW_ACTUAL}}', 'NOT_FOUND')}")
        print(f"DEBUG: CW_ERROR replacement: {replacements.get('{{CW_ERROR}}', 'NOT_FOUND')}")
        print(f"DEBUG: CW_TOLERANCE_LIMIT replacement: {replacements.get('{{CW_TOLERANCE_LIMIT}}', 'NOT_FOUND')}")
        print(f"DEBUG: CCW_SET replacement: {replacements.get('{{CCW_SET}}', 'NOT_FOUND')}")
        print(f"DEBUG: CCW_ACTUAL replacement: {replacements.get('{{CCW_ACTUAL}}', 'NOT_FOUND')}")
        print(f"DEBUG: CCW_ERROR replacement: {replacements.get('{{CCW_ERROR}}', 'NOT_FOUND')}")
        print(f"DEBUG: CCW_TOLERANCE_LIMIT replacement: {replacements.get('{{CCW_TOLERANCE_LIMIT}}', 'NOT_FOUND')}")
        
        # Debug: แสดงข้อมูลแถวที่ 2 และ 3
        print(f"DEBUG: CW_SET_2: {getattr(cal, 'cwset_2', None)}")
        print(f"DEBUG: CW_ACTUAL_2: {getattr(cal, 'cw_actual_2', None)}")
        print(f"DEBUG: CW_ERROR_2: {getattr(cal, 'cw_error_2', None)}")
        print(f"DEBUG: CW_SET_3: {getattr(cal, 'cwset_3', None)}")
        print(f"DEBUG: CW_ACTUAL_3: {getattr(cal, 'cw_actual_3', None)}")
        print(f"DEBUG: CW_ERROR_3: {getattr(cal, 'cw_error_3', None)}")
        
        print(f"DEBUG: CCW_SET_2: {getattr(cal, 'ccwset_2', None)}")
        print(f"DEBUG: CCW_ACTUAL_2: {getattr(cal, 'ccw_actual_2', None)}")
        print(f"DEBUG: CCW_ERROR_2: {getattr(cal, 'ccw_error_2', None)}")
        print(f"DEBUG: CCW_SET_3: {getattr(cal, 'ccwset_3', None)}")
        print(f"DEBUG: CCW_ACTUAL_3: {getattr(cal, 'ccw_actual_3', None)}")
        print(f"DEBUG: CCW_ERROR_3: {getattr(cal, 'ccw_error_3', None)}")
        
        # Debug: แสดงค่าที่ส่งไปใน replacements สำหรับแถวที่ 2 และ 3
        print(f"DEBUG: CCW_ACTUAL_2 replacement: {replacements.get('{{CCW_ACTUAL_2}}', 'NOT_FOUND')}")
        print(f"DEBUG: CCW_ERROR_2 replacement: {replacements.get('{{CCW_ERROR_2}}', 'NOT_FOUND')}")
        print(f"DEBUG: CCW_ACTUAL_3 replacement: {replacements.get('{{CCW_ACTUAL_3}}', 'NOT_FOUND')}")
        print(f"DEBUG: CCW_ERROR_3 replacement: {replacements.get('{{CCW_ERROR_3}}', 'NOT_FOUND')}")
        
        # Debug: แสดงข้อมูล APPROVER และ CALIBRATOR
        print(f"DEBUG: APPROVER: {replacements.get('{{APPROVER}}', 'NOT_FOUND')}")
        print(f"DEBUG: CALIBRATOR: {replacements.get('{{CALIBRATOR}}', 'NOT_FOUND')}")
        print(f"DEBUG: Certificate issuer: {cal.certificate_issuer}")
        print(f"DEBUG: Calibrator: {cal.calibrator}")
        
        # แทนค่าในเอกสาร
        replace_text_in_document(doc, replacements)
        print("DEBUG: Document replacement completed")
        
        # สร้าง response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="Torque_Certificate_{cal_id}.docx"'
        
        # บันทึกเอกสาร
        doc.save(response)
        print("DEBUG: Creating response for Torque certificate")
        print("DEBUG: Response created successfully")
        
        return response
        
    except Exception as e:
        print(f"DEBUG: Error in export_torque_certificate_docx: {e}")
        return HttpResponse(f"Error: {e}", status=500)


def export_microwave_certificate_docx(request, cal_id):
    """Export Microwave certificate as DOCX"""
    try:
        print(f"DEBUG: Attempting to export Microwave certificate for ID: {cal_id}")
        
        # ดึงข้อมูลการสอบเทียบ
        cal = get_object_or_404(MicrowaveCalibration, id=cal_id)
        print(f"DEBUG: Found calibration - Status: {cal.status}, Machine: {cal.machine.name}")
        
        # ข้อมูลเครื่องมือ
        machine = cal.machine
        print(f"DEBUG: Machine info - Name: {machine.name}, Model: {machine.model}, Serial: {machine.serial_number}")
        
        # ข้อมูลมาตรฐาน
        standard_equipment = cal.calibration_equipment_used.first()
        print(f"DEBUG: Standard equipment found: {standard_equipment.equipment.name if standard_equipment else 'None'}")
        
        # Template path
        template_path = os.path.join(settings.BASE_DIR, 'cert_templates', 'Microwave_template.docx')
        print(f"DEBUG: Template path: {template_path}")
        
        # เปิด template
        doc = Document(template_path)
        
        # สร้าง replacements dictionary
        replacements = {
            # ข้อมูลเครื่องมือ
            '{{MACHINE_NAME}}': machine.name or '-',
            '{{MACHINE_MODEL}}': machine.model or '-',
            '{{MACHINE_SERIAL}}': machine.serial_number or '-',
            '{{MACHINE_MANUFACTURER}}': '-',
            '{{MACHINE_RANGE}}': cal.readings or '-',
            
            # ข้อมูลการสอบเทียบ
            '{{CALIBRATION_DATE}}': cal.date_calibration.strftime('%d/%m/%Y') if cal.date_calibration else '-',
            '{{NEXT_DUE_DATE}}': cal.next_due.strftime('%d/%m/%Y') if cal.next_due else '-',
            '{{CERTIFICATE_NUMBER}}': cal.certificate_number or '-',
            '{{STATUS}}': cal.status or '-',
            
            # ข้อมูลมาตรฐาน
            '{{STANDARD_NAME}}': standard_equipment.equipment.name if standard_equipment else '-',
            '{{STANDARD_MODEL}}': standard_equipment.equipment.model if standard_equipment else '-',
            '{{STANDARD_SERIAL}}': standard_equipment.equipment.serial_number if standard_equipment else '-',
            '{{STANDARD_CERTIFICATE}}': '-',
            '{{STANDARD_DUE_DATE}}': '-',
            
            # ข้อมูลผู้สอบเทียบ
            '{{CALIBRATOR}}': cal.calibrator.get_full_name() if cal.calibrator else '-',
            '{{APPROVER}}': cal.certificate_issuer.get_full_name() if cal.certificate_issuer else '-',
            
            # เพิ่ม placeholder ที่ขาดหายไป
            '{{MODEL}}': machine.model or '-',
            '{{DESCRIPTION}}': machine.machine_type.name if machine.machine_type else '-',
            '{{DUE_DATE}}': cal.next_due.strftime('%d/%m/%Y') if cal.next_due else '-',
            '{{SERIAL_NUMBER}}': machine.serial_number or '-',
            '{{MANUFACTURER}}': '-',
            '{{RANGE}}': cal.dc_uuc_range or '-',
            
            # เพิ่ม placeholder ที่ขาดหายไป
            '{{CUSTOMER_ADDRESS}}': machine.organize.name if machine.organize else '-',
            '{{DATE_OF_CALIBRATION}}': cal.date_calibration.strftime('%d/%m/%Y') if cal.date_calibration else '-',
            
            # ข้อมูลมาตรฐานเพิ่มเติม
            '{{STANDARD_ASSET_NO}}': standard_equipment.equipment.asset_number if standard_equipment and hasattr(standard_equipment.equipment, 'asset_number') else '-',
            '{{STANDARD_DESCRIPTION}}': standard_equipment.equipment.name if standard_equipment else '-',
            '{{STANDARD_MAKER_MODEL}}': standard_equipment.equipment.model if standard_equipment else '-',
            '{{STANDARD_SERIAL}}': standard_equipment.equipment.serial_number if standard_equipment else '-',
            '{{STANDARD_ASSET_NO_2}}': '-',
            '{{STANDARD_DESCRIPTION_2}}': '-',
            '{{STANDARD_MAKER_MODEL_2}}': '-',
            '{{STANDARD_SERIAL_2}}': '-',
            
            # ข้อมูลตาราง DC Voltage
            '{{DC_UUC_RANGE}}': cal.dc_uuc_range or '-',
            '{{DC_UUC_SETTING}}': cal.dc_uuc_setting or '-',
            '{{DC_MEASURED_VALUE}}': cal.dc_measured_value or '-',
            '{{DC_UNCERTAINTY}}': cal.dc_uncertainty or '-',
            '{{DC_TOLERANCE_LIMIT}}': cal.dc_tolerance_limit or '-',
            
            # ข้อมูลตาราง DC Voltage แถวที่ 2
            '{{DC_UUC_RANGE_2}}': cal.dc_uuc_range_2 or '-',
            '{{DC_UUC_SETTING_2}}': cal.dc_uuc_setting_2 or '-',
            '{{DC_MEASURED_VALUE_2}}': cal.dc_measured_value_2 or '-',
            '{{DC_UNCERTAINTY_2}}': cal.dc_uncertainty_2 or '-',
            '{{DC_TOLERANCE_LIMIT_2}}': cal.dc_tolerance_limit_2 or '-',
            
            # ข้อมูลตาราง DC Voltage แถวที่ 3
            '{{DC_UUC_RANGE_3}}': cal.dc_uuc_range_3 or '-',
            '{{DC_UUC_SETTING_3}}': cal.dc_uuc_setting_3 or '-',
            '{{DC_MEASURED_VALUE_3}}': cal.dc_measured_value_3 or '-',
            '{{DC_UNCERTAINTY_3}}': cal.dc_uncertainty_3 or '-',
            '{{DC_TOLERANCE_LIMIT_3}}': cal.dc_tolerance_limit_3 or '-',
            
            # ข้อมูลตาราง DC Voltage แถวที่ 4
            '{{DC_UUC_RANGE_4}}': cal.dc_uuc_range_4 or '-',
            '{{DC_UUC_SETTING_4}}': cal.dc_uuc_setting_4 or '-',
            '{{DC_MEASURED_VALUE_4}}': cal.dc_measured_value_4 or '-',
            '{{DC_UNCERTAINTY_4}}': cal.dc_uncertainty_4 or '-',
            '{{DC_TOLERANCE_LIMIT_4}}': cal.dc_tolerance_limit_4 or '-',
            
            # ข้อมูลตาราง DC Voltage แถวที่ 5
            '{{DC_UUC_RANGE_5}}': cal.dc_uuc_range_5 or '-',
            '{{DC_UUC_SETTING_5}}': cal.dc_uuc_setting_5 or '-',
            '{{DC_MEASURED_VALUE_5}}': cal.dc_measured_value_5 or '-',
            '{{DC_UNCERTAINTY_5}}': cal.dc_uncertainty_5 or '-',
            '{{DC_TOLERANCE_LIMIT_5}}': cal.dc_tolerance_limit_5 or '-',
        }
        
        print(f"DEBUG: Total replacements: {len(replacements)}")
        
        # แทนค่าในเอกสาร
        replace_text_in_document(doc, replacements)
        print("DEBUG: Document replacement completed")
        
        # สร้าง response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="Microwave_Certificate_{cal_id}.docx"'
        
        # บันทึกเอกสาร
        doc.save(response)
        print("DEBUG: Creating response for Microwave certificate")
        print("DEBUG: Response created successfully")
        
        return response
        
    except Exception as e:
        print(f"DEBUG: Error in export_microwave_certificate_docx: {e}")
        return HttpResponse(f"Error: {e}", status=500)
